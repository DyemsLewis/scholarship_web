from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(r"D:\XAMPP\htdocs\scholarship_web\deliverables\Scholarship_Providers_and_Applicant_Requirements_Guide.docx")
NAVY = "17324D"
BLUE = "2E74B5"
MID_BLUE = "1F4D78"
SLATE = "5B6573"
LIGHT_BLUE = "E8F1F8"
LIGHT_GRAY = "F4F6F8"


def set_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_paragraph(paragraph, fill, border_color=None):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    if border_color:
        borders = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "18")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), border_color)
        borders.append(left)
        p_pr.append(borders)


def set_cell_free_page_geometry(document):
    for section in document.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.header_distance = Inches(0.35)
        section.footer_distance = Inches(0.35)


def add_page_number(paragraph):
    paragraph.add_run(" | Page ")
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def configure_styles(document):
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, MID_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.widow_control = True

    for name in ["List Bullet", "List Number"]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.widow_control = True


def add_bullets(document, items):
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        if isinstance(item, tuple):
            label, detail = item
            run = paragraph.add_run(label)
            set_font(run, bold=True)
            paragraph.add_run(detail)
        else:
            paragraph.add_run(item)


def add_numbered(document, items):
    for item in items:
        paragraph = document.add_paragraph(style="List Number")
        if isinstance(item, tuple):
            label, detail = item
            run = paragraph.add_run(label)
            set_font(run, bold=True)
            paragraph.add_run(detail)
        else:
            paragraph.add_run(item)


def add_note(document, label, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.right_indent = Inches(0.18)
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.paragraph_format.line_spacing = 1.2
    paragraph.paragraph_format.keep_together = True
    shade_paragraph(paragraph, LIGHT_BLUE, BLUE)
    run = paragraph.add_run(f"{label}: ")
    set_font(run, bold=True, color=NAVY)
    run = paragraph.add_run(text)
    set_font(run, color=NAVY)


def add_labeled_paragraph(document, label, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.keep_together = True
    run = paragraph.add_run(f"{label}: ")
    set_font(run, bold=True, color=NAVY)
    paragraph.add_run(text)


def build_document():
    document = Document()
    set_cell_free_page_geometry(document)
    configure_styles(document)

    section = document.sections[0]
    section.different_first_page_header_footer = True
    header = section.header
    header_paragraph = header.paragraphs[0]
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header_paragraph.add_run("SCHOLARSHIP REFERENCE GUIDE")
    set_font(header_run, size=8.5, color=SLATE, bold=True)

    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_paragraph.add_run("Scholarship Providers and Applicant Requirements")
    set_font(footer_run, size=8.5, color=SLATE)
    add_page_number(footer_paragraph)
    for run in footer_paragraph.runs[1:]:
        set_font(run, size=8.5, color=SLATE)

    kicker = document.add_paragraph()
    kicker.paragraph_format.space_after = Pt(5)
    run = kicker.add_run("PRACTICAL REFERENCE GUIDE")
    set_font(run, size=9, color=BLUE, bold=True)

    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(5)
    title.paragraph_format.keep_with_next = True
    run = title.add_run("Scholarship Providers and Applicant Requirements")
    set_font(run, size=24, color=NAVY, bold=True)

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    run = subtitle.add_run("What providers are, what scholarships contain, what applicants submit, and how the complete process works")
    set_font(run, size=11.5, color=SLATE)

    add_note(
        document,
        "Purpose",
        "This guide explains the common responsibilities of scholarship providers and applicants. Exact requirements may differ by program, education level, funding source, and applicable policy.",
    )

    document.add_heading("1. What Is a Scholarship Provider?", level=1)
    document.add_paragraph(
        "A scholarship provider is an organization or authorized group that funds, manages, or delivers educational assistance. The provider defines who the program is intended for, what support is available, how applications are reviewed, and who receives the benefit."
    )

    document.add_heading("Common provider types", level=2)
    add_bullets(document, [
        ("Government agencies and local government units. ", "They may support learners based on residence, income, education level, sector, or public policy priorities."),
        ("Schools, colleges, universities, and training institutions. ", "They may provide tuition discounts, grants, allowances, or institution-specific assistance."),
        ("Foundations and non-profit organizations. ", "They commonly support financial need, community development, academic merit, or specific learner groups."),
        ("Companies and corporate social responsibility programs. ", "They may fund education, training, internships, equipment, or career pathways related to their mission."),
        ("Community, faith-based, cooperative, and professional groups. ", "They often serve a local area, membership group, profession, or community cause."),
        ("Private donors and family foundations. ", "They may sponsor a small number of learners under clearly documented rules and an accountable selection process."),
    ])

    document.add_heading("Why providers run scholarship programs", level=2)
    add_bullets(document, [
        "Advance an educational, social, community, institutional, or public-service mission.",
        "Direct limited resources toward learners who fit the intended purpose of the fund.",
        "Create a fair and repeatable application and review process.",
        "Keep evidence for sponsors, boards, donors, auditors, and partner organizations.",
        "Measure whether the assistance reached the intended learners and produced useful outcomes.",
        "Build trust by publishing clear benefits, requirements, decisions, and contact information.",
    ])

    document.add_heading("Provider verification", level=2)
    document.add_paragraph(
        "A platform should verify the provider's identity and authority before allowing public program publication. Typical proof may include organization registration, a permit, an authorization letter, institutional identification, an official address, and verified contact details. Verification reduces scams but does not replace the applicant's responsibility to read the program terms."
    )

    document.add_heading("2. What Is a Scholarship?", level=1)
    document.add_paragraph(
        "A scholarship is organized educational support given under stated conditions. It is not limited to a cash award, and it does not automatically guarantee full tuition or approval. A program may be based on financial need, academic merit, talent, leadership, location, course, education level, community membership, disability, or another lawful target."
    )

    document.add_heading("Possible scholarship benefits", level=2)
    add_bullets(document, [
        "Full or partial tuition and school-fee assistance.",
        "One-time or recurring cash grants and living allowances.",
        "Books, uniforms, school supplies, laboratory materials, or learning kits.",
        "Transportation, meals, boarding, housing, internet, or device support.",
        "Examination, certification, training, or admission fees.",
        "Mentoring, tutorials, leadership development, internships, or career support.",
        "Emergency education assistance or support for a specific school term.",
    ])

    document.add_heading("Information every program should state clearly", level=2)
    add_bullets(document, [
        "Provider name, verified contact information, and program title.",
        "Target applicants, education levels, locations, and any strict eligibility limits.",
        "Each benefit, its value or description, frequency, duration, and release method.",
        "Number of available slots, application opening date, deadline, and expected timeline.",
        "Required applicant information and documents.",
        "Screening, examination, interview, and distribution stages that actually apply.",
        "On-site, online, or hybrid arrangements, including separate event addresses when needed.",
        "Renewal conditions, possible recipient obligations, and reasons support may be ended.",
        "Privacy notice, consent terms, correction process, and official support channel.",
    ])

    document.add_heading("Coverage and location rules", level=2)
    add_labeled_paragraph(document, "Required coverage", "Applicants outside the listed residence or service area are not eligible. Providers should use this only when funding or mission rules truly require it.")
    add_labeled_paragraph(document, "Priority coverage", "Applicants in the preferred area may receive priority, but applicants elsewhere may still apply.")
    add_labeled_paragraph(document, "Event location", "This only tells applicants where an exam, interview, orientation, or distribution will occur. It should not be treated as a residence requirement.")
    add_labeled_paragraph(document, "Nationwide or online", "No local residence restriction applies unless another rule is stated.")

    document.add_heading("3. What Providers Need From Applicants", level=1)
    document.add_paragraph(
        "Providers need enough information to confirm identity, eligibility, readiness, and the correct use of limited scholarship resources. They should request only information that is necessary for the stated program and protect sensitive files from unrelated access."
    )

    document.add_heading("Applicant profile information", level=2)
    add_bullets(document, [
        ("Identity and contact. ", "Full name, birthdate, contact number, email, and account manager when a parent or guardian is responsible."),
        ("Education. ", "Current education level, school or learning institution, grade or year level, strand, course, training program, enrollment status, and grading scale."),
        ("Academic record. ", "General average, GWA, GPA, pass/fail result, competency rating, or another scale appropriate to the learner's institution."),
        ("Household context. ", "Income bracket, household size, support needs, and other need-based information only when relevant."),
        ("Location. ", "Barangay, city or municipality, province, and region when coverage or travel matters."),
        ("Guardian details. ", "Guardian name, relationship, contact details, and consent for minors or younger learners."),
        ("Preferences and goals. ", "Desired support, preferred program type, relocation willingness, and education goal for matching and ranking."),
    ])

    document.add_heading("Common supporting documents", level=2)
    add_bullets(document, [
        "School ID or another valid identity document.",
        "Certificate of enrollment, registration form, or admission letter.",
        "Latest report card, transcript of records, grades, or competency result.",
        "Proof of household income, certificate of indigency, or a similar need document.",
        "Birth certificate and parent or guardian ID for younger applicants when required.",
        "Certificate of residency when the scholarship has strict local coverage.",
        "Good moral certificate, recommendation, portfolio, essay, or service record only when relevant.",
        "Program-specific documents that cannot reasonably be prepared in advance.",
    ])

    document.add_heading("Requirements should fit the learner", level=2)
    add_labeled_paragraph(document, "Preschool and elementary", "Focus on guardian information, birth or identity proof, enrollment, school record, residence, and household need. Course or strand information is not appropriate.")
    add_labeled_paragraph(document, "Junior and senior high school", "Use grade level, curriculum or strand, report-card average, enrollment, guardian context, and location when relevant.")
    add_labeled_paragraph(document, "College and university", "Course, year level, enrollment, GWA or GPA scale, transcript, and institution type may be important.")
    add_labeled_paragraph(document, "TVET and vocational", "Use training program, qualification level, competency status, certification readiness, and training-center information.")
    add_labeled_paragraph(document, "ALS and alternative learning", "Use ALS level, learning-center verification, assessment status, location, and support needs rather than college-specific fields.")

    add_note(
        document,
        "Document quality",
        "Files should be readable, current, complete, and consistent with the applicant profile. A provider should request clarification or replacement when a file is unclear instead of immediately treating the applicant as dishonest.",
    )

    document.add_heading("4. What Scholarship Providers Do", level=1)
    document.add_heading("Before applications open", level=2)
    add_bullets(document, [
        "Confirm the funding source, budget, available slots, and authorized decision makers.",
        "Define the program purpose, target applicants, benefits, strict rules, priorities, and exclusions.",
        "Choose only necessary documents and set a realistic deadline.",
        "Prepare a consistent review rubric and declare examination or interview stages in advance.",
        "Publish clear contacts, locations, online links, privacy terms, and recipient obligations.",
        "Submit the program for platform or administrative review before publication.",
    ])

    document.add_heading("While applications are open", level=2)
    add_bullets(document, [
        "Answer applicant questions through an official channel.",
        "Review eligibility, profile information, and required documents consistently.",
        "Use matching or decision-support scores as guidance, not as the final decision.",
        "Request replacement files with a clear reason when evidence is unreadable or incomplete.",
        "Apply the same published rubric and record reviewer actions.",
        "Publish shared exam, interview, and distribution schedules once per program when possible.",
        "Protect applicant data and restrict access to authorized staff with appropriate permissions.",
    ])

    document.add_heading("Decision and selection", level=2)
    add_bullets(document, [
        "The provider makes the final selection, not the platform or matching score.",
        "Approval should consider eligibility, valid documents, rubric results, stage outcomes, available slots, and funding.",
        "Rejection or replacement requests should use understandable reasons and avoid unnecessary sensitive details.",
        "Conflicts of interest should be disclosed, and reviewers should not favor applicants outside the stated process.",
        "Applicants should be notified when their status or required action changes.",
    ])

    document.add_heading("After approval", level=2)
    add_bullets(document, [
        "Confirm all benefits, not only the cash amount, and explain how each will be released.",
        "Provide the distribution date, address or online arrangement, contact person, and required instructions.",
        "Explain any renewal, grade-maintenance, attendance, orientation, reporting, or service conditions before acceptance.",
        "Keep award and distribution records for accountability and permitted reporting.",
        "Monitor continuation or renewal only according to the terms originally disclosed.",
        "Close the application cycle and retain or remove personal data according to the stated policy.",
    ])

    document.add_heading("5. How Applications Are Evaluated", level=1)
    add_numbered(document, [
        ("Eligibility check. ", "Confirm education level, course or strand, year level, grades, income, coverage, and other strict rules."),
        ("Completeness check. ", "Confirm that each required item is present and linked to the correct applicant."),
        ("Document review. ", "Check readability, validity, consistency, and whether replacement is required."),
        ("Structured review. ", "Apply the published rubric for merit, need, fit, document quality, or other lawful criteria."),
        ("Additional stages. ", "Conduct provider-managed exams or interviews when the program declares them."),
        ("Final decision. ", "Approve, reject, waitlist, or move the applicant to the next declared stage based on evidence and available slots."),
        ("Outcome communication. ", "Notify the applicant and provide the next action, schedule, distribution details, or reason where appropriate."),
    ])

    add_note(
        document,
        "Decision support",
        "A matching or DSS score helps organize evidence and explain fit. It must not replace human review, provider judgment, document verification, or the provider's final decision.",
    )

    document.add_heading("Fair and responsible review", level=2)
    add_bullets(document, [
        "Use the same published rules for applicants in the same program.",
        "Separate strict eligibility rules from preferences and logistical information.",
        "Do not introduce hidden requirements after submission.",
        "Allow reasonable correction of profile details and replacement of unclear files.",
        "Provide guardian-appropriate communication for younger applicants.",
        "Use only lawful, relevant criteria and avoid discrimination or personal favoritism.",
        "Keep an audit trail of reviews, status changes, document decisions, and schedules.",
    ])

    document.add_heading("6. Applicant Responsibilities and Rights", level=1)
    document.add_heading("Applicant responsibilities", level=2)
    add_bullets(document, [
        "Provide truthful, current, and consistent profile information.",
        "Use an authorized parent or guardian account manager when the learner is too young to manage the account independently.",
        "Upload readable documents and replace files when a valid reason is provided.",
        "Read benefits, eligibility, coverage, schedules, privacy terms, and possible recipient obligations.",
        "Meet deadlines and attend declared examinations, interviews, orientations, or distributions.",
        "Monitor official platform notifications and the registered email address.",
        "Report important changes such as school, course, address, contact information, or enrollment status.",
        "Keep original documents and do not submit altered, borrowed, or misleading evidence.",
    ])

    document.add_heading("Applicant rights", level=2)
    add_bullets(document, [
        "Know who operates the program and how to contact the provider.",
        "See the benefits, eligibility rules, documents, stages, and deadline before applying.",
        "Understand how personal data and uploaded files will be used.",
        "Correct inaccurate profile information and replace rejected files when permitted.",
        "Receive understandable status updates and next-step instructions.",
        "Report suspicious programs, inappropriate document requests, hidden fees, or platform problems.",
        "Be evaluated under the stated rules, while recognizing that eligibility does not guarantee selection.",
    ])

    document.add_heading("Parents and guardians", level=2)
    document.add_paragraph(
        "A parent or guardian may register and manage the account for a younger learner, but the learner remains the scholarship applicant and beneficiary. The profile should identify who manages the account, and providers should use guardian consent and age-appropriate communication when legally or practically required."
    )

    document.add_heading("7. Complete Platform Process", level=1)
    add_numbered(document, [
        "The provider registers, verifies its email, completes its organization profile, and uploads proof of authority.",
        "An administrator reviews and approves or rejects the provider verification request.",
        "An approved provider creates a scholarship with benefits, applicants, documents, selection stages, schedules, location, and terms.",
        "An administrator reviews the program before it becomes publicly available.",
        "The applicant registers, verifies the email, completes the learner profile, and may submit identity or enrollment proof for verification.",
        "The platform compares the profile with published programs and explains matches, missing details, and blocking rules.",
        "The applicant prepares common documents, selects a scholarship, reviews its details, accepts the terms, and submits an application.",
        "The provider reviews the profile, supporting files, DSS explanation, and rubric, then requests replacement or moves the applicant through declared stages.",
        "Shared exam or interview information is announced to affected applicants; the provider conducts those activities outside the portal when applicable.",
        "The provider makes and records the final approval or rejection decision.",
        "Approved applicants receive the complete benefit package and distribution instructions, followed by any disclosed renewal or recipient obligations.",
    ])

    document.add_heading("8. Common Status Meanings", level=1)
    add_labeled_paragraph(document, "Draft", "The provider is still preparing the scholarship and applicants cannot see it.")
    add_labeled_paragraph(document, "Pending review", "The provider or program is waiting for administrative verification.")
    add_labeled_paragraph(document, "Published", "The program is visible and may accept applications until its deadline or closure.")
    add_labeled_paragraph(document, "Submitted", "The applicant successfully sent an application.")
    add_labeled_paragraph(document, "Under review", "The provider is checking eligibility, profile information, and documents.")
    add_labeled_paragraph(document, "Needs replacement", "A document must be replaced using the provider's stated reason.")
    add_labeled_paragraph(document, "Qualified for exam or interview", "The applicant may proceed to the next declared selection stage.")
    add_labeled_paragraph(document, "Approved", "The applicant was selected, subject to the disclosed terms and release process.")
    add_labeled_paragraph(document, "Rejected or not selected", "The application will not proceed in the current cycle.")
    add_labeled_paragraph(document, "Distribution scheduled", "The benefit-release details are available to the approved applicant.")
    add_labeled_paragraph(document, "Closed", "The provider is no longer accepting new applications.")

    document.add_heading("9. Privacy, Security, and Accountability", level=1)
    add_bullets(document, [
        "Applicants should consent before submitting personal data, documents, or an application.",
        "Providers should collect only information required for the program and limit staff access by role and permission.",
        "Verification proofs should not be exposed to unrelated providers or applications.",
        "Sensitive files should use protected storage, authorized viewing, activity logs, and a retention policy.",
        "Providers should never request an applicant's password or use private links outside approved communication channels without explanation.",
        "Any legitimate applicant fee must be lawful, necessary, and disclosed before application; hidden processing or convenience fees are a warning sign.",
        "Platform logs, status histories, and reviewer records support accountability but do not replace formal legal or financial records where those are required.",
    ])

    document.add_heading("10. Practical Checklists", level=1)
    document.add_heading("Provider checklist before publishing", level=2)
    add_bullets(document, [
        "Verified organization identity and authorized account owner.",
        "Confirmed budget, benefits, frequency, duration, and slots.",
        "Target applicants and strict eligibility rules clearly separated from preferences.",
        "Reasonable documents selected for the correct education level.",
        "Deadline and selection timeline are realistic.",
        "Exam, interview, and distribution stages included only when needed.",
        "Event modes, addresses, online links, and instructions are complete.",
        "Review rubric totals 100 percent and uses relevant criteria.",
        "Privacy, contact, renewal, and recipient terms are understandable.",
        "Program submitted for administrative review before publication.",
    ])

    document.add_heading("Applicant checklist before submitting", level=2)
    add_bullets(document, [
        "Email and contact details are active.",
        "Required profile sections are complete and accurate.",
        "Education level, grading scale, location, and guardian details are correct.",
        "The provider and program appear legitimate and verified.",
        "Benefits, coverage, deadline, stages, and possible obligations are understood.",
        "Every required document is readable, current, and saved under the correct type.",
        "The match explanation has been reviewed, including any strict blocking rule.",
        "Application and document terms have been read and accepted.",
        "The submitted application appears in Applications and notifications are enabled.",
        "Original documents and official provider contact details are retained.",
    ])

    add_note(
        document,
        "Key principle",
        "A responsible scholarship process connects a clearly verified provider, a transparent program, a prepared applicant, consistent human review, and accountable delivery of every promised benefit.",
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()

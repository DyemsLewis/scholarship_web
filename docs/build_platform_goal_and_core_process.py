from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(__file__).with_name("Platform_Goal_and_Core_Process.docx")
NAVY = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
MUTED = RGBColor(89, 102, 122)
TEXT = RGBColor(20, 28, 42)


def set_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_spacing(paragraph, before=0, after=6, line=1.10):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_page_number(paragraph):
    run = paragraph.add_run("Page ")
    set_font(run, 9, MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def add_label(document, label, text):
    label_paragraph = document.add_paragraph()
    set_spacing(label_paragraph, before=1, after=2, line=1.0)
    label_paragraph.paragraph_format.keep_with_next = True
    set_font(label_paragraph.add_run(label), 10, NAVY, True)

    body = document.add_paragraph()
    set_spacing(body, after=7)
    body.paragraph_format.widow_control = True
    set_font(body.add_run(text), 11, TEXT)


def add_section(document, title, purpose, process, value, safeguard):
    heading = document.add_paragraph(style="Heading 1")
    heading.paragraph_format.keep_with_next = True
    heading.add_run(title)
    add_label(document, "Purpose", purpose)
    add_label(document, "How it works", process)
    add_label(document, "Why it matters", value)
    add_label(document, "Process boundary", safeguard)


def configure(document):
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading = document.styles["Heading 1"]
    heading.font.name = "Calibri"
    heading._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    heading._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    heading.font.size = Pt(16)
    heading.font.bold = True
    heading.font.color.rgb = BLUE
    heading.paragraph_format.space_before = Pt(16)
    heading.paragraph_format.space_after = Pt(8)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_font(header.add_run("Scholarship Finder Platform | Goal and Core Process"), 10, MUTED, True)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(footer.add_run("Research and capstone documentation | August 2026 | "), 9, MUTED)
    add_page_number(footer)


def build():
    document = Document()
    configure(document)

    eyebrow = document.add_paragraph()
    set_spacing(eyebrow, after=10)
    set_font(eyebrow.add_run("SCHOLARSHIP FINDER AND ELIGIBILITY PLATFORM"), 11, BLUE, True)

    title = document.add_paragraph()
    set_spacing(title, after=10)
    set_font(title.add_run("Platform Goal and Core Process"), 28, NAVY, True)

    subtitle = document.add_paragraph()
    set_spacing(subtitle, after=18, line=1.2)
    set_font(subtitle.add_run("A clear explanation of what the platform is designed to achieve and how applicants, providers, and administrators move through a responsible scholarship process."), 15, MUTED)

    summary = document.add_paragraph()
    set_spacing(summary, after=20)
    set_font(summary.add_run("Platform summary. "), 11, NAVY, True, True)
    set_font(summary.add_run("The Scholarship Finder and Eligibility Platform helps learners discover appropriate opportunities, prepare a usable profile and documents, and apply through a structured process. It helps providers publish and manage genuine opportunities, while administrators protect the integrity of accounts, programs, and reported concerns. It is a decision-support and process-management platform, not a substitute for a provider's final judgment."), 11, TEXT, italic=True)

    add_section(
        document,
        "Platform goal",
        "The platform exists to make scholarship access clearer, more organized, and more trustworthy for learners across different education levels. It brings scholarship discovery, profile preparation, eligibility guidance, documents, applications, schedules, and updates into one process instead of leaving applicants to search across disconnected posts, messages, and files.",
        "Applicants create an account, may complete their profile when ready, browse available programs, and view a program before starting an application. Providers create and manage scholarship programs with their own benefits, eligibility, requirements, schedules, and review decisions. Administrators verify platform-facing accounts and programs, handle oversight, and respond to platform-wide concerns.",
        "A single process gives applicants a more understandable journey and gives providers a more consistent way to receive and review information. It reduces missed requirements, unclear deadlines, repeated questions, and accidental use of outdated scholarship details.",
        "The platform can recommend and organize, but it does not promise that an applicant will receive a scholarship. The provider remains responsible for the real opportunity, selection criteria, review, and final decision."
    )

    add_section(
        document,
        "Users and their responsibilities",
        "The core process serves three primary user groups: applicants, scholarship providers, and platform administrators. Each group has a different role so that no one person or system has unnecessary access or authority.",
        "Applicants maintain their own profile, upload relevant documents, review program details, submit applications, and follow published updates. Providers manage their organization profile, create programs, set criteria and schedules, review submissions, and communicate official decisions. Administrators oversee account and program verification, user management, reports, and platform rules.",
        "Role separation supports fairness and data protection. An applicant should not see provider-only review information, a provider should see only records needed for its programs, and administrators should use higher access only for platform operations and oversight.",
        "A matching result is guidance, not an automatic acceptance. Staff access should be role-based, and final scholarship decisions must remain accountable to the provider."
    )

    add_section(
        document,
        "Account registration and access",
        "Registration establishes a traceable account before a person can save data, submit a program, or apply. Separate applicant and provider registration paths keep the information appropriate to each role.",
        "A new user enters basic account details and verifies the registered email through a verification code before the account is created. After sign-in, applicants can explore the platform even if their profile is incomplete, while providers can complete organization details and submit proof for verification before publishing programs. Users can manage their account details and receive notifications through the platform.",
        "Email verification reduces avoidable errors, helps prevent duplicate or unreachable accounts, and provides a reliable address for official updates. Allowing applicants to explore before completing a profile keeps discovery accessible while still requiring the needed information before an application is submitted.",
        "The platform should collect only information appropriate to the user role. Passwords remain protected credentials, and account access does not give automatic authority to publish programs or view unrelated applicant records."
    )

    add_section(
        document,
        "Applicant profile and document readiness",
        "The applicant profile helps a learner present the information that scholarship providers commonly need for matching and review. It supports different education levels rather than assuming every applicant is a college student.",
        "Applicants can add personal, contact, education, location, academic, household, guardian, and relevant supporting information over time. Common documents can be prepared in the document area before a specific application. Where academic proof is needed, the form can ask for the grading format, grade level, average or grade point, and a supporting grade record. Profile information can be updated when circumstances change.",
        "A prepared profile prevents repeated form-filling and makes it easier to compare program criteria with the learner's actual situation. It also lets the platform show readiness clearly: an applicant may browse at any time, but the system can identify missing information or required files before submission.",
        "Sensitive evidence should be requested only when it is relevant to a scholarship or verification purpose. Updating a proof can require the related verification to be reviewed again, while separately stored general documents should not be deleted without the applicant choosing to remove them."
    )

    add_section(
        document,
        "Scholarship discovery and eligibility guidance",
        "Scholarship discovery helps applicants find opportunities that are understandable and relevant, not merely a large unfiltered list. Each program can describe its provider, intended learners, benefits, eligibility, document requirements, deadline, coverage, and later-stage activities.",
        "Applicants browse program cards, open a quick information view or full details page, save programs, and start an application only after reviewing the requirements. The decision-support feature compares available program criteria with profile information such as education level, school type, course or strand where relevant, academic measure, location, and income range when the provider uses those criteria. Programs open to all learners should not be treated as a mismatch because a narrow criterion is absent.",
        "Clear guidance helps applicants spend time on programs they are more likely to fit and helps providers receive applications that better match their stated purpose. It also makes visible why a program may need more profile information before a useful guidance result can be shown.",
        "The score or recommendation is not a decision and must not be presented as one. It should explain its basis in plain language, respect provider-defined criteria, and leave the final outcome to the provider's human review."
    )

    add_section(
        document,
        "Application preparation and submission",
        "The application process turns a program's published requirements into a simple, trackable submission. It is designed to separate preparation from provider review so applicants can correct missing items before final submission.",
        "When an applicant starts a program application, the system displays the program-specific checklist and lets the applicant upload or select the required files in context. The application flow checks whether the profile information and documents needed for that program are ready. After review, the applicant submits the application and can return to the application details page to see the current stage, files, provider messages, and published schedules.",
        "This flow reduces the need for the provider to repeatedly explain basic requirements and gives applicants a clearer record of what they sent. It also avoids asking applicants to choose unrelated requirements manually because the program already defines what is needed.",
        "An application is not complete merely because a file exists. The provider may request clarification, accept or reject individual proof, or make a final decision according to the published criteria and its lawful internal process."
    )

    add_section(
        document,
        "Provider program setup and review",
        "Provider program setup gives organizations a structured way to present an opportunity before applicants see it. The program form captures only the details needed to explain the offer, select appropriate applicants, and run any later stage.",
        "A provider creates a program with its title, provider identity, logo where available, target applicants, benefits, eligibility, document requirements, deadline, location or coverage, and relevant schedule information. The provider can describe exams, interviews, orientations, or award distribution in advance when those are known. The platform's review process can require program approval before publication, depending on administrator policy.",
        "A structured program page makes the provider's offer easier to understand and easier to update. It gives applicants a single source of truth and gives the provider a consistent review workspace rather than relying only on posts, messages, and spreadsheets.",
        "Providers should publish only benefits, schedules, and requirements they can reasonably deliver. Program changes should be communicated through the platform, particularly when an active application cycle could be affected."
    )

    add_section(
        document,
        "Review, schedules, and final decisions",
        "The review stage is where a provider checks whether an applicant meets the program's real requirements. It keeps program-level selection with the provider while making each next step understandable to the applicant.",
        "Providers review the applicant profile and submitted files from the application workspace, record a decision, and use clear statuses such as review, approved to next stage, rejected, passed, or not selected when appropriate. When a program uses an exam, interview, orientation, or distribution activity, the provider publishes the shared schedule information for relevant applicants. A provider can mark a schedule complete and then review the applicable applicants for the next decision rather than managing attendance as a separate decision for every person.",
        "Keeping schedules program-based makes large applicant groups more practical to manage. Applicants see the information that affects their own application, while providers retain one consistent place to publish dates, addresses, instructions, and results.",
        "Providers conduct external examinations and interviews under their own procedures. The platform should publish official logistics and outcomes, but it should not claim to administer external exams or automatically promote an applicant without a provider decision."
    )

    add_section(
        document,
        "Notifications, reports, and accountability",
        "Notifications and reports make the process responsive when something changes or something goes wrong. They support communication without turning every page into a long list of alerts.",
        "The system can notify users in the platform and, when email settings are enabled, through email for important events such as verification, account updates, program review outcomes, application decisions, schedule publication, document feedback, and service updates. Applicants can report a platform or program-related concern through a compact reporting tool. Relevant reports are visible to the provider and to administrators, while platform-wide resolution remains an administrator responsibility.",
        "Timely updates reduce missed dates and make the record of communication easier to trace. Issue reporting creates a clear path for concerns instead of forcing applicants to search for individual contacts or use informal messages.",
        "Notifications are informative, not a replacement for the official application record. Providers should avoid placing confidential details in broad messages, and resolution ownership should be clear so a provider cannot close a platform-wide complaint without administrator oversight."
    )

    add_section(
        document,
        "Administrative oversight and platform trust",
        "Administrative oversight protects the shared platform without taking over the provider's scholarship decision. It focuses on account legitimacy, program review, appropriate access, reports, and consistency of platform rules.",
        "Administrators can review provider verification, review programs before publication when required, review applicant account verification where applicable, manage users and staff permissions, inspect relevant audit records, and handle reports that concern the platform or require escalation. Administrators can create staff accounts with role-based permissions, while controls prevent an administrator from suspending their own account through normal user management.",
        "This oversight helps applicants trust that the platform is not only a collection of unverified postings. It also gives providers a clear route for support and creates basic accountability for sensitive actions, notifications, and review outcomes.",
        "Administrative approval of a program is not a guarantee that a provider will fulfill every promise. The provider remains responsible for its scholarship, while the platform maintains reasonable process, privacy, and moderation controls."
    )

    closing = document.add_paragraph(style="Heading 1")
    closing.add_run("Core process in one view")
    add_label(document, "Core flow", "Applicants verify an account, explore programs, prepare relevant information and files, receive eligibility guidance, and submit a program-specific application. Providers verify their organization, publish clear programs, review applications, publish schedules, and record the official outcome. Administrators provide role-based oversight, verification, moderation, and support across the shared platform.")
    add_label(document, "Expected outcome", "The platform makes scholarship access more discoverable and understandable while giving providers an organized, accountable process. Its value comes from clearer information, appropriate guidance, safer handling of applicant data, and consistent communication across the full scholarship journey.")

    document.core_properties.title = "Platform Goal and Core Process"
    document.core_properties.subject = "Goal, roles, and core scholarship platform workflow"
    document.core_properties.author = "Scholarship Finder Platform Project Team"
    document.save(OUTPUT)


if __name__ == "__main__":
    build()
    print(OUTPUT)

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(__file__).with_name("Scholarship_Platform_Process_Justification_and_References.docx")

NAVY = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
MUTED = RGBColor(89, 102, 122)
BLACK = RGBColor(20, 28, 42)


def set_font(run, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.1):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_page_number(paragraph):
    run = paragraph.add_run("Page ")
    set_font(run, size=9, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def add_labeled_paragraph(document, label, text):
    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph, after=7, line=1.15)
    label_run = paragraph.add_run(f"{label}: ")
    set_font(label_run, size=11, color=NAVY, bold=True)
    body_run = paragraph.add_run(text)
    set_font(body_run, size=11, color=BLACK)
    return paragraph


def add_section(document, title, process, justification, reference):
    heading = document.add_paragraph(style="Heading 1")
    heading.paragraph_format.keep_with_next = True
    heading.add_run(title)
    add_labeled_paragraph(document, "Platform process", process)
    add_labeled_paragraph(document, "Justification", justification)
    add_labeled_paragraph(document, "Reference basis", reference)


def configure_document(document):
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, NAVY, 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.1

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(header, after=0)
    header_run = header.add_run("Scholarship Finder Platform | Process Justification and Reference Basis")
    set_font(header_run, size=9, color=MUTED, bold=True)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(footer, after=0)
    footer_run = footer.add_run("Research and capstone documentation | August 2026 | ")
    set_font(footer_run, size=9, color=MUTED)
    add_page_number(footer)


def add_title(document):
    kicker = document.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_spacing(kicker, after=4)
    run = kicker.add_run("SCHOLARSHIP FINDER AND ELIGIBILITY PLATFORM")
    set_font(run, size=10, color=BLUE, bold=True)

    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("Process Justification and Reference Basis")
    set_font(run, size=24, color=NAVY, bold=True)

    subtitle = document.add_paragraph()
    set_paragraph_spacing(subtitle, after=16, line=1.15)
    run = subtitle.add_run(
        "A plain-language explanation of how the platform processes work, why they are structured this way, and the official or professional references that informed them."
    )
    set_font(run, size=12, color=MUTED)

    note = document.add_paragraph()
    note.paragraph_format.space_after = Pt(14)
    note.paragraph_format.line_spacing = 1.1
    note_run = note.add_run(
        "Source boundary. The platform is an independent academic project. DOST-SEI and CHED materials are used as comparative references for common scholarship practices only. They do not make the platform an official DOST or CHED system, and every provider remains responsible for its own rules and final decisions."
    )
    set_font(note_run, size=10.5, color=BLACK, italic=True)


def add_reference_entry(document, title, detail, url):
    paragraph = document.add_paragraph()
    set_paragraph_spacing(paragraph, after=8, line=1.1)
    title_run = paragraph.add_run(title)
    set_font(title_run, size=10.5, color=NAVY, bold=True)
    detail_run = paragraph.add_run(f". {detail} ")
    set_font(detail_run, size=10.5, color=BLACK)
    url_run = paragraph.add_run(url)
    set_font(url_run, size=9.5, color=BLUE)


def build_document():
    document = Document()
    configure_document(document)
    add_title(document)

    intro_heading = document.add_paragraph(style="Heading 1")
    intro_heading.add_run("How to read this document")
    intro = document.add_paragraph()
    set_paragraph_spacing(intro, after=8, line=1.15)
    intro.add_run(
        "The platform combines common scholarship application practices with privacy, security, and human-review safeguards. Each section identifies the implemented process, explains why it exists, and names the source used as a reference. A reference may support the overall pattern rather than every field or screen. Where no external scholarship agency prescribes a particular interface choice, the document states that the feature is a platform design decision."
    )

    add_section(
        document,
        "Account registration and email verification",
        "Applicants and providers create an account with a small set of identity and contact details. The account is activated only after the user verifies a code sent to the registered email address. Password reset and account recovery use the same controlled email channel.",
        "A verified email reduces accidental or false registrations, gives the platform a reliable communication address for deadlines and decisions, and prevents an applicant from starting a workflow using an address they do not control. Delaying activation until verification also protects role-specific pages and improves the integrity of later application records.",
        "The DOST-SEI online scholarship process requires a valid email, sends a verification step, then lets an applicant proceed to eligibility and application modules. OWASP guidance also recommends time-limited, single-use verification and reset mechanisms before account activation or sensitive changes.",
    )

    add_section(
        document,
        "Role separation and access control",
        "The platform separates applicant, provider, and administrator accounts. Applicants can maintain their own profile, documents, saved programs, and applications. Providers manage only their organization, programs, applicants, and service requests. Administrators verify providers and applicants, review programs, manage accounts, and view platform records according to assigned permissions.",
        "Scholarship information is not equally relevant or appropriate for every user. Separating roles prevents a provider from viewing another provider's applicants, prevents applicants from changing program rules, and limits administrative actions to authorized staff. Permissions also make review and accountability clearer when a team has more than one provider or administrator account.",
        "This is a platform governance design supported by the National Privacy Commission requirement to use reasonable organizational and technical safeguards, and by OWASP authentication guidance on protecting sensitive internal accounts and applying access controls around authenticated functions.",
    )

    add_section(
        document,
        "Applicant profile, guardian support, and profile readiness",
        "The applicant profile collects only information needed to describe the learner's education level, school context, academic performance, contact details, location, scholarship goals, and, when appropriate, parent or guardian details. A guardian can support a younger learner's account. The profile is optional for browsing, but profile readiness is shown before an application can be submitted.",
        "The platform is designed for learners across grade levels, not only college students. Education level changes the information that is meaningful for matching. Parent or guardian support is included because younger applicants may need help providing accurate information and receiving notices. Allowing browsing before profile completion reduces pressure, while application readiness protects providers from receiving applications that lack the information needed for a fair review.",
        "CHED scholarship forms demonstrate the use of applicant identity, academic information, and attached requirements for evaluating scholarship eligibility. The National Privacy Commission requires that collection be transparent, proportionate, and tied to a legitimate stated purpose; this is why highly sensitive proof is not demanded during ordinary browsing or basic profile completion.",
    )

    add_section(
        document,
        "Document library and academic proof",
        "Applicants can prepare common school-related documents in a reusable document library. Program-specific requirements appear when the applicant starts a particular application. Academic proof is requested in a format appropriate to the learner's selected education level and grading approach, such as GWA, grade average, or report-card evidence. Updating verification proof does not automatically remove the reusable document copy.",
        "A reusable library avoids repeated uploads for every scholarship and lets applicants prepare before a deadline. Keeping program-specific requirements separate prevents the platform from asking every learner for documents that only some providers need. The design also recognizes that grading systems and school records differ across education levels, so one fixed GWA field is not enough for all applicants.",
        "CHED scholarship forms and current CHED merit guidance use academic qualifications and documentary requirements, while DOST-SEI guidance requires applicants to complete requirements through its application process. The platform adapts the pattern of evidence-based review while leaving each provider in control of the documents it actually needs.",
    )

    add_section(
        document,
        "Provider registration, verification, and program creation",
        "Providers register through a separate organization account. They provide organization details and verification proof, then wait for administrator verification before publishing or managing scholarship programs. The program form separates basic information, target learners, eligibility, document requirements, benefits, location, schedules, and possible commitments into smaller sections.",
        "Providers need a different workflow from applicants because they publish opportunities and handle personal applicant information. Verification helps reduce misleading or duplicate scholarship postings. Dividing the program form into sections lowers cognitive load and makes it less likely that a provider will omit deadline, eligibility, benefit, or contact information that applicants need before they decide to apply.",
        "This is a platform design decision informed by the National Privacy Commission's accountability and security expectations for organizations that process personal information. It also mirrors the way official scholarship calls present eligibility, requirements, benefits, and application instructions as distinct decision points.",
    )

    add_section(
        document,
        "Administrative program review before publication",
        "A provider-created program enters an administrative review queue before it becomes publicly available. Administrators can review the provider's verification state, program details, eligibility rules, document requirements, benefit information, schedule details, and any needed corrections. Programs can be published, returned for changes, or rejected.",
        "Publication review is a quality and trust control. It helps ensure that applicants see opportunities with understandable requirements, a real provider identity, and enough information to make an informed decision. It also gives the platform a way to prevent accidental publication of incomplete or conflicting rules.",
        "The process is a platform governance decision. Its rationale is consistent with the National Privacy Commission emphasis on accountability and with official scholarship calls that publish clear requirements and deadlines before accepting applicants. The provider retains ownership of the scholarship, but the platform manages the quality of information shown through its shared portal.",
    )

    add_section(
        document,
        "Scholarship finder and decision support guidance",
        "The finder shows available programs and uses profile information such as education level, school type, course or strand, academic performance, location, and stated preferences to explain likely fit. The matching score is guidance for the applicant; it does not approve, reject, rank for award, or replace the provider's decision. A program marked open to any learner is not treated as a mismatch merely because the learner has a particular strand or course.",
        "Applicants often face many opportunities with different target groups and requirements. A transparent fit explanation helps them decide where to spend time, identify missing information, and avoid applying to programs that clearly do not match. Keeping the score advisory avoids treating incomplete self-reported information as an automatic decision about a person's eligibility.",
        "DOST-SEI's online flow uses an eligibility module before the full application, showing the value of early fit checks. NIST's AI Risk Management Framework supports clearly defining human roles and oversight when a system provides decision support. The platform therefore keeps providers responsible for final scholarship decisions and makes the reasoning visible to the applicant.",
    )

    add_section(
        document,
        "Starting an application and submitting requirements",
        "The application starts as a short guided workflow. The applicant first confirms the selected program, reviews the requirements and deadlines, then uploads only the required files. The application page continues to show the requirement checklist, uploaded files, progress, and relevant schedule information after submission.",
        "A guided flow reduces the chance that an applicant selects the wrong program or misses a requirement hidden in a long form. Showing the checklist after submission gives the applicant a single place to correct files when a provider requests an update. It also makes the application status understandable without forcing the applicant to remember different instructions across pages.",
        "DOST-SEI's E-Scholarship process uses registration, eligibility, application information, document completion, and submission stages. Official CHED forms also pair academic eligibility with documentary attachments. The platform uses the same staged idea while presenting it in a reusable web and mobile-friendly format.",
    )

    add_section(
        document,
        "Provider applicant review and final decision",
        "Providers review the applicant's submitted application, profile summary, evidence files, document status, and program-specific criteria. The provider makes the final decision to advance, reject, or request clarification. Administrators may verify the authenticity or completeness of platform records, but they do not replace the provider's scholarship decision.",
        "The scholarship provider is the party that knows the funding purpose, available slots, legal obligations, and final selection criteria. Separating verification from award selection prevents the platform from falsely appearing to decide who receives a scholarship. It also provides a clear audit trail when a provider needs to explain what information was considered.",
        "CHED scholarship guidance identifies qualifications, academic standards, and documentary requirements, while official application forms warn against false or withheld information. NIST guidance on decision support reinforces the need for defined human oversight when systems use data-driven recommendations. The platform therefore supports review, but does not automate final approval.",
    )

    add_section(
        document,
        "Exams, interviews, and award distribution",
        "Providers add program-level exam, interview, screening, orientation, or distribution schedules with date, time, mode, venue or address, instructions, and relevant benefit information. Applicants can view eligible schedules in their application details. Providers manage schedule completion and stage results in groups where possible, rather than manually repeating the same action for every applicant.",
        "Many scholarships require a later assessment or orientation after document review. Putting schedule details at program level gives every eligible applicant the same official information and prevents contradictory messages. Showing the full benefit package at distribution avoids reducing a scholarship to cash amount only, because providers may also offer supplies, mentoring, training, transport, or other support.",
        "DOST-SEI materials describe a sequence that includes online application, assessment or examination, and email notification of qualifiers. The platform adapts this staged scholarship pattern but leaves the actual exam and interview administration to the provider, which is appropriate because methods differ by organization and program.",
    )

    add_section(
        document,
        "Notifications, meeting coordination, and provider services",
        "Important actions create platform notifications and can use the configured email channel, including verification, application decisions, document requests, program schedules, provider account verification, and service updates. Optional provider services use a shared workspace for the request brief, milestones, files, updates, and a provider-proposed meeting time that an administrator confirms or declines.",
        "Scholarship deadlines and results are time-sensitive. Delivering updates in the account and through email reduces the risk that an applicant or provider misses a change because they were not actively logged in. A confirmed meeting request avoids treating a provider's preferred time as final before the administrator has accepted it. The shared workspace gives both sides one record of the service rather than relying only on untraceable messages.",
        "DOST-SEI announcements state that applicants register with an email address and receive process-related notices through their registered email. OWASP guidance supports controlled, time-limited email verification and secure account-recovery communications. The meeting workspace is a platform process decision designed for clear coordination and auditability.",
    )

    add_section(
        document,
        "Privacy notice, consent, audit records, and problem reports",
        "The platform presents privacy information and terms in context, records acceptance, limits sensitive proof to the stage where it is needed, and keeps role-based activity records for important administrative actions. Applicants can report a problem, with program-related reports visible to the relevant provider and the administrator, while general platform concerns are handled by the administrator.",
        "Scholarship applications may involve learner, school, contact, household, academic, and document information. Users need a clear explanation of why information is collected, who may review it, and how to raise a concern. Audit records and scoped reports help the platform investigate problems without exposing one user's information to unrelated users.",
        "Republic Act No. 10173 requires transparency, legitimate purpose, and proportionality in personal-data processing. National Privacy Commission consent guidance states that privacy information should be accessible and understandable to the target audience. These principles support the platform's contextual notices, limited collection, role restrictions, and report handling.",
    )

    reference_heading = document.add_paragraph(style="Heading 1")
    reference_heading.add_run("Reference list")
    reference_intro = document.add_paragraph()
    set_paragraph_spacing(reference_intro, after=8, line=1.1)
    reference_intro.add_run(
        "The references below were used as comparative process, privacy, security, or decision-support sources. They should be checked again before public deployment because government calls, forms, and dates can change by academic year."
    )

    references = [
        (
            "DOST Region VII Citizen's Charter, Processing of DOST-SEI Undergraduate Scholarship Application",
            "Official process reference for online registration, email verification, eligibility screening, application ID issuance, and later application completion.",
            "https://region7.dost.gov.ph/wp-content/uploads/2023/11/DOST-7-Citizens-Charter-2023-RP.pdf",
        ),
        (
            "DOST Region VII, 2026 Undergraduate Scholarships Application: Deadline Extended",
            "Official notice showing that registered applicants complete requirements through the E-Scholarship Application System and must review incomplete submissions before the deadline.",
            "https://region7.dost.gov.ph/2026-undergraduate-scholarships-application-deadline-extended/",
        ),
        (
            "DOST CAR, 2026 DOST-SEI Undergraduate Scholarship Examination",
            "Official example of an examination stage and qualifier notification through the email address used in the application system.",
            "https://car.dost.gov.ph/the-2026-dost-science-education-institute-dost-sei-undergraduate-scholarship-examination-was-successfully-held-on-february-21-22-2026/",
        ),
        (
            "Commission on Higher Education, CHED Scholarship Application Form",
            "Official comparative reference for applicant information, academic requirements, documentary attachments, and applicant declarations.",
            "https://ched.gov.ph/wp-content/uploads/CHED-Scholarship-Program-Application-Form.pdf",
        ),
        (
            "Commission on Higher Education, CHED Merit Scholarship Program",
            "Official comparative reference for qualification requirements, ranking or selection criteria, and program-level scholarship guidance.",
            "https://ched.gov.ph/merit-scholarship/",
        ),
        (
            "National Privacy Commission, Republic Act No. 10173: Data Privacy Act of 2012",
            "Primary Philippine legal reference for transparency, legitimate purpose, proportionality, lawful processing, and data-subject rights.",
            "https://privacy.gov.ph/data-privacy-act/",
        ),
        (
            "National Privacy Commission, Circular No. 2023-04: Guidelines on Consent",
            "Reference for accessible, understandable privacy information and proportional collection of personal data.",
            "https://privacy.gov.ph/wp-content/uploads/2023/11/NPC-Circular-No.-2023-04_Guidelines-on-Consent_07Nov2023.pdf",
        ),
        (
            "OWASP Cheat Sheet Series, Email Validation and Verification",
            "Security reference for email ownership verification, time-limited single-use tokens, and safer email-related account flows.",
            "https://cheatsheetseries.owasp.org/cheatsheets/Email_Validation_and_Verification_Cheat_Sheet.html",
        ),
        (
            "OWASP Cheat Sheet Series, Authentication",
            "Security reference for account authentication, consistent error handling, protected sensitive accounts, and authenticated access control.",
            "https://cheatsheetseries.owasp.org/cheatsheets/Authentication_Cheat_Sheet.html",
        ),
        (
            "National Institute of Standards and Technology, AI Risk Management Framework 1.0",
            "Decision-support reference for defined human oversight, clear roles, and documenting the limits of data-driven recommendations.",
            "https://doi.org/10.6028/NIST.AI.100-1",
        ),
    ]

    for title, detail, url in references:
        add_reference_entry(document, title, detail, url)

    document.core_properties.title = "Scholarship Platform Process Justification and Reference Basis"
    document.core_properties.subject = "Process justifications and external reference basis for the scholarship finder platform"
    document.core_properties.author = "Scholarship Finder Platform Project Team"
    document.save(OUTPUT)


if __name__ == "__main__":
    build_document()
    print(OUTPUT)

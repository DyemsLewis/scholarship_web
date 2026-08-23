from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(__file__).with_name("Provider_Guide_and_Reference_Basis.docx")
NAVY = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
MUTED = RGBColor(89, 102, 122)
TEXT = RGBColor(20, 28, 42)


def set_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def paragraph_spacing(paragraph, before=0, after=7, line=1.15):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_page_number(paragraph):
    run = paragraph.add_run("Page ")
    set_font(run, 9, MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def add_labeled_paragraph(document, label, body):
    label_paragraph = document.add_paragraph()
    paragraph_spacing(label_paragraph, before=1, after=2, line=1.0)
    label_paragraph.paragraph_format.keep_with_next = True
    run = label_paragraph.add_run(label)
    set_font(run, 10, NAVY, True)

    body_paragraph = document.add_paragraph()
    paragraph_spacing(body_paragraph, after=7)
    body_paragraph.paragraph_format.keep_together = True
    body_paragraph.paragraph_format.widow_control = True
    run = body_paragraph.add_run(body)
    set_font(run, 11, TEXT)
    return body_paragraph


def add_section(document, title, provider_role, reason, platform_practice, reference):
    heading = document.add_paragraph(style="Heading 1")
    heading.paragraph_format.keep_with_next = True
    heading.add_run(title)
    add_labeled_paragraph(document, "Provider role", provider_role)
    add_labeled_paragraph(document, "Why this matters", reason)
    add_labeled_paragraph(document, "Platform practice", platform_practice)
    add_labeled_paragraph(document, "Reference basis", reference)


def add_reference(document, title, detail, url):
    paragraph = document.add_paragraph()
    paragraph_spacing(paragraph, after=5)
    run = paragraph.add_run(title + ". ")
    set_font(run, 10, NAVY, True)
    run = paragraph.add_run(detail + " ")
    set_font(run, 10, TEXT)
    run = paragraph.add_run(url)
    set_font(run, 10, BLUE)


def configure(document):
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)

    heading = document.styles["Heading 1"]
    heading.font.name = "Calibri"
    heading._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    heading._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    heading.font.size = Pt(19)
    heading.font.bold = True
    heading.font.color.rgb = BLUE
    heading.paragraph_format.space_before = Pt(18)
    heading.paragraph_format.space_after = Pt(8)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run("Scholarship Finder Platform | Provider Guide")
    set_font(run, 10, MUTED, True)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Research and capstone documentation | August 2026 | ")
    set_font(run, 9, MUTED)
    add_page_number(footer)


def build():
    document = Document()
    configure(document)

    eyebrow = document.add_paragraph()
    paragraph_spacing(eyebrow, after=10)
    run = eyebrow.add_run("SCHOLARSHIP FINDER AND ELIGIBILITY PLATFORM")
    set_font(run, 11, BLUE, True)

    title = document.add_paragraph()
    paragraph_spacing(title, after=10)
    run = title.add_run("Provider Guide and Reference Basis")
    set_font(run, 28, NAVY, True)

    subtitle = document.add_paragraph()
    paragraph_spacing(subtitle, after=18, line=1.2)
    run = subtitle.add_run("A practical guide for scholarship providers on using the platform responsibly, publishing clear opportunities, and managing applicant review without replacing provider authority.")
    set_font(run, 15, MUTED)

    scope = document.add_paragraph()
    paragraph_spacing(scope, after=20)
    run = scope.add_run("Scope note. ")
    set_font(run, 11, NAVY, True, True)
    run = scope.add_run("The platform is an independent academic project. DOST-SEI, CHED, National Privacy Commission, OWASP, and NIST materials are comparative references for common scholarship, privacy, security, and decision-support practices. They do not make the platform an official government service.")
    set_font(run, 11, TEXT, italic=True)

    opening = document.add_paragraph(style="Heading 1")
    opening.add_run("The provider's place in the platform")
    add_labeled_paragraph(document, "Provider purpose", "A provider is an organization, institution, foundation, business, community group, or authorized team that offers educational support. The provider controls the scholarship purpose, target learners, available slots, benefits, documents, schedules, criteria, and final award decision.")
    add_labeled_paragraph(document, "What the platform does", "The platform gives providers one managed space to publish opportunities, receive structured applications, review applicant information and evidence, communicate schedules and decisions, and keep a record of the process. It does not take ownership of the scholarship fund or make the final selection on the provider's behalf.")
    add_labeled_paragraph(document, "Reference basis", "CHED and DOST-SEI materials show that scholarship opportunities commonly have defined qualifications, requirements, stages, and responsible administering organizations. The platform adapts that structure for different provider types.")

    add_section(document, "Provider registration and verification", "Providers register through a separate organization account and submit basic organization details and verification proof. Before they can publish a program, an administrator checks that the organization and account are appropriate for the platform.", "Applicants need confidence that a scholarship post comes from a genuine organization. Verification also helps prevent duplicate, misleading, or incomplete provider listings, while allowing staff accounts to use only the permissions assigned to them.", "A provider can begin account setup, manage its profile, upload verification proof when requested, and create staff accounts with role-based permissions. Provider verification is separate from scholarship approval: verification confirms the account for platform access, while program review checks the individual opportunity before it becomes public.", "This workflow is a platform governance decision informed by the National Privacy Commission's expectation for reasonable organizational safeguards and OWASP guidance on protecting role-based accounts.")

    add_section(document, "Creating a clear scholarship program", "The provider describes the program title, short description, target learners, scholarship type, eligibility, requirements, benefit package, deadline, application process, location coverage, and relevant contact or schedule details.", "Applicants should see the information needed to decide whether an opportunity is worth pursuing before they start an application. Separating the form into meaningful sections reduces provider workload and makes it less likely that a key detail, such as a deadline or benefit, is omitted.", "The create and edit program workspace groups core information, eligibility, documents, benefits, process schedules, and possible commitments. Providers select common eligibility choices and can add a custom rule when their objective requires it. The platform displays benefits beyond cash, such as supplies, mentoring, training, or transport support.", "CHED scholarship forms and merit guidance show the practical need to publish qualifications and documentary requirements. Official scholarship calls also distinguish eligibility, requirements, benefits, and application instructions so applicants can assess fit before submitting.")

    add_section(document, "Fair eligibility and document requirements", "Providers set only the eligibility conditions and files that are genuinely relevant to their scholarship purpose. Academic standards may use GWA, percentage average, grade points, report-card evidence, or another grading format that matches the target learner's school context.", "One fixed academic measure is not appropriate for every school, grade level, or education pathway. Requiring excessive or highly sensitive proof before it is needed can also discourage eligible learners and weaken trust in the service.", "The platform lets providers target appropriate education levels and school contexts, use common eligibility selections, and add custom requirements when necessary. Applicants can prepare reusable school-related files and then provide program-specific requirements within the application. The system makes it clear that a matching score is guidance, not an automatic determination of eligibility or award.", "CHED scholarship application materials provide an example of academic qualifications and documentary attachments. The Data Privacy Act and National Privacy Commission consent guidance support proportional, transparent collection connected to a legitimate purpose.")

    add_section(document, "Program review before publication", "Before a provider program is publicly listed, it enters an administrator review queue. The administrator checks that provider verification is current and that the scholarship information is understandable, internally consistent, and suitable for publication.", "A publication review is a quality and trust safeguard. It reduces the risk that an applicant sees a program with a missing deadline, contradictory requirement, unclear award, or unverified provider identity. It does not allow the administrator to take over the provider's selection decisions.", "A program may be published, returned for clarification, or rejected with a reason. Once published, a material change can be reviewed again before it becomes visible to applicants. Providers retain ownership of the scholarship while the platform is responsible for the quality and safety of shared information.", "This is a platform governance practice consistent with accountability and safeguard principles in National Privacy Commission guidance. The distinction between platform verification and provider selection is also consistent with the need for a clear accountable decision-maker.")

    add_section(document, "Reviewing applications and reaching a decision", "Providers review applications against their published criteria. The review can include the applicant profile, academic context, required documents, program-specific proof, document status, clarification requests, and review notes. The provider decides whether to advance, decline, or request additional information.", "A structured review prevents staff from relying on incomplete memory or scattered messages. It also protects applicants by making the reason for each requested file or stage visible in one place. The provider, not the platform algorithm or administrator, remains responsible for the final scholarship decision.", "The application workspace places the decision area before detailed supporting information so reviewers can see the action they need to take, then inspect the applicant profile, files, history, and schedules as needed. Providers can work through applicant tables and review pages without downloading every file first. The platform records important actions to support accountability.", "CHED scholarship guidance demonstrates that qualifications and documents are reviewed as part of scholarship selection. NIST guidance supports defining human oversight clearly when a system provides decision-support information, which is why matching guidance does not replace provider judgment.")

    add_section(document, "Managing exams, interviews, and award distribution", "When a scholarship includes an assessment, interview, orientation, or award distribution, the provider sets the program-level date, time, mode, address or venue, instructions, and relevant benefit information. Providers then advance or record results for eligible applicant groups according to their own process.", "The same schedule should not be typed separately for every applicant. Program-level scheduling reduces repetitive work and ensures that eligible applicants receive the same official information. It also lets a provider update a schedule without forgetting one applicant.", "The platform supports exam, interview, screening, orientation, and distribution schedules as program details. Applicants see the relevant schedule after reaching the correct stage, and notifications can be sent through their account and configured email channel. Bulk actions are limited to eligible applicants so a provider cannot accidentally advance an unrelated group.", "DOST-SEI materials provide a comparative example of online application followed by scholarship examination and notification of qualifiers. The platform adopts the communication pattern but leaves all examination and interview administration to the provider.")

    add_section(document, "Privacy, data access, and communication", "Providers may view only applicant information and documents connected to their own programs and review responsibilities. They should use that data only for scholarship administration, keep account credentials secure, avoid sharing files outside authorized channels, and communicate only what applicants need to know.", "Scholarship applications can include personal, household, academic, school, and document information. Clear access boundaries protect applicants and help a provider meet its responsibility to use information only for the stated scholarship purpose.", "The platform separates applicant, provider, and administrator accounts, restricts document views by role and program, records significant activity, provides privacy information in context, and lets applicants report a concern. Providers can send decisions, schedules, clarification requests, and service updates through the platform rather than relying only on informal messages.", "Republic Act No. 10173 and National Privacy Commission consent guidance support transparency, legitimate purpose, proportionality, and safeguards. OWASP authentication guidance supports secured accounts and protected access to sensitive functions.")

    add_section(document, "Provider services and support", "Provider services are a shared workspace for an organization and the platform support team. A service can include a request brief, assigned contact, milestones, target dates, shared files, updates, feedback, and an optional meeting request that an administrator confirms or declines.", "A provider may need help with onboarding, program preparation, data quality, process guidance, or an approved platform service. Keeping those conversations and files in one workspace makes the support process traceable and reduces misunderstandings that can happen when information is spread across personal messages.", "The service page separates the provider request, milestones, team files, platform deliverables, service history, and updates. Providers can propose a meeting time, but it becomes final only after an administrator confirms it. This protects both parties from treating a preferred time as a confirmed appointment.", "The service workspace is a platform design decision for transparent coordination and accountability. Its access and record-keeping approach is informed by privacy safeguards and role-based access principles.")

    ref_heading = document.add_paragraph(style="Heading 1")
    ref_heading.add_run("Reference list")
    reference_intro = document.add_paragraph()
    paragraph_spacing(reference_intro, after=8)
    run = reference_intro.add_run("These references explain the common scholarship, privacy, security, and decision-support patterns used as comparative guidance. They should be reviewed again before public deployment because government forms, scholarship calls, and guidance can change.")
    set_font(run, 11, TEXT)

    references = [
        ("DOST Region VII Citizen's Charter, Processing of DOST-SEI Undergraduate Scholarship Application", "Official comparative source for online registration, email verification, eligibility screening, application ID issuance, and application completion.", "https://region7.dost.gov.ph/wp-content/uploads/2023/11/DOST-7-Citizens-Charter-2023-RP.pdf"),
        ("DOST CAR, 2026 DOST-SEI Undergraduate Scholarship Examination", "Official comparative source for scholarship examination stages and qualifier notification through the registered email address.", "https://car.dost.gov.ph/the-2026-dost-science-education-institute-dost-sei-undergraduate-scholarship-examination-was-successfully-held-on-february-21-22-2026/"),
        ("Commission on Higher Education, CHED Scholarship Application Form", "Official comparative source for applicant information, academic qualifications, documentary attachments, and declarations.", "https://ched.gov.ph/wp-content/uploads/CHED-Scholarship-Program-Application-Form.pdf"),
        ("Commission on Higher Education, CHED Merit Scholarship Program", "Official comparative source for qualification requirements and scholarship program guidance.", "https://ched.gov.ph/merit-scholarship/"),
        ("National Privacy Commission, Republic Act No. 10173: Data Privacy Act of 2012", "Primary Philippine legal source for transparency, legitimate purpose, proportionality, safeguards, and data-subject rights.", "https://privacy.gov.ph/data-privacy-act/"),
        ("National Privacy Commission, Circular No. 2023-04: Guidelines on Consent", "Reference for accessible privacy information and proportional personal-data collection.", "https://privacy.gov.ph/wp-content/uploads/2023/11/NPC-Circular-No.-2023-04_Guidelines-on-Consent_07Nov2023.pdf"),
        ("OWASP Cheat Sheet Series, Authentication", "Security reference for protected accounts, authentication, and access control around sensitive functions.", "https://cheatsheetseries.owasp.org/cheatsheets/Authentication_Cheat_Sheet.html"),
        ("National Institute of Standards and Technology, AI Risk Management Framework 1.0", "Reference for clear human oversight and documented limits for decision-support outputs.", "https://doi.org/10.6028/NIST.AI.100-1"),
    ]
    for reference in references:
        add_reference(document, *reference)

    document.core_properties.title = "Provider Guide and Reference Basis"
    document.core_properties.subject = "Provider responsibilities, platform processes, and comparative references"
    document.core_properties.author = "Scholarship Finder Platform Project Team"
    document.save(OUTPUT)


if __name__ == "__main__":
    build()
    print(OUTPUT)

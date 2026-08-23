from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(__file__).with_name("Scholarship_Provider_Information_and_Role.docx")
NAVY = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
MUTED = RGBColor(89, 102, 122)
TEXT = RGBColor(20, 28, 42)


def font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def spacing(paragraph, before=0, after=7, line=1.15):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def page_number(paragraph):
    run = paragraph.add_run("Page ")
    font(run, 9, MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def labeled(document, label, body):
    label_p = document.add_paragraph()
    spacing(label_p, before=1, after=2, line=1.0)
    label_p.paragraph_format.keep_with_next = True
    run = label_p.add_run(label)
    font(run, 10, NAVY, True)

    body_p = document.add_paragraph()
    spacing(body_p, after=7)
    body_p.paragraph_format.keep_together = True
    body_p.paragraph_format.widow_control = True
    run = body_p.add_run(body)
    font(run, 11, TEXT)
    return body_p


def section(document, title, definition, contribution, value, safeguards):
    heading = document.add_paragraph(style="Heading 1")
    heading.paragraph_format.keep_with_next = True
    heading.add_run(title)
    labeled(document, "What this means", definition)
    labeled(document, "What providers contribute", contribution)
    labeled(document, "Why it matters", value)
    labeled(document, "Good practice", safeguards)


def reference(document, title, detail, url):
    paragraph = document.add_paragraph()
    spacing(paragraph, after=5)
    run = paragraph.add_run(title + ". ")
    font(run, 10, NAVY, True)
    run = paragraph.add_run(detail + " ")
    font(run, 10, TEXT)
    run = paragraph.add_run(url)
    font(run, 10, BLUE)


def configure(document):
    sec = document.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.49)
    sec.footer_distance = Inches(0.49)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)

    heading = document.styles["Heading 1"]
    heading.font.name = "Calibri"
    heading._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    heading._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    heading.font.size = Pt(19)
    heading.font.bold = True
    heading.font.color.rgb = BLUE
    heading.paragraph_format.space_before = Pt(18)
    heading.paragraph_format.space_after = Pt(8)

    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run("Scholarship Finder Platform | Provider Information")
    font(run, 10, MUTED, True)

    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Research and capstone documentation | August 2026 | ")
    font(run, 9, MUTED)
    page_number(footer)


def build():
    document = Document()
    configure(document)

    eyebrow = document.add_paragraph()
    spacing(eyebrow, after=10)
    run = eyebrow.add_run("SCHOLARSHIP FINDER AND ELIGIBILITY PLATFORM")
    font(run, 11, BLUE, True)

    title = document.add_paragraph()
    spacing(title, after=10)
    run = title.add_run("Scholarship Provider Information")
    font(run, 28, NAVY, True)

    subtitle = document.add_paragraph()
    spacing(subtitle, after=18, line=1.2)
    run = subtitle.add_run("An explanation of who scholarship providers are, why they invest in learners, the value they create, and the responsibilities that make their support credible.")
    font(run, 15, MUTED)

    scope = document.add_paragraph()
    spacing(scope, after=20)
    run = scope.add_run("Scope note. ")
    font(run, 11, NAVY, True, True)
    run = scope.add_run("This document discusses providers in general and uses DOST-SEI, CHED, Philippine privacy guidance, and international education sources as comparative references. It does not describe any specific organization as a provider in this platform.")
    font(run, 11, TEXT, italic=True)

    overview = document.add_paragraph(style="Heading 1")
    overview.add_run("Who scholarship providers are")
    labeled(document, "Definition", "A scholarship provider is the person, organization, institution, or group that makes educational support available to learners. The provider establishes the scholarship purpose, makes resources available, sets appropriate rules, reviews applicants, and decides how awards or other support will be delivered.")
    labeled(document, "Common provider types", "Providers can include national or local government agencies, schools, universities, foundations, charities, community organizations, religious organizations, businesses, industry associations, professional groups, alumni associations, cooperatives, donors, and employer-led programs. A provider may be large or small; what matters is that it has a clear purpose, capacity to deliver what it promises, and a responsible process for applicants.")
    labeled(document, "Provider authority", "The provider owns its scholarship goals, available slots, selection criteria, funds, benefits, schedules, and final selection decision. A platform can organize the process and protect information, but it should not pretend to decide who deserves an award on behalf of the provider.")

    section(document, "Why providers offer scholarships", "Providers generally offer scholarships because they want to remove educational barriers, invest in people, respond to a community need, strengthen a field of study, support a workforce pipeline, honor a donor's purpose, or help learners who have limited opportunity to continue their education.", "A provider contributes financial assistance, services, access to learning resources, mentoring, academic support, training, connections, or a combination of these. The support can be designed for a broad learner group or for a clearly defined community, school, course, location, or social purpose.", "Scholarships help providers turn a stated mission into a practical outcome. A school can improve access and completion, a company can develop future talent, a community organization can support local mobility, and a foundation can advance the cause it was created to serve. The impact is not limited to the award amount because educational support can influence persistence, skills, confidence, and future opportunity.", "A provider should state its objective honestly, match requirements to that objective, and avoid making claims it cannot support. The selection process should be clear enough that applicants understand what the provider is trying to achieve and what information will be considered.")

    section(document, "What providers gain from supporting learners", "Provider value is not only direct income. A provider can gain measurable social impact, stronger relationships with schools and communities, progress toward its education or corporate-social-responsibility goals, better visibility for a legitimate cause, a future talent network, and evidence that its resources are reaching intended learners.", "The provider contributes a structured opportunity rather than only a public announcement. By publishing clear eligibility, benefits, requirements, and dates, it creates a fairer route for learners to seek support and gives itself a consistent record of applications and outcomes.", "The main return is strategic and mission-based. A community group can demonstrate local benefit, a school can support student success, a business can build relevant skills and reputation, and a donor can see its support aligned with a defined purpose. The platform can reduce administrative friction, but it does not change the fact that the provider's strongest return is the credibility and impact of a well-run scholarship.", "Providers should not overstate their impact. They can track basic program information, applicant demand, completion of review stages, and award delivery, while protecting personal data and avoiding claims that a scholarship alone guarantees employment or academic success.")

    section(document, "Support providers can offer", "Scholarship support can be money, tuition assistance, school supplies, books, transportation, food or accommodation support, devices or internet access, mentoring, tutoring, review classes, training, internship exposure, certification support, examination assistance, or other learning-related help.", "Providers contribute by recognizing that learners face different barriers. For one learner, a cash grant may be enough; for another, supplies, transport, mentoring, or access to a training opportunity can be the difference between participating and dropping out.", "A complete benefit description helps applicants understand the real offer before applying. It also helps providers explain whether a benefit is one-time, recurring, conditional, in kind, or available only after a later stage. This prevents a scholarship from being described as cash only when the actual support package is broader.", "Providers should publish the benefits they can realistically deliver and clarify important limits. If a provider has later contracts, orientations, or renewal conditions, it should explain that these are handled after selection rather than presenting vague obligations to every applicant.")

    section(document, "Provider responsibilities to applicants", "A provider has responsibility for the accuracy and fairness of the opportunity it posts. Applicants need clear eligibility, a realistic deadline, understandable documents, a description of benefits, a consistent review process, timely schedule updates, and respectful communication about decisions or next stages.", "Providers contribute trust by treating applications as more than data. They explain why a document is needed, review it only for the scholarship purpose, request clarification when appropriate, and avoid asking for sensitive evidence that is not relevant to a fair selection decision.", "A reliable provider process reduces applicant anxiety and protects the provider from confusion, complaints, and inconsistent staff actions. It also gives a provider a better chance of selecting applicants who genuinely fit its stated purpose because applicants can decide early whether the program is appropriate for them.", "Providers should use only relevant criteria, keep decisions under accountable human review, document important actions, and communicate changes through official channels. They should never treat a platform matching score as a final decision or share applicant files outside authorized access.")

    section(document, "Provider data responsibilities", "Providers may receive personal information, academic records, household context, and supporting documents from applicants. That information exists to administer a specific scholarship, not for unrelated marketing, public sharing, or uncontrolled reuse.", "Providers contribute safe handling by limiting access to staff who need it, protecting their accounts, reviewing documents inside authorized tools, and keeping applicant information only for an appropriate purpose and period. A smaller provider has the same basic responsibility to respect applicant privacy as a larger institution.", "Responsible data handling strengthens trust in the scholarship itself. Learners and parents or guardians are more likely to participate when the provider is clear about what it collects, why it collects it, who can view it, and how to raise a concern.", "Providers should apply transparency, legitimate purpose, proportionality, and reasonable safeguards. They should use role-based access, avoid requesting highly sensitive proof too early, and coordinate with the platform administrator if a privacy concern or suspected account issue arises.")

    section(document, "How providers work with a scholarship platform", "A scholarship platform gives providers a shared place to turn their process into a clear applicant experience. It can support provider registration, program publication, matching guidance, document collection, application review, schedules, notifications, issue reporting, and a record of actions.", "Providers contribute the real opportunity and the final judgment. The platform contributes structure, repeatable communication, discovery by appropriate learners, access boundaries, and reduced manual administration. Neither side should claim the other side's responsibility: the provider controls the scholarship while the platform protects the process around it.", "The provider benefits most when it uses the platform to present one consistent source of truth instead of relying on scattered social-media posts, personal messages, and separate files. A shared process makes it easier to update information, show what applicants need, review submissions, and preserve an accountable record without requiring the provider to build its own complete system.", "Providers should keep their profile current, review programs before publishing, explain benefits and eligibility plainly, and use the platform's review and communication tools consistently. The platform should remain supportive and transparent rather than adding unnecessary barriers to the provider's own scholarship process.")

    ref_heading = document.add_paragraph(style="Heading 1")
    ref_heading.add_run("Reference list")
    intro = document.add_paragraph()
    spacing(intro, after=8)
    run = intro.add_run("The references below support the scholarship, education-access, privacy, security, and decision-support principles discussed in this document. They are comparative sources and should be checked again before public deployment or policy writing.")
    font(run, 11, TEXT)

    references = [
        ("DOST Region VII Citizen's Charter, Processing of DOST-SEI Undergraduate Scholarship Application", "Official comparative source for a structured scholarship application process, eligibility screening, requirement completion, and administration.", "https://region7.dost.gov.ph/wp-content/uploads/2023/11/DOST-7-Citizens-Charter-2023-RP.pdf"),
        ("Commission on Higher Education, CHED Scholarship Application Form", "Official comparative source for qualifications, academic information, documentary attachments, and scholarship declarations.", "https://ched.gov.ph/wp-content/uploads/CHED-Scholarship-Program-Application-Form.pdf"),
        ("Commission on Higher Education, CHED Merit Scholarship Program", "Official comparative source for qualification requirements and scholarship program guidance.", "https://ched.gov.ph/merit-scholarship/"),
        ("United Nations Sustainable Development Goal 4: Quality Education", "International reference describing the goal of inclusive and equitable quality education and lifelong learning opportunities.", "https://sdgs.un.org/goals/goal4"),
        ("National Privacy Commission, Republic Act No. 10173: Data Privacy Act of 2012", "Primary Philippine legal source for transparency, legitimate purpose, proportionality, safeguards, and data-subject rights.", "https://privacy.gov.ph/data-privacy-act/"),
        ("National Privacy Commission, Circular No. 2023-04: Guidelines on Consent", "Reference for accessible, understandable privacy information and proportional personal-data collection.", "https://privacy.gov.ph/wp-content/uploads/2023/11/NPC-Circular-No.-2023-04_Guidelines-on-Consent_07Nov2023.pdf"),
        ("OWASP Cheat Sheet Series, Authentication", "Security reference for protected accounts and access control around sensitive applicant information.", "https://cheatsheetseries.owasp.org/cheatsheets/Authentication_Cheat_Sheet.html"),
        ("National Institute of Standards and Technology, AI Risk Management Framework 1.0", "Reference for maintaining human oversight and clear limits when systems provide decision-support information.", "https://doi.org/10.6028/NIST.AI.100-1"),
    ]
    for item in references:
        reference(document, *item)

    document.core_properties.title = "Scholarship Provider Information and Role"
    document.core_properties.subject = "Scholarship provider roles, value, responsibilities, and reference basis"
    document.core_properties.author = "Scholarship Finder Platform Project Team"
    document.save(OUTPUT)


if __name__ == "__main__":
    build()
    print(OUTPUT)

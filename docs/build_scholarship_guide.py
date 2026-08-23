from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(__file__).with_name("Scholarship_Guide_and_Reference_Basis.docx")
NAVY = RGBColor(11, 37, 69)
BLUE = RGBColor(46, 116, 181)
MUTED = RGBColor(89, 102, 122)
TEXT = RGBColor(20, 28, 42)


def font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def spacing(paragraph, before=0, after=7, line=1.15):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def page_number(paragraph):
    run = paragraph.add_run("Page ")
    font(run, 9, MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def labeled(document, label, body):
    paragraph = document.add_paragraph()
    spacing(paragraph)
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.widow_control = True
    run = paragraph.add_run(label + ": ")
    font(run, 11, NAVY, True)
    run = paragraph.add_run(body)
    font(run, 11, TEXT)
    return paragraph


def section(document, title, meaning, purpose, practice, reference, page_break_before=False):
    heading = document.add_paragraph(style="Heading 1")
    heading.paragraph_format.keep_with_next = True
    heading.paragraph_format.page_break_before = page_break_before
    heading.add_run(title)
    labeled(document, "What it means", meaning)
    labeled(document, "Why it matters", purpose)
    labeled(document, "Platform practice", practice)
    reference_paragraph = labeled(document, "Reference basis", reference)


def add_reference(document, title, description, url):
    paragraph = document.add_paragraph()
    spacing(paragraph, after=9)
    run = paragraph.add_run(title + ". ")
    font(run, 10.5, NAVY, True)
    run = paragraph.add_run(description + " ")
    font(run, 10.5, TEXT)
    run = paragraph.add_run(url)
    font(run, 10.5, BLUE)


def configure(document):
    sec = document.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.49)
    sec.footer_distance = Inches(0.49)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)

    heading = document.styles["Heading 1"]
    heading.font.name = "Calibri"
    heading._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    heading._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    heading.font.size = Pt(19)
    heading.font.bold = True
    heading.font.color.rgb = BLUE
    heading.paragraph_format.space_before = Pt(18)
    heading.paragraph_format.space_after = Pt(8)

    header = sec.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header.add_run("Scholarship Finder Platform | Scholarship Guide")
    font(run, 10, MUTED, True)

    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Research and capstone documentation | August 2026 | ")
    font(run, 9, MUTED)
    page_number(footer)


def build():
    document = Document()
    configure(document)

    eyebrow = document.add_paragraph()
    spacing(eyebrow, after=10)
    run = eyebrow.add_run("SCHOLARSHIP FINDER AND ELIGIBILITY PLATFORM")
    font(run, 11, BLUE, True)

    title = document.add_paragraph()
    spacing(title, after=10)
    run = title.add_run("Scholarship Guide and Reference Basis")
    font(run, 28, NAVY, True)

    subtitle = document.add_paragraph()
    spacing(subtitle, after=18, line=1.2)
    run = subtitle.add_run("A plain-language guide to scholarships, their common requirements and benefits, and the external sources that informed the platform's scholarship model.")
    font(run, 15, MUTED)

    notice = document.add_paragraph()
    spacing(notice, after=20, line=1.15)
    run = notice.add_run("Scope note. ")
    font(run, 11, NAVY, True, True)
    run = notice.add_run("This guide is for an independent academic platform. DOST-SEI and CHED materials are comparative references for scholarship practices. Individual providers set their own eligibility rules, benefits, schedules, final decisions, and award conditions.")
    font(run, 11, TEXT, italic=True)

    intro = document.add_paragraph(style="Heading 1")
    intro.add_run("Understanding scholarships")
    labeled(document, "Core idea", "A scholarship is financial or non-financial support that helps a learner enter, continue, or complete education. It may be offered by a government agency, school, community organization, foundation, private company, professional group, or individual donor. A scholarship is different from a loan because it is normally not repaid, but it can include conditions that the recipient must meet.")
    labeled(document, "How this platform uses the term", "The platform treats a scholarship or program as a published opportunity with a provider, target learners, eligibility criteria, benefits, requirements, application period, selection process, and possible post-award commitments. This structure helps applicants compare opportunities without pretending that every scholarship has identical rules.")
    labeled(document, "Reference basis", "CHED's scholarship materials present scholarships as structured programs with qualifications, documentary requirements, and selection conditions. DOST-SEI scholarship guidance similarly uses a defined application process, eligibility checks, documents, and later assessment stages.")

    section(document, "Types of scholarships and target learners", "Scholarships may be merit-based, need-based, course or strand-based, place-based, talent-based, advocacy-based, employment-linked, or designed for a particular learner group. Some are open to all eligible learners, while others focus on elementary, junior high, senior high, technical-vocational, college, graduate, or community-based education.", "A specific target makes a scholarship fairer and easier to administer because the provider can state the purpose of its funding and evaluate applications against comparable criteria. At the same time, an overly narrow target can exclude capable learners, so providers should use only restrictions that are genuinely relevant to their objective.", "The provider form records education level, school type, course or strand, location coverage, academic standard, income context, and custom eligibility choices separately. Programs marked open to any learner do not treat a learner's different course or strand as a mismatch merely because the field is not relevant.", "CHED merit guidance demonstrates that scholarship qualifications can be tied to academic and applicant conditions. The National Privacy Commission's proportionality principle supports requesting only the learner characteristics needed for the provider's stated purpose.")

    section(document, "Scholarship benefits beyond cash", "A scholarship benefit can be a cash grant, tuition support, school supplies, books, transportation aid, meal support, internet support, boarding support, mentoring, training, review classes, internship access, examination support, or a combination of these. The full package matters because a learner's barrier to staying in school may not be tuition alone.", "Showing benefits clearly lets applicants understand the real value and limitations of the opportunity before they apply. It also prevents a cash amount from being mistaken as the whole award when the provider offers useful support in kind or through services.", "The platform records an award amount when applicable and also lets a provider describe other benefits. The scholarship details and distribution information display the complete benefit package rather than describing the award as cash only.", "Official scholarship calls commonly state both financial assistance and other conditions or support arrangements. The platform adapts this practice so providers can communicate their own benefit package in a comparable format.")

    section(document, "Eligibility and academic standards", "Eligibility is the set of conditions that determine whether an applicant may be considered. Conditions can include education level, grade level, school enrollment, academic performance, course or strand, location, household circumstances, citizenship where relevant, or provider-specific rules. Academic standards may use GWA, percentage average, grade points, report cards, or another school-appropriate measure.", "A single GWA field is not fair across all school systems and grade levels. Allowing an appropriate grading basis and evidence makes it easier for younger learners and non-college applicants to represent their performance accurately, while still giving providers information needed for screening.", "Providers select common eligibility options and may add a custom condition when justified. Applicants choose the grading approach that matches their school record, then upload academic proof only when it is necessary for an application or verification. The decision-support score explains likely fit but never makes the final decision.", "CHED application forms and merit guidance use academic qualifications and documentary proof. DOST-SEI's eligibility module provides a comparable example of checking fit before a full application. NIST guidance supports keeping human review in control when a system provides decision-support information.")

    section(document, "Requirements, documents, and proof", "Document requirements are the evidence a provider uses to confirm information that affects eligibility or selection. Common examples include proof of enrollment, a report card or grades, an application form, proof of income when relevant, identification, recommendation letters, portfolios, or provider-specific proof.", "Requirements should be connected to a real selection need. Asking for sensitive records too early can reduce trust and creates unnecessary privacy risk. A reusable document library reduces repeat uploads, while program-specific requirements prevent applicants from submitting files that a provider does not actually need.", "Applicants can prepare common school-related files in the document area and then see the exact required files when starting an application. The profile verification flow focuses on school and academic proof instead of requesting highly sensitive documents by default. Providers review the submitted files in context and can request clarification when needed.", "CHED scholarship forms illustrate the use of applicant information and documentary attachments. The Data Privacy Act and National Privacy Commission consent guidance support transparency, legitimate purpose, and proportional collection of personal information.")

    section(document, "Application, review, and selection", "An application is the learner's formal request to be considered for a particular scholarship. It usually includes an eligibility check, required information, evidence files, declarations, and a status that changes as the provider reviews the submission. Selection is the provider's process of considering qualified applicants against its published criteria and available slots.", "A guided application reduces errors and helps learners see missing requirements before submission. Separating platform verification from the provider's award decision protects the provider's authority and avoids presenting the system's matching score or administrative checks as automatic scholarship approval.", "The platform guides the applicant from program selection to requirements and submission, then keeps the checklist, uploaded files, status, and schedule information in one application page. Providers review the applicant profile, proof, program criteria, and any requested clarification before they approve, reject, or advance an applicant. Administrators verify platform and provider records but do not replace the provider's final selection decision.", "The DOST-SEI online flow shows registration, eligibility, application information, document completion, and submission stages. CHED scholarship guidance identifies qualifications and supporting documents. These sources inform the staged pattern, while the final decision remains a provider responsibility.")

    section(document, "Exams, interviews, screening, and schedules", "Some scholarships use additional stages after initial document review, such as an exam, interview, screening activity, orientation, or portfolio review. These stages may be on-site, online, or managed outside the portal, depending on the provider's process.", "Schedule details must be consistent for all eligible applicants. Publishing program-level dates, times, locations, instructions, and benefit details avoids conflicting messages and prevents staff from manually sending the same information to many applicants. A provider should schedule a stage only when it is genuinely part of the program process.", "Providers add exam, interview, screening, orientation, and distribution schedules at the program level. Applicants who have reached the relevant stage can view the schedule in their application details and receive a notification. Providers can record results in grouped workflows instead of having to perform repetitive attendance actions for every applicant.", "DOST-SEI materials provide an official example of a scholarship examination after application processing and of notifying qualifiers through the email used in the system. The platform adapts the staged approach but does not conduct the provider's actual exam or interview.")

    section(document, "Award distribution, commitments, and renewal", "Award distribution is the stage where the provider releases cash, supplies, services, or other stated benefits to selected learners. Some programs may include conditions after award, such as maintaining satisfactory academic standing, attending an orientation, providing progress updates, or following a provider's separate agreement. Renewable scholarships may require a later review of continuing eligibility.", "Applicants need to know what happens after approval, what benefits are included, and whether there may be future obligations. Providers need freedom to communicate their own legitimate conditions without making a generic portal agreement appear to be the final contract for every scholarship.", "The platform displays distribution schedules, locations or addresses, instructions, and the whole benefit package. It presents possible commitments as information and makes clear that detailed terms are handled by the provider with accepted applicants. Renewal is represented as a program condition when relevant, not as an automatic promise of future funding.", "Government scholarship programs commonly publish continuation conditions, documentary checks, and award processes. The platform adopts the communication pattern while leaving legal obligations, award agreements, and renewal decisions to the provider.", page_break_before=True)

    section(document, "Applicant rights and responsibilities", "Applicants have the right to clear information about a program, a reason for collecting their data, access to their own submitted information, and a way to report a concern. They are responsible for providing truthful information, protecting their account, submitting files on time, and communicating with the provider when clarification is requested.", "A fair scholarship process depends on both clear provider rules and responsible applicant participation. Clear rights build trust, while applicant responsibilities protect other applicants and providers from false or incomplete records.", "The platform provides privacy notices, terms acknowledgements, application checklists, status notifications, document update tools, and a report option. It does not guarantee an award, replace provider rules, or allow an applicant to bypass eligibility, document, or decision stages.", "Republic Act No. 10173 provides Philippine privacy principles and data-subject rights. CHED and DOST-SEI materials show the importance of truthful, complete information and requirement completion in scholarship processing.")

    section(document, "Provider responsibilities and platform safeguards", "Providers are responsible for publishing accurate opportunity details, using relevant criteria, handling applicant information carefully, communicating changes promptly, reviewing applications fairly, and honoring the benefits and conditions they publish. Administrators are responsible for verifying provider and program records, applying access controls, and responding to platform concerns.", "A shared scholarship finder is useful only when opportunities are trustworthy. Provider verification and administrative publication review reduce misleading posts, while role-specific access protects applicants from unrelated staff viewing their information.", "The platform separates provider, applicant, and administrator access. A provider account is verified before it can publish programs. A program enters administrative review before it becomes public. Activity records, scoped document access, privacy notices, and reports support accountability without allowing administrators to make provider award decisions.", "National Privacy Commission guidance calls for reasonable organizational and technical safeguards. OWASP authentication guidance supports protected accounts and access control. These references inform the platform safeguards, while the exact review workflow is a platform governance decision.")

    reference_heading = document.add_paragraph(style="Heading 1")
    reference_heading.add_run("Reference list")
    intro = document.add_paragraph()
    spacing(intro, after=8)
    run = intro.add_run("These sources were used to explain common scholarship practices, privacy safeguards, account security, and human-reviewed decision support. They should be checked again before deployment because forms, requirements, and scholarship calls can change.")
    font(run, 11, TEXT)

    references = [
        ("DOST Region VII Citizen's Charter, Processing of DOST-SEI Undergraduate Scholarship Application", "Official comparative source for online registration, email verification, eligibility screening, application ID issuance, and application completion.", "https://region7.dost.gov.ph/wp-content/uploads/2023/11/DOST-7-Citizens-Charter-2023-RP.pdf"),
        ("DOST Region VII, 2026 Undergraduate Scholarships Application: Deadline Extended", "Official notice showing requirement completion and review through an E-Scholarship Application System.", "https://region7.dost.gov.ph/2026-undergraduate-scholarships-application-deadline-extended/"),
        ("DOST CAR, 2026 DOST-SEI Undergraduate Scholarship Examination", "Official comparative source for a scholarship examination stage and qualifier notification through the registered email address.", "https://car.dost.gov.ph/the-2026-dost-science-education-institute-dost-sei-undergraduate-scholarship-examination-was-successfully-held-on-february-21-22-2026/"),
        ("Commission on Higher Education, CHED Scholarship Application Form", "Official comparative source for applicant information, academic qualifications, documentary attachments, and declarations.", "https://ched.gov.ph/wp-content/uploads/CHED-Scholarship-Program-Application-Form.pdf"),
        ("Commission on Higher Education, CHED Merit Scholarship Program", "Official comparative source for qualification requirements and scholarship program guidance.", "https://ched.gov.ph/merit-scholarship/"),
        ("National Privacy Commission, Republic Act No. 10173: Data Privacy Act of 2012", "Primary Philippine legal source for transparency, legitimate purpose, proportionality, safeguards, and data-subject rights.", "https://privacy.gov.ph/data-privacy-act/"),
        ("National Privacy Commission, Circular No. 2023-04: Guidelines on Consent", "Reference for accessible, understandable privacy information and proportional collection of personal data.", "https://privacy.gov.ph/wp-content/uploads/2023/11/NPC-Circular-No.-2023-04_Guidelines-on-Consent_07Nov2023.pdf"),
        ("OWASP Cheat Sheet Series, Authentication", "Security reference for account authentication and access control around authenticated functions.", "https://cheatsheetseries.owasp.org/cheatsheets/Authentication_Cheat_Sheet.html"),
        ("National Institute of Standards and Technology, AI Risk Management Framework 1.0", "Reference for clear human oversight and documented limits for decision-support outputs.", "https://doi.org/10.6028/NIST.AI.100-1"),
    ]
    for item in references:
        add_reference(document, *item)

    document.core_properties.title = "Scholarship Guide and Reference Basis"
    document.core_properties.subject = "Scholarship concepts, processes, responsibilities, and comparative references"
    document.core_properties.author = "Scholarship Finder Platform Project Team"
    document.save(OUTPUT)


if __name__ == "__main__":
    build()
    print(OUTPUT)

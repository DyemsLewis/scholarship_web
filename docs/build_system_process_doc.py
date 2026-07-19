from pathlib import Path
from textwrap import wrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "deliverables" / "Scholarship_Portal_Complete_System_Process.docx"
TMP = ROOT / ".tmp-docx"
FLOW_IMAGE = TMP / "system_process_flow.png"


# compact_reference_guide preset tokens
FONT = "Calibri"
BODY_SIZE = 11
BODY_AFTER = 6
BODY_LINE = 1.25
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "5E6B78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"
GOLD = "7A5A00"
RED = "9B1C1C"
GREEN = "1F3A5F"
CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_run_font(run, size=None, color=None, bold=None, italic=None, name=FONT):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, margins=CELL_MARGIN_DXA):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side in ("top", "start", "bottom", "end"):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(margins[side]))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color="C9D1D9", size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths_dxa, indent_dxa=TABLE_INDENT_DXA):
    assert sum(widths_dxa) == CONTENT_DXA, (widths_dxa, sum(widths_dxa))
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            set_cell_border(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def paragraph_border_bottom(paragraph, color="B7C6D6", size="6", space="3"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color=MUTED)


def add_numbering_definition(doc, ordered=False):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(el.get(qn("w:abstractNumId"))) for el in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(el.get(qn("w:numId"))) for el in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal" if ordered else "bullet")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1." if ordered else "•")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend([tabs, ind, spacing])
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT)
    fonts.set(qn("w:hAnsi"), FONT)
    r_pr.append(fonts)
    level.extend([start, num_fmt, lvl_text, lvl_jc, p_pr, r_pr])
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])


def add_list_item(doc, text, num_id, bold_lead=None):
    p = doc.add_paragraph()
    apply_numbering(p, num_id)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = BODY_LINE
    if bold_lead and text.startswith(bold_lead):
        first = p.add_run(bold_lead)
        set_run_font(first, size=BODY_SIZE, bold=True, color=INK)
        rest = p.add_run(text[len(bold_lead):])
        set_run_font(rest, size=BODY_SIZE)
    else:
        run = p.add_run(text)
        set_run_font(run, size=BODY_SIZE)
    return p


def add_body(doc, text, bold_lead=None, color=None, italic=False, keep=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(BODY_AFTER)
    p.paragraph_format.line_spacing = BODY_LINE
    p.paragraph_format.keep_together = keep
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, size=BODY_SIZE, bold=True, color=color or INK)
        rest = p.add_run(text[len(bold_lead):])
        set_run_font(rest, size=BODY_SIZE, color=color, italic=italic)
    else:
        run = p.add_run(text)
        set_run_font(run, size=BODY_SIZE, color=color, italic=italic)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.keep_together = True
    run = p.add_run(text)
    set_run_font(run, size={1: 16, 2: 13, 3: 12}[level], bold=True,
                 color={1: BLUE, 2: BLUE, 3: DARK_BLUE}[level])
    return p


def add_callout(doc, label, text, tone="info"):
    fill = {"info": CALLOUT, "warning": "FFF7E0", "risk": "FDEEEE", "success": "EDF4F7"}[tone]
    accent = {"info": DARK_BLUE, "warning": GOLD, "risk": RED, "success": GREEN}[tone]
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.keep_together = True
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), accent)
    p_bdr.append(left)
    p_pr.append(p_bdr)
    r1 = p.add_run(f"{label}: ")
    set_run_font(r1, size=10.5, bold=True, color=accent)
    r2 = p.add_run(text)
    set_run_font(r2, size=10.5, color=INK)
    return p


def add_table(doc, headers, rows, widths_dxa, alignments=None):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa)
    set_repeat_table_header(table.rows[0])
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = (alignments[i] if alignments else WD_ALIGN_PARAGRAPH.LEFT)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.05
        run = p.add_run(str(header))
        set_run_font(run, size=9.5, bold=True, color=INK)
    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            p = cells[i].paragraphs[0]
            p.alignment = (alignments[i] if alignments else WD_ALIGN_PARAGRAPH.LEFT)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            run = p.add_run(str(value))
            set_run_font(run, size=9.5, color="263442")
    set_table_geometry(table, widths_dxa)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    return table


def add_page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def build_flow_image(path):
    width, height = 1600, 860
    image = Image.new("RGB", (width, height), f"#{WHITE}")
    draw = ImageDraw.Draw(image)
    font_regular_path = Path(r"C:\Windows\Fonts\calibri.ttf")
    font_bold_path = Path(r"C:\Windows\Fonts\calibrib.ttf")
    fallback = Path(r"C:\Windows\Fonts\arial.ttf")
    regular = ImageFont.truetype(str(font_regular_path if font_regular_path.exists() else fallback), 30)
    bold = ImageFont.truetype(str(font_bold_path if font_bold_path.exists() else fallback), 34)
    small = ImageFont.truetype(str(font_regular_path if font_regular_path.exists() else fallback), 22)

    def box(x, y, w, h, title, subtitle, fill, outline=INK):
        draw.rounded_rectangle((x, y, x + w, y + h), radius=22, fill=f"#{fill}", outline=f"#{outline}", width=4)
        lines = wrap(title, 20)
        ty = y + 20
        for line in lines:
            bbox = draw.textbbox((0, 0), line, font=bold)
            draw.text((x + (w - (bbox[2] - bbox[0])) / 2, ty), line, font=bold, fill=f"#{INK}")
            ty += 42
        max_chars = max(16, int(w / 17))
        for line in wrap(subtitle, max_chars):
            bbox = draw.textbbox((0, 0), line, font=small)
            draw.text((x + (w - (bbox[2] - bbox[0])) / 2, ty + 4), line, font=small, fill=f"#{MUTED}")
            ty += 33

    def arrow(x1, y1, x2, y2, color=BLUE):
        draw.line((x1, y1, x2, y2), fill=f"#{color}", width=7)
        if x2 >= x1:
            pts = [(x2, y2), (x2 - 24, y2 - 14), (x2 - 24, y2 + 14)]
        else:
            pts = [(x2, y2), (x2 + 24, y2 - 14), (x2 + 24, y2 + 14)]
        draw.polygon(pts, fill=f"#{color}")

    title = "Scholarship Portal: End-to-End System Process"
    draw.text((70, 35), title, font=bold, fill=f"#{INK}")
    draw.line((70, 92, 1530, 92), fill=f"#{BLUE}", width=4)

    box(70, 150, 420, 170, "Applicant", "Verify email • complete profile • prepare documents", "EDF4F7")
    box(590, 150, 420, 170, "Provider", "Verify organization • create and manage programs", "FFF7E0")
    box(1110, 150, 420, 170, "Administrator", "Govern accounts • verify actors • publish programs", "F2F4F7")

    box(70, 430, 330, 165, "Discover", "Match, save, and review requirements", "F4F6F9")
    box(455, 430, 330, 165, "Apply", "Eligibility gate, terms, document attachment", "E8EEF5")
    box(840, 430, 330, 165, "Review", "Documents, DSS, rubric, exam or interview", "F4F6F9")
    box(1225, 430, 305, 165, "Outcome", "Award, schedule, distribute, renew", "EDF4F7")

    arrow(400, 512, 455, 512)
    arrow(785, 512, 840, 512)
    arrow(1170, 512, 1225, 512)
    draw.line((280, 320, 280, 430), fill=f"#{BLUE}", width=5)
    draw.line((800, 320, 1005, 430), fill=f"#{GOLD}", width=5)
    draw.line((1320, 320, 1320, 400), fill=f"#{MUTED}", width=5)
    draw.line((1320, 400, 1005, 400), fill=f"#{MUTED}", width=5)
    draw.line((1005, 400, 1005, 430), fill=f"#{MUTED}", width=5)

    draw.rounded_rectangle((70, 680, 1460, 95 + 680), radius=16, fill="#F8FAFC", outline="#C9D1D9", width=3)
    footer = "Shared services: authentication • eligibility matching • private documents • notifications • audit logs • analytics snapshots"
    draw.text((105, 710), footer, font=regular, fill=f"#{DARK_BLUE}")
    path.parent.mkdir(parents=True, exist_ok=True)
    image.save(path, dpi=(180, 180))


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(BODY_SIZE)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(BODY_AFTER)
    normal.paragraph_format.line_spacing = BODY_LINE

    for level, size, color, before, after in [
        (1, 16, BLUE, 18, 10),
        (2, 13, BLUE, 14, 7),
        (3, 12, DARK_BLUE, 10, 5),
    ]:
        style = doc.styles[f"Heading {level}"]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    header = section.header
    hp = header.paragraphs[0]
    hp.paragraph_format.space_after = Pt(4)
    run = hp.add_run("SCHOLARSHIP PORTAL  |  SYSTEM PROCESS GUIDE")
    set_run_font(run, size=8.5, bold=True, color=MUTED)
    paragraph_border_bottom(hp)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.paragraph_format.space_before = Pt(3)
    fp.paragraph_format.tab_stops.add_tab_stop(Inches(6.0))
    left = fp.add_run("Complete System Process  •  July 2026")
    set_run_font(left, size=8.5, color=MUTED)
    fp.add_run("\t")
    page_label = fp.add_run("Page ")
    set_run_font(page_label, size=8.5, color=MUTED)
    add_field(fp, "PAGE")

    doc.core_properties.title = "Scholarship Portal - Complete System Process"
    doc.core_properties.subject = "End-to-end business and technical process guide"
    doc.core_properties.author = "Scholarship Portal Project"
    doc.core_properties.keywords = "scholarship, applicant, provider, administrator, workflow, DSS"


def add_cover(doc):
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(92)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    r = kicker.add_run("SYSTEM PROCESS GUIDE")
    set_run_font(r, size=11, bold=True, color=BLUE)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    r = title.add_run("Scholarship Portal")
    set_run_font(r, size=30, bold=True, color=INK)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(8)
    r = subtitle.add_run("Complete End-to-End System Process")
    set_run_font(r, size=16, bold=True, color=DARK_BLUE)

    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc.paragraph_format.space_after = Pt(54)
    r = desc.add_run("Business workflow, role responsibilities, decision support, data flow, controls, and operations")
    set_run_font(r, size=11, italic=True, color=MUTED)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(5)
    r = meta.add_run("Prepared from the current Laravel, Vue, and Flutter codebase")
    set_run_font(r, size=10, color=MUTED)
    meta2 = doc.add_paragraph()
    meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta2.add_run("Version 1.0  |  July 2026")
    set_run_font(r, size=10, bold=True, color=BLUE)

    add_page_break(doc)


def add_contents(doc, bullet_id):
    add_heading(doc, "Guide map", 1)
    add_body(doc, "This guide follows the system from account creation through scholarship publication, applicant submission, provider review, award distribution, reporting, and background operations.")
    for item in [
        "1. System overview and role model",
        "2. Account registration, authentication, and safeguards",
        "3. Provider verification and scholarship publication",
        "4. Applicant onboarding, discovery, and application submission",
        "5. Document handling, review, and status management",
        "6. Decision-support scoring and provider rubric",
        "7. Administrator governance, notifications, and auditability",
        "8. Mobile process, data model, deployment, and implementation notes",
    ]:
        add_list_item(doc, item, bullet_id)
    add_callout(doc, "Scope", "This document describes behavior currently implemented in the repository. Where the user interface suggests a preferred sequence but the backend permits broader actions, that distinction is stated explicitly.", "info")


def add_overview(doc, bullet_id):
    add_heading(doc, "1. System overview", 1)
    add_body(doc, "The Scholarship Portal coordinates applicants, scholarship providers, and administrators in one platform. Applicants discover and apply for programs; providers publish programs and make selection decisions; administrators govern accounts, verify organizations and applicant proofs, and control program publication.")
    picture = doc.add_picture(str(FLOW_IMAGE), width=Inches(6.45))
    picture._inline.docPr.set("descr", "Flow diagram showing applicants, providers, and administrators moving through scholarship discovery, application, review, and outcome stages.")
    picture._inline.docPr.set("title", "Scholarship Portal end-to-end process")
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)
    r = caption.add_run("Figure 1. Core roles and the scholarship lifecycle")
    set_run_font(r, size=9, italic=True, color=MUTED)

    add_heading(doc, "1.1 Role responsibilities", 2)
    add_table(doc, ["Role", "Primary responsibility", "Main workspace"], [
        ("Applicant", "Build a profile, discover suitable programs, prepare documents, submit applications, and track outcomes.", "/dashboard"),
        ("Provider", "Verify the organization, create programs, review applicants and documents, and record outcomes.", "/provider"),
        ("Administrator", "Manage accounts, verify actors, publish programs, export records, and review audit logs.", "/admin"),
    ], [1650, 5560, 2150], [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER])

    add_heading(doc, "1.2 Application architecture", 2)
    for text in [
        "Web interface: Blade pages mount Vue 3 components and use JSON endpoints for interactive data and actions.",
        "Backend: Laravel 12 validates requests, authorizes role access, applies business rules, persists records, and creates notifications and audit events.",
        "Mobile interface: Flutter supports applicant-only actions through bearer-token APIs under /api/mobile.",
        "Storage: scholarship images are public; applicant, provider, and application documents are stored privately.",
        "Background services: the queue handles notification email and the scheduler creates deadline reminders.",
    ]:
        add_list_item(doc, text, bullet_id)


def add_auth(doc, bullet_id, num_id):
    add_page_break(doc)
    add_heading(doc, "2. Accounts, authentication, and safeguards", 1)
    add_heading(doc, "2.1 Web registration", 2)
    for text in [
        "Validate identity, contact, email, username, password, role, and acceptance of current terms.",
        "Create the user and a role-specific applicant or provider profile.",
        "Hash the password, start a web session, and regenerate the session identifier.",
        "Send an email-verification link and create an in-portal verification reminder.",
        "Record the registration in the activity log and redirect the user to the role dashboard.",
    ]:
        add_list_item(doc, text, num_id)

    add_heading(doc, "2.2 Login and account protection", 2)
    add_table(doc, ["Control", "System behavior"], [
        ("Incorrect credentials", "Reject the attempt and record a failed-login audit event."),
        ("Suspended account", "End the session or reject the mobile token and deny access."),
        ("Forced password reset", "Block login until the one-hour reset process is completed."),
        ("Role access", "Admin and provider routes require dedicated role middleware; applicant controllers verify the applicant role."),
        ("Rate limiting", "Registration, login, password reset, verification, and selected upload/submission endpoints are throttled."),
        ("Terms evidence", "Web registration, application submission, program publishing, and document uploads record acceptance timestamps and terms version."),
    ], [2300, 7060])

    add_heading(doc, "2.3 Password and email recovery", 2)
    add_body(doc, "Password-reset requests always return a neutral response to avoid exposing whether an email exists. When an account is found, a random token is stored as a hash and the reset link expires after one hour. Email verification can be completed through a signed link, resent by the user, or marked complete by an administrator.")
    add_callout(doc, "Current terms version", "The web application records Terms.VERSION = 2026-07-09 for supported acceptance events.", "info")


def add_provider(doc, bullet_id, num_id):
    add_page_break(doc)
    add_heading(doc, "3. Provider process", 1)
    add_heading(doc, "3.1 Organization verification", 2)
    add_body(doc, "A provider may update its organization profile and upload registration evidence, an authorization letter, identification, school or office proof, or another supporting document. Uploading proof notifies administrators. The administrator sets the provider status to pending, approved, or rejected; a rejection requires an explanation.")
    add_callout(doc, "Publishing gate", "A provider must have a verified email and an admin-approved provider profile before it can create, update, or duplicate a scholarship.", "warning")

    add_heading(doc, "3.2 Scholarship definition", 2)
    for text in [
        "Program identity: title, category, description, image, contacts, and location.",
        "Eligibility: education levels, courses or strands, school types, year levels, locations, income preference, and academic requirement.",
        "Offer: award amount, available slots, application mode, deadline, renewal policy, and contract terms.",
        "Requirements: document list and a configurable review rubric whose weights must total 100 percent.",
        "Provider evidence: acceptance of current program terms is stored with the scholarship record.",
    ]:
        add_list_item(doc, text, bullet_id)

    add_heading(doc, "3.3 Program lifecycle", 2)
    add_table(doc, ["Action", "Resulting state", "Administrator involvement"], [
        ("Save new program", "Draft", "None until submitted"),
        ("Submit new program", "Pending review", "Admin is notified"),
        ("Approve review", "Published", "Program becomes visible to applicants"),
        ("Reject review", "Rejected", "Admin must provide a review note"),
        ("Close published program", "Closed", "No further applications accepted"),
        ("Make material edit to published program", "Pending review", "Program returns for admin approval"),
        ("Duplicate existing program", "Draft copy", "Copy must follow the normal review process"),
    ], [3000, 2100, 4260])

    add_heading(doc, "3.4 Provider-managed assessments", 2)
    add_body(doc, "Providers can maintain qualifying-exam or screening-assessment information, including title, type, duration, passing score, delivery mode, venue, instructions, and active status. The provider then uses application statuses to record the candidate's exam progression.")


def add_applicant(doc, bullet_id, num_id):
    add_page_break(doc)
    add_heading(doc, "4. Applicant process", 1)
    add_heading(doc, "4.1 Profile completion", 2)
    add_body(doc, "An applicant must complete all dynamically required profile fields before applying. Requirements vary by education level, grading method, age, and account-management arrangement.")
    add_table(doc, ["Profile condition", "Additional requirement"], [
        ("All applicants", "Name, contact, birthdate, education level, institution, year level, income bracket, city, province, and region"),
        ("Senior high school, college, or TVET", "Track, strand, course, or program"),
        ("Numeric grading scale", "GWA or general average"),
        ("Minor, basic-education learner, or guardian-managed account", "Account manager and guardian name, relationship, and contact"),
    ], [3400, 5960])

    add_heading(doc, "4.2 Optional applicant verification", 2)
    add_body(doc, "Applicants may upload up to three school IDs, government IDs, enrollment certificates, birth certificates, or other proofs. Each upload moves the profile verification state to pending and notifies administrators. An administrator may approve or reject the proofs.")
    add_callout(doc, "Important distinction", "Admin approval of the applicant profile is visible to providers but is not currently an application-submission requirement. The enforced gates are email verification, profile completeness, program availability, duplicate prevention, and eligibility.", "warning")

    add_heading(doc, "4.3 Scholarship discovery and matching", 2)
    for text in [
        "Display only published programs whose deadline is today or later, or programs with no deadline.",
        "Order programs by nearest deadline and expose provider, award, requirement, location, and contract information.",
        "Evaluate academic, education-level, course, school-type, year-level, location, income, and prepared-document criteria.",
        "Show preference matching, map links, approximate distance when coordinates exist, and document readiness.",
        "Allow applicants to save and unsave programs for later review.",
    ]:
        add_list_item(doc, text, num_id)

    add_heading(doc, "4.4 Reusable document library", 2)
    add_body(doc, "Applicants can upload prepared documents before applying. A document is keyed by its requirement name, and uploading a replacement removes the previous stored file. Supported formats are PDF, JPG, JPEG, PNG, DOC, and DOCX, with a maximum size of 5 MB.")


def add_submission(doc, bullet_id, num_id):
    add_page_break(doc)
    add_heading(doc, "5. Application submission and review", 1)
    add_heading(doc, "5.1 Submission gates", 2)
    gates = [
        ("1", "Applicant and account", "The user is an active applicant and is not blocked by a required password reset."),
        ("2", "Email", "The applicant's email is verified."),
        ("3", "Profile", "All context-dependent required profile fields are complete."),
        ("4", "Program", "The scholarship is published and its deadline has not passed."),
        ("5", "Uniqueness", "The applicant has not already applied for the same scholarship."),
        ("6", "Eligibility", "No explicit non-document criterion has a failing match."),
        ("7", "Terms", "The web applicant accepts current submission terms."),
    ]
    add_table(doc, ["Gate", "Area", "Rule"], gates, [800, 2100, 6460], [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT])
    add_callout(doc, "Eligibility behavior", "An explicit mismatch blocks submission. A missing value is not treated as the same kind of blocker, although many missing values are already caught by the profile-completeness gate. Missing prepared documents affect readiness rather than eligibility.", "info")

    add_heading(doc, "5.2 Records created at submission", 2)
    for text in [
        "Create the scholarship application with submitted status, applicant notes, confirmed checklist, and submission timestamp.",
        "Snapshot the eligibility evaluation and scholarship review rubric.",
        "Create the first application-status history entry.",
        "Copy confirmed reusable documents into private application-document records.",
        "Calculate and persist the decision-support score and its versioned input snapshot.",
        "Record analytics funnel and activity-log events.",
        "Notify the provider and confirm successful submission to the applicant.",
    ]:
        add_list_item(doc, text, num_id)

    add_heading(doc, "5.3 Document review", 2)
    add_table(doc, ["Document state", "Meaning", "Provider note"], [
        ("Pending", "Uploaded and waiting for provider review", "Optional"),
        ("Accepted", "Provider considers the file satisfactory", "Optional"),
        ("Rejected", "File does not satisfy the requirement", "Required"),
        ("Needs replacement", "Applicant should upload a corrected or updated file", "Required"),
    ], [2200, 4700, 2460])
    add_body(doc, "Every changed document decision is logged, recalculates document readiness, and notifies the applicant. Only the applicant, the owning provider, or an administrator may access an application document.")


def add_statuses(doc):
    add_page_break(doc)
    add_heading(doc, "6. Application status lifecycle", 1)
    add_heading(doc, "6.1 Standard review path", 2)
    add_table(doc, ["Stage", "Applicant-facing meaning", "Typical provider action"], [
        ("Submitted", "Application was received", "Begin review"),
        ("Under review", "Eligibility and evidence are being checked", "Review criteria, documents, and notes"),
        ("Qualified", "Initial requirements were met", "Prioritize for shortlist or approval"),
        ("Shortlisted", "Candidate advances to the next screening step", "Arrange follow-up screening"),
        ("Interview", "Interview or follow-up is required", "Record result and next decision"),
        ("Approved", "Application is approved", "Prepare award and distribution details"),
        ("Awarded", "An award outcome is recorded", "Set amount and schedule"),
        ("Distribution scheduled", "A future or current release date is posted", "Provide distribution instructions"),
        ("Distributed", "Scholarship support was released", "Maintain outcome evidence"),
        ("Renewed", "Ongoing support was renewed", "Update renewal outcome"),
    ], [2050, 4150, 3160])

    add_heading(doc, "6.2 Exam-based review path", 2)
    add_body(doc, "An exam-oriented program can progress through Qualified for exam, Exam scheduled, Exam taken, and Exam passed before approval. Exam failed is a closed negative outcome and requires a decision reason.")
    add_heading(doc, "6.3 Negative outcomes", 2)
    add_table(doc, ["Outcome", "Use"], [
        ("Rejected", "Application is closed during eligibility or review."),
        ("Not awarded", "Candidate completed review but was not selected for an award."),
        ("Exam failed", "Candidate did not pass the provider-managed assessment."),
    ], [2600, 6760])
    add_body(doc, "Negative outcomes require a decision reason such as missing documents, unmet academic requirement, outside eligibility, limited funds, not selected, failed exam, or another stated reason.")

    add_heading(doc, "6.4 Enforced transition rules", 2)
    add_callout(doc, "Distribution sequencing", "Distribution may be scheduled only from approved, awarded, or already scheduled. A distribution date cannot be earlier than today. Distributed status is allowed only after scheduling, and never before the scheduled date.", "success")
    add_callout(doc, "Status flexibility", "Most other status values can technically be selected directly by the provider. The displayed timeline communicates the preferred sequence, but the backend does not enforce every intermediate transition.", "warning")
    add_body(doc, "There is no required in-platform applicant acceptance step after an award. Contract signing, confirmation, and reward release are handled directly by the provider; the portal records and displays the relevant terms, amount, schedule, and instructions.")


def add_dss(doc):
    add_page_break(doc)
    add_heading(doc, "7. Decision support and provider rubric", 1)
    add_heading(doc, "7.1 Suitability score", 2)
    add_table(doc, ["DSS criterion", "Weight", "Purpose"], [
        ("Structured eligibility match", "65%", "Measures match against scholarship profile criteria."),
        ("Academic merit", "20%", "Compares the applicant's academic record with the program requirement and grading scale."),
        ("Financial need", "15%", "Prioritizes declared income brackets for assistance-focused programs."),
    ], [3500, 1400, 4460], [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT])
    add_table(doc, ["Score", "Recommendation", "Interpretation"], [
        ("85-100", "Strong match", "High suitability based on current structured data"),
        ("70-84", "Potential match", "Good suitability with items for provider confirmation"),
        ("55-69", "Needs review", "Incomplete or uncertain criteria require manual review"),
        ("Below 55", "Limited match", "Several inputs are weak, missing, or mismatched"),
    ], [1600, 2600, 5160], [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.LEFT])
    add_callout(doc, "Decision boundary", "The DSS supports screening only. It does not automatically approve or reject an applicant, and document readiness and review progress are presented separately from the suitability score.", "info")

    add_heading(doc, "7.2 Recalculation and traceability", 2)
    add_body(doc, "The system recalculates DSS results after application submission, applicant-profile changes, document uploads or deletions, provider document decisions, status changes, and selected data reads. Each distinct input state is hashed and stored as a methodology-versioned snapshot with applicant, scholarship, academic, eligibility, and score information.")

    add_heading(doc, "7.3 Provider review rubric", 2)
    add_table(doc, ["Default criterion", "Weight"], [
        ("Eligibility fit", "35%"),
        ("Academic merit", "25%"),
        ("Financial need", "20%"),
        ("Document quality", "20%"),
    ], [7100, 2260], [WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.CENTER])
    add_body(doc, "A provider may define up to six unique criteria with weights totaling exactly 100 percent. Scores range from 0 to 100. The weighted total promotes consistency but remains advisory and separate from the DSS.")


def add_admin_notifications(doc, bullet_id):
    add_page_break(doc)
    add_heading(doc, "8. Administration, notifications, and auditability", 1)
    add_heading(doc, "8.1 Administrator governance", 2)
    for text in [
        "Create and edit applicant, provider, and administrator accounts; change roles and profile data.",
        "Suspend or reactivate accounts and require a password reset.",
        "Manually verify email or resend verification links.",
        "Approve or reject provider organizations and applicant profile proofs.",
        "Publish or reject submitted scholarship programs.",
        "View platform metrics, recent applications, DSS explanations, pending documents, and deadlines.",
        "Export users and applications to CSV and inspect activity logs with actor, role, IP address, and metadata.",
    ]:
        add_list_item(doc, text, bullet_id)
    add_callout(doc, "Administrative safeguard", "An administrator cannot suspend their own account or perform a role/status change that would remove the last active administrator.", "success")

    add_heading(doc, "8.2 Notification lifecycle", 2)
    add_table(doc, ["Trigger", "Recipient", "Channel"], [
        ("Verification proof uploaded", "Administrator", "Portal and email"),
        ("Provider or applicant verification updated", "Affected user", "Portal and email"),
        ("Scholarship submitted or reviewed", "Administrator or provider", "Portal and email"),
        ("Application submitted", "Provider and applicant", "Portal and email"),
        ("Document or application status changed", "Applicant", "Portal and email"),
        ("Award or distribution updated", "Applicant", "Portal and email"),
        ("Account action", "Affected user", "Portal and email where possible"),
    ], [3420, 2460, 3480])
    add_body(doc, "Users can mark notifications individually or all at once. Most newly created portal notifications are also queued for email; the verification reminder itself is excluded because verification email is handled separately.")

    add_heading(doc, "8.3 Deadline reminders", 2)
    add_body(doc, "At 8:00 AM daily, the scheduler checks published scholarships closing in 7, 3, 1, or 0 days. It creates deduplicated reminders for providers and for applicants who saved a program but have not yet applied. Applicants who already applied are excluded from saved-program reminders.")


def add_mobile_data_ops(doc, bullet_id):
    add_page_break(doc)
    add_heading(doc, "9. Mobile, data, and operations", 1)
    add_heading(doc, "9.1 Applicant mobile application", 2)
    add_body(doc, "The Flutter client is restricted to applicant accounts. It uses bearer tokens whose hashes are stored in the database; tokens expire after 30 days and become unusable if the account is suspended or requires a password reset.")
    for text in [
        "Register and log in as an applicant.",
        "Update the applicant profile and view profile readiness.",
        "Browse, match, save, and unsave scholarships.",
        "Upload or delete reusable prepared documents.",
        "Submit applications and track current application information.",
        "Read and mark notifications and log out by deleting the active token.",
    ]:
        add_list_item(doc, text, bullet_id)
    add_callout(doc, "Mobile limitation", "Mobile registration does not currently send an email-verification link, while mobile application submission still requires a verified email. Verification may therefore need to be completed through the web flow or by an administrator.", "risk")

    add_heading(doc, "9.2 Core data relationships", 2)
    add_table(doc, ["Entity", "Key relationships"], [
        ("User", "One role profile; many notifications, bookmarks, documents, funnel events, and DSS snapshots"),
        ("Provider", "Many scholarships, provider verification documents, and assessments"),
        ("Applicant", "Student profile, profile proofs, document library, bookmarks, and applications"),
        ("Scholarship", "Owning provider, applications, bookmarks, criteria, rubric, and analytics events"),
        ("Application", "Applicant, scholarship, documents, status history, rubric scores, DSS snapshots, and outcomes"),
    ], [2300, 7060])

    add_heading(doc, "9.3 Runtime and deployment", 2)
    add_table(doc, ["Component", "Local/production responsibility"], [
        ("Laravel web server", "Serve routes and JSON endpoints; production document root points to public/."),
        ("Database", "SQLite is supported locally; production guidance recommends MySQL and regular backups."),
        ("Vite build", "Compile Vue and Tailwind assets with npm run build."),
        ("Queue worker", "Deliver notification email reliably and retry failures."),
        ("Scheduler", "Run schedule:run every minute so the daily reminder command executes."),
        ("Private storage", "Protect and back up uploaded applicant and provider documents outside the web root."),
        ("Health check", "Use /up to verify that the hosted Laravel application responds."),
    ], [2800, 6560])


def add_security_notes(doc, bullet_id):
    add_page_break(doc)
    add_heading(doc, "10. Security, controls, and current implementation notes", 1)
    add_heading(doc, "10.1 Implemented controls", 2)
    for text in [
        "Passwords and reset tokens are stored as hashes; web sessions regenerate after authentication.",
        "Web requests receive CSRF protection and sensitive routes use authentication and role checks.",
        "Private document access is authorized against administrator role, applicant ownership, or provider ownership.",
        "Input validation constrains file types, sizes, statuses, numeric values, dates, and rubric structure.",
        "Activity logs record important account, application, document, verification, and outcome actions.",
        "Notification and reminder deduplication prevents repeated scheduled messages.",
        "DSS input hashing preserves calculation history without duplicating an identical snapshot.",
    ]:
        add_list_item(doc, text, bullet_id)

    add_heading(doc, "10.2 Current behavior to understand", 2)
    add_table(doc, ["Area", "Current implementation"], [
        ("Account setup page", "The /account/setup page is informational; actual applicant profile editing occurs under /dashboard/profile."),
        ("Applicant profile verification", "Admin approval is displayed but is not required for application submission."),
        ("Prepared documents", "Missing files reduce readiness but do not automatically block an otherwise eligible submission."),
        ("Application status sequence", "Only distribution sequencing has strict transition enforcement; other stages can be selected directly."),
        ("Award acceptance", "No in-platform applicant acceptance is required; providers handle confirmation and contracts directly."),
        ("Contract fields", "Legacy acceptance and student-response columns remain in the data model but current payloads explicitly disable that workflow."),
        ("Mobile terms and verification", "Mobile submission does not record the same terms evidence as the web flow, and mobile registration does not initiate verification email."),
    ], [2500, 6860])

    add_heading(doc, "10.3 Source map", 2)
    add_body(doc, "The following code areas are the primary implementation references for this guide:")
    for text in [
        "routes/web.php and routes/api.php - web and mobile entry points",
        "app/Http/Controllers/AuthController.php - registration, login, verification, and password recovery",
        "app/Http/Controllers/ApplicantDashboardController.php - applicant workflow and application submission",
        "app/Http/Controllers/ProviderController.php - programs, document review, status decisions, and distribution",
        "app/Http/Controllers/AdminController.php - verification, publication, account controls, exports, and logs",
        "app/Services/ScholarshipEligibilityService.php - applicant-program matching and blockers",
        "app/Services/DecisionSupportService.php - suitability score, explanations, progress, and snapshots",
        "app/Console/Commands/SendScholarshipReminders.php - deadline reminders",
    ]:
        add_list_item(doc, text, bullet_id)

    add_callout(doc, "End state", "The complete system process ends when the provider records distribution or renewal, the applicant can see the final outcome, administrators retain exportable and auditable records, and background services continue to deliver notifications and reminders.", "success")


def build_document():
    TMP.mkdir(parents=True, exist_ok=True)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    build_flow_image(FLOW_IMAGE)

    doc = Document()
    configure_document(doc)
    bullet_id = add_numbering_definition(doc, ordered=False)
    number_id = add_numbering_definition(doc, ordered=True)

    add_cover(doc)
    add_contents(doc, bullet_id)
    add_overview(doc, bullet_id)
    add_auth(doc, bullet_id, number_id)
    add_provider(doc, bullet_id, number_id)
    add_applicant(doc, bullet_id, number_id)
    add_submission(doc, bullet_id, number_id)
    add_statuses(doc)
    add_dss(doc)
    add_admin_notifications(doc, bullet_id)
    add_mobile_data_ops(doc, bullet_id)
    add_security_notes(doc, bullet_id)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()

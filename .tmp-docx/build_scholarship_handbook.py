from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"D:\XAMPP\htdocs\scholarship_web")
OUTPUT = ROOT / "deliverables" / "Scholarship_Platform_Complete_Process_Gap_Analysis_and_Improvement_Guide.docx"

# Resolved compact_reference_guide token map.
PAGE_WIDTH_DXA = 12240
PAGE_HEIGHT_DXA = 15840
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_TOP_BOTTOM_DXA = 80
CELL_LEFT_RIGHT_DXA = 120

FONT = "Calibri"
NAVY = "0B172A"
INK = "1F2937"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
GOLD = "B7791F"
MUTED = "64748B"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
ALT_ROW = "F8FAFC"
CALLOUT = "F4F6F9"
PALE_GOLD = "FFF8E6"
PALE_GREEN = "EDF7F1"
PALE_RED = "FDF2F2"
BORDER = "D8E0EA"
WHITE = "FFFFFF"


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_run_font(run, size: float | None = None, color: str | None = None,
                 bold: bool | None = None, italic: bool | None = None,
                 name: str = FONT) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = CELL_TOP_BOTTOM_DXA,
                     start: int = CELL_LEFT_RIGHT_DXA,
                     bottom: int = CELL_TOP_BOTTOM_DXA,
                     end: int = CELL_LEFT_RIGHT_DXA) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = BORDER, size: str = "4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa: Sequence[int]) -> None:
    if sum(widths_dxa) != CONTENT_WIDTH_DXA:
        raise ValueError(f"Table widths must total {CONTENT_WIDTH_DXA}: {widths_dxa}")

    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[idx] / 1440)
            set_cell_margins(cell)


def add_page_field(paragraph, field_name: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = field_name
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    display = OxmlElement("w:t")
    display.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, display, end])
    set_run_font(run, size=8.5, color=MUTED)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    relationship_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.extend([color, underline])
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.extend([run_properties, text_node])
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def configure_numbering(document: Document) -> tuple[int, int]:
    numbering = document.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    next_abstract = max(abstract_ids, default=-1) + 1

    def abstract_num(abstract_id: int, fmt: str, text_value: str) -> None:
        abstract = OxmlElement("w:abstractNum")
        abstract.set(qn("w:abstractNumId"), str(abstract_id))
        multi = OxmlElement("w:multiLevelType")
        multi.set(qn("w:val"), "singleLevel")
        abstract.append(multi)
        level = OxmlElement("w:lvl")
        level.set(qn("w:ilvl"), "0")
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), fmt)
        level_text = OxmlElement("w:lvlText")
        level_text.set(qn("w:val"), text_value)
        justify = OxmlElement("w:lvlJc")
        justify.set(qn("w:val"), "left")
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), "540")
        tabs.append(tab)
        indentation = OxmlElement("w:ind")
        indentation.set(qn("w:left"), "540")
        indentation.set(qn("w:hanging"), "270")
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.extend([tabs, indentation, spacing])
        level.extend([start, num_fmt, level_text, justify, p_pr])
        abstract.append(level)
        numbering.append(abstract)

    abstract_num(next_abstract, "bullet", "\u2022")
    abstract_num(next_abstract + 1, "decimal", "%1.")

    def add_num(abstract_id: int) -> int:
        num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
        num_id = max(num_ids, default=0) + 1
        num = OxmlElement("w:num")
        num.set(qn("w:numId"), str(num_id))
        abstract_ref = OxmlElement("w:abstractNumId")
        abstract_ref.set(qn("w:val"), str(abstract_id))
        num.append(abstract_ref)
        numbering.append(num)
        return num_id

    return add_num(next_abstract), add_num(next_abstract + 1)


def configure_document() -> tuple[Document, int, int]:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.widow_control = True

    heading_specs = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in heading_specs.items():
        style = doc.styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    # Quiet running header without a divider rule.
    header = section.header
    header_p = header.paragraphs[0]
    header_p.paragraph_format.space_after = Pt(0)
    header_p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    left = header_p.add_run("SCHOLARSHIP PLATFORM | COMPREHENSIVE HANDBOOK")
    set_run_font(left, size=8, color=MUTED, bold=True)
    header_p.add_run("\t")
    right = header_p.add_run("AS-BUILT REVIEW | JULY 2026")
    set_run_font(right, size=8, color=MUTED)

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_p.paragraph_format.space_before = Pt(0)
    label = footer_p.add_run("Page ")
    set_run_font(label, size=8.5, color=MUTED)
    add_page_field(footer_p, "PAGE")
    label = footer_p.add_run(" of ")
    set_run_font(label, size=8.5, color=MUTED)
    add_page_field(footer_p, "NUMPAGES")

    doc.core_properties.title = "Scholarship Platform Complete Process, Gap Analysis, and Improvement Guide"
    doc.core_properties.subject = "Platform process, scholarship reference, governance, sustainability, and roadmap"
    doc.core_properties.author = "Scholarship Portal Project Team"
    doc.core_properties.keywords = "scholarship, applicant, provider, admin, DSS, Laravel, Vue, Flutter, Philippines"
    doc.core_properties.comments = "As-built reference prepared 20 July 2026"

    bullet_num, decimal_num = configure_numbering(doc)
    return doc, bullet_num, decimal_num


doc, BULLET_NUM, DECIMAL_NUM = configure_document()


def add_spacer(points: float) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(points)
    p.paragraph_format.line_spacing = 1
    r = p.add_run(" ")
    set_run_font(r, size=1, color=WHITE)


def add_body(text: str, *, bold_lead: str | None = None, italic: bool = False,
             align: WD_ALIGN_PARAGRAPH | None = None, after: float = 6) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    if align is not None:
        p.alignment = align
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, bold=True)
        rest = p.add_run(text[len(bold_lead):])
        set_run_font(rest, italic=italic)
    else:
        r = p.add_run(text)
        set_run_font(r, italic=italic)


def add_labelled(label: str, text: str, *, after: float = 6) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(f"{label}: ")
    set_run_font(r, bold=True, color=NAVY)
    r = p.add_run(text)
    set_run_font(r)


def add_heading(text: str, level: int = 1) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    r = p.add_run(text)
    set_run_font(r, bold=True, color=BLUE if level < 3 else DARK_BLUE,
                 size={1: 16, 2: 13, 3: 12}[level])


def add_kicker(text: str, *, align=WD_ALIGN_PARAGRAPH.LEFT, after: float = 5) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text.upper())
    set_run_font(r, size=9, color=GOLD, bold=True)
    r.font.letter_spacing = Pt(1.2) if hasattr(r.font, "letter_spacing") else None


def apply_num(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])


def new_decimal_num() -> int:
    numbering = doc.part.numbering_part.element
    base_num = next(node for node in numbering.findall(qn("w:num")) if int(node.get(qn("w:numId"))) == DECIMAL_NUM)
    abstract_id = int(base_num.find(qn("w:abstractNumId")).get(qn("w:val")))
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    num_id = max(num_ids, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_bullets(items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph()
        apply_num(p, BULLET_NUM)
        r = p.add_run(item)
        set_run_font(r)


def add_numbered(items: Iterable[tuple[str, str] | str]) -> None:
    num_id = new_decimal_num()
    for item in items:
        p = doc.add_paragraph()
        apply_num(p, num_id)
        if isinstance(item, tuple):
            label, detail = item
            r = p.add_run(f"{label}. ")
            set_run_font(r, bold=True, color=NAVY)
            r = p.add_run(detail)
            set_run_font(r)
        else:
            r = p.add_run(item)
            set_run_font(r)


def add_callout(label: str, text: str, kind: str = "info") -> None:
    fill = {"info": CALLOUT, "success": PALE_GREEN, "warning": PALE_GOLD, "risk": PALE_RED}[kind]
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    set_table_borders(table, color=BORDER, size="5")
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(f"{label.upper()}  ")
    set_run_font(r, size=9, color=GOLD if kind == "warning" else DARK_BLUE, bold=True)
    r = p.add_run(text)
    set_run_font(r, size=10.5, color=INK)
    add_spacer(4)


def add_data_table(headers: Sequence[str], rows: Sequence[Sequence[str]], widths_dxa: Sequence[int],
                   *, font_size: float = 9.2, header_fill: str = LIGHT_BLUE) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths_dxa)
    set_table_borders(table)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    header_row = table.rows[0]
    set_repeat_table_header(header_row)
    prevent_row_split(header_row)
    for idx, text in enumerate(headers):
        cell = header_row.cells[idx]
        shade_cell(cell, header_fill)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(text)
        set_run_font(r, size=font_size, color=NAVY, bold=True)
    for row_idx, values in enumerate(rows):
        row = table.add_row()
        prevent_row_split(row)
        for col_idx, value in enumerate(values):
            cell = row.cells[col_idx]
            if row_idx % 2 == 1:
                shade_cell(cell, ALT_ROW)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.1
            r = p.add_run(str(value))
            set_run_font(r, size=font_size, color=INK)
    set_table_geometry(table, widths_dxa)
    add_spacer(5)


def chapter(title: str, part: str | None = None) -> None:
    doc.add_page_break()
    if part:
        add_kicker(part)
    add_heading(title, 1)


def add_source(label: str, title: str, url: str, note: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run(f"{label}. {title}. ")
    set_run_font(r, bold=True, color=NAVY)
    add_hyperlink(p, url, url)
    r = p.add_run(f" Accessed 20 July 2026. {note}")
    set_run_font(r, size=10, color=MUTED)


# ---------------------------------------------------------------------------
# Cover and front matter
# ---------------------------------------------------------------------------

add_spacer(74)
add_kicker("Platform handbook and capstone reference", align=WD_ALIGN_PARAGRAPH.CENTER, after=20)

cover_title = doc.add_paragraph()
cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover_title.paragraph_format.space_after = Pt(16)
cover_title.paragraph_format.keep_together = True
r = cover_title.add_run(
    "DEVELOPMENT OF A WEB- AND MOBILE-BASED SCHOLARSHIP FINDER AND ELIGIBILITY PLATFORM "
    "FOR FILIPINO STUDENTS WITH A DECISION SUPPORT SYSTEM"
)
set_run_font(r, size=24, color=NAVY, bold=True)

cover_subtitle = doc.add_paragraph()
cover_subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover_subtitle.paragraph_format.space_after = Pt(32)
r = cover_subtitle.add_run("Complete System Process, Scholarship Gap Analysis, and Improvement Guide")
set_run_font(r, size=15, color=DARK_BLUE, bold=True)

cover_line = doc.add_paragraph()
cover_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover_line.paragraph_format.space_after = Pt(8)
r = cover_line.add_run("ACCESS | TRUST | ELIGIBILITY | TRANSPARENCY")
set_run_font(r, size=9.5, color=GOLD, bold=True)

add_spacer(50)
add_body("Prepared from an as-built review of the Laravel, Vue, Tailwind CSS, and Flutter project",
         align=WD_ALIGN_PARAGRAPH.CENTER, after=5)
add_body("System snapshot date: 20 July 2026 | Document version 1.2",
         align=WD_ALIGN_PARAGRAPH.CENTER, after=5)
add_body("For capstone defense, stakeholder orientation, pilot planning, and future development",
         align=WD_ALIGN_PARAGRAPH.CENTER, after=0)

doc.add_page_break()
add_kicker("Document orientation")
add_heading("Purpose, Scope, and How to Use This Handbook", 1)
add_body(
    "This handbook explains the platform's end-to-end process, why each control exists, what applicants, "
    "parents or guardians, scholarship providers, and administrators gain from using it, how scholarship "
    "programs should be designed and evaluated, how the current technical system stores and protects data, "
    "and which improvements should be prioritized after the current implementation."
)
add_body(
    "It is deliberately comprehensive. It can support a capstone manuscript, system demonstration, provider "
    "orientation, applicant help content, administrator procedures, security and hosting preparation, and a "
    "future product roadmap. It should be updated whenever eligibility logic, privacy policy, program workflow, "
    "or production infrastructure changes."
)

add_heading("Evidence Labels", 2)
add_labelled("AS-BUILT", "Functionality confirmed in the current source code, routes, migrations, models, tests, and seed data.")
add_labelled("RECOMMENDED POLICY", "A process or safeguard that should be formally approved before a real pilot or public launch.")
add_labelled("FUTURE FEATURE", "A useful capability that is not represented as part of the current production-ready behavior.")
add_labelled("DEMO DATA", "Fictional accounts and scholarship records used only for testing and presentation.")

add_callout(
    "Important scope note",
    "This is a system and scholarship operations reference, not legal, tax, financial, or regulatory advice. "
    "Actual scholarship rules must be confirmed with the sponsoring organization and current official guidance. "
    "The platform supports selection and communication; it does not itself guarantee an award.",
    "warning",
)

add_heading("System Snapshot", 2)
add_data_table(
    ["Area", "As-built position"],
    [
        ["Web roles", "Applicant/student, scholarship provider, and administrator with role-based workspaces."],
        ["Mobile scope", "Flutter application for applicants only; provider and administrator operations remain web-only."],
        ["Backend", "Laravel 12 on PHP 8.2 with Vue 3, Tailwind CSS 4, Vite, queues, scheduler, and role middleware."],
        ["Database", "SQLite for the current local environment; MySQL is the recommended production database."],
        ["Functional breadth", "131 application routes grouped across public/authentication, applicant, provider, administrator, and mobile API functions."],
        ["Data model", "28 current database tables, including 19 scholarship-domain tables plus Laravel framework tables."],
        ["Automated verification", "73 tests with 591 assertions passed at the audit point; automated tests must still be supplemented by user acceptance testing."],
    ],
    [2500, 6860],
    font_size=9.4,
)

add_heading("Contents at a Glance", 2)
add_numbered([
    ("Executive summary", "The platform's purpose, central value, and immediate priorities."),
    ("Problem, objectives, and scope", "The scholarship-access problem and the boundaries of the proposed solution."),
    ("Stakeholders and value", "Why applicants, parents, providers, administrators, and communities should use it."),
    ("Process design rationale", "Why browsing, verification, reusable documents, DSS guidance, and human decisions are separated."),
    ("Applicant process", "Registration through profile, matching, application, schedules, decision, and award tracking."),
    ("Provider process", "Verification, program creation, review, publication, applicant assessment, and award management."),
    ("Administrator process", "Trust, moderation, user controls, audit logs, and program governance."),
    ("Notifications and mobile", "In-web, email, scheduler, and applicant-only mobile behavior."),
    ("Scholarship knowledge guide", "Definitions, types, criteria, documents, selection, awards, contracts, and warning signs."),
    ("Decision support system", "Method, weights, transparency, limits, fairness, and final-decision responsibility."),
    ("Technical architecture", "Web, mobile, API, maps, queues, storage, and internet dependencies."),
    ("Database and data lifecycle", "Where each category of information is stored and how records relate."),
    ("Security, privacy, and governance", "Current safeguards and the controls still needed before public hosting."),
    ("Hosting, operations, and quality", "Production checklist, backups, monitoring, and testing."),
    ("Sustainability and provider adoption", "Why providers benefit, how the service can remain sustainable, and ethical revenue options."),
    ("Current limitations and gap audit", "What is implemented, partial, missing, deferred, or dependent on policy and operations."),
    ("Future roadmap", "Prioritized improvements from pilot readiness to institutional scale."),
    ("Research and evaluation", "Measures for usability, matching quality, fairness, impact, and provider value."),
])
add_body(
    "Appendices provide a master scholarship-listing template, applicant and provider review checklists, status "
    "definitions, a database dictionary, an access matrix, deployment and testing guides, a risk register, a "
    "glossary, and official references.",
    italic=True,
)


# ---------------------------------------------------------------------------
# Executive summary and platform rationale
# ---------------------------------------------------------------------------

chapter("Executive Summary")
add_body(
    "The proposed platform is a single, structured place where Filipino learners and their parents or guardians "
    "can discover scholarship opportunities, understand whether a program fits their circumstances, prepare "
    "reusable documents, submit applications, receive schedules and decisions, and track progress. Providers can "
    "publish reviewed programs, receive more complete applications, assess evidence consistently, communicate "
    "shared or individual schedules, and preserve an auditable record. Administrators protect the marketplace by "
    "verifying providers and applicants, reviewing programs before publication, managing accounts, and reviewing "
    "logs."
)
add_body(
    "The key design decision is to separate discovery from commitment. A newly registered applicant may explore "
    "the platform before completing every profile field, reducing early abandonment. The system only enforces "
    "profile, document, email, eligibility, deadline, and terms checks when an application is actually submitted. "
    "Likewise, a provider may create a profile and upload proof, but cannot publish programs until the provider is "
    "verified and each submitted program passes administrator review."
)
add_body(
    "The decision support system does not replace provider judgment. It calculates explainable suitability using "
    "eligibility, academic, and financial-need information, while showing document readiness separately. It treats "
    "open criteria such as any course, all strands, or nationwide eligibility as non-restrictive. The provider "
    "remains responsible for the final decision, while snapshots and status histories make the process more "
    "traceable."
)
add_callout(
    "Central recommendation",
    "Use the current system as a controlled pilot platform, not yet as an unrestricted national production service. "
    "Before a public launch, finalize privacy and retention policies, strengthen file security, establish provider "
    "verification service levels, test backups and restoration, and complete accessibility and user acceptance testing.",
    "success",
)

add_heading("What Success Looks Like", 2)
add_bullets([
    "Applicants spend less time searching across unrelated social-media posts and repeatedly preparing the same information.",
    "Parents or guardians can support younger learners without forcing elementary-level applicants into college-specific forms.",
    "Providers receive a more relevant and review-ready applicant pool instead of large volumes of incomplete messages and attachments.",
    "Administrators can explain who changed a status, when a proof was reviewed, why a program was published, and how a recommendation was produced.",
    "Scholarship decisions remain human-led, while technology improves discovery, consistency, communication, and accountability.",
    "The mobile experience expands applicant access without exposing provider and administrator controls on mobile devices.",
])

add_heading("Immediate Priorities Before a Pilot", 2)
add_numbered([
    ("Approve operating policies", "Finalize privacy notice, retention schedule, complaints process, provider proof standards, program-review criteria, and guardian consent rules."),
    ("Harden document handling", "Add malware scanning, stronger file-content validation, upload quarantine, and documented deletion procedures."),
    ("Secure privileged roles", "Add multi-factor authentication for administrators and providers, review permissions, and conduct an account-recovery test."),
    ("Prove operations", "Run backup restoration, queue and scheduler monitoring, mail-delivery testing, and incident-response exercises."),
    ("Validate with people", "Conduct applicant, parent or guardian, provider, and administrator usability tests across different education levels and device sizes."),
])


chapter("1. Problem, Objectives, and Scope", "Part I | Platform Foundation")
add_heading("1.1 The Problem Being Addressed", 2)
add_body(
    "Scholarship information is often fragmented across provider websites, school announcements, social-media "
    "posts, printed notices, and informal referrals. Requirements may be unclear, deadlines may change, and a "
    "learner may only discover a mismatch after spending time collecting documents. Providers face the opposite "
    "problem: applications arrive through inconsistent channels, evidence is difficult to compare, status updates "
    "are repetitive, and there may be limited documentation of how decisions were reached."
)
add_body(
    "The problem is not merely a lack of listings. It is a coordination and decision problem involving discovery, "
    "trust, eligibility interpretation, document readiness, communication, review consistency, and outcome tracking. "
    "A scholarship finder becomes more useful when it can explain fit, preserve evidence, and connect all parties "
    "through a controlled process."
)

add_heading("1.2 General Objective", 2)
add_body(
    "Develop a web- and mobile-based scholarship finder and eligibility platform for Filipino learners that makes "
    "opportunities easier to discover, provides explainable decision support, supports structured applications and "
    "document preparation, and gives verified providers and administrators appropriate tools for review and governance."
)

add_heading("1.3 Specific Objectives", 2)
add_bullets([
    "Provide searchable, structured, and provider-attributed scholarship information.",
    "Support applicants from preschool through elementary, junior high school, senior high school, TVET, ALS, college, and other appropriate learner categories.",
    "Allow applicants to explore before completing the entire profile while preventing incomplete or ineligible application submission.",
    "Reuse common documents across programs while keeping provider-specific requirements visible at application time.",
    "Compute an explainable suitability result without transferring final decision authority from the provider.",
    "Give providers tools to define target applicants, requirements, stages, schedules, review criteria, and award-related obligations.",
    "Require administrator verification of providers and review of programs before public publication.",
    "Maintain notifications, status histories, logs, and calculation snapshots for traceability.",
    "Serve applicants through both web and mobile while reserving sensitive provider and administrator work for the web.",
])

add_heading("1.4 Current Scope", 2)
add_data_table(
    ["Included", "Excluded or handled externally"],
    [
        ["Applicant discovery, profile, saved programs, reusable documents, eligibility matching, application wizard, status tracking, schedules, and notifications.", "The platform does not guarantee acceptance, replace official provider rules, or pay scholarship benefits."],
        ["Provider verification, program drafting, program review submission, applicant review, document decisions, rubric scoring, scheduling, and status management.", "Online examinations are not administered by the platform; providers conduct examinations and interviews using their chosen methods."],
        ["Administrator verification, publication review, account management, logs, exports, and notification oversight.", "National identity verification, bank disbursement, tax advice, legal adjudication, and institutional accreditation are not current functions."],
        ["Applicant-only Flutter mobile app using the same Laravel backend and database.", "Provider and administrator mobile panels are intentionally excluded from the present scope."],
    ],
    [4680, 4680],
    font_size=9.0,
)

add_heading("1.5 Guiding Boundaries", 2)
add_bullets([
    "A match is guidance, not an award decision.",
    "A verified provider badge means the platform reviewed submitted proof under its policy; it is not a government endorsement unless an authorized government registry or agency confirms it.",
    "Only information necessary for a defined purpose should be collected, especially from minors.",
    "Sensitive documents should never become public profile content.",
    "Program terms, funding, award distribution, examination administration, and final selection remain the provider's responsibility.",
])


chapter("2. Stakeholders and Why They Should Use the Platform", "Part I | Platform Foundation")
add_heading("2.1 Stakeholder Map", 2)
add_data_table(
    ["Stakeholder", "Primary need", "Platform value"],
    [
        ["Applicant or learner", "Find suitable, trustworthy opportunities and complete applications efficiently.", "Centralized discovery, clear criteria, explainable matches, reusable files, status tracking, and reminders."],
        ["Parent or guardian", "Support a minor, understand obligations, and protect the learner's information.", "Guardian-aware profile fields, age-appropriate forms, visible requirements, terms, schedules, and controlled documents."],
        ["Scholarship provider", "Reach qualified applicants and administer selection consistently.", "Structured programs, verified identity, review queues, evidence access, rubrics, schedules, notifications, exports, and audit history."],
        ["Administrator", "Protect trust, enforce policy, and resolve operational issues.", "Provider, applicant, and program review; user controls; logs; status histories; and oversight."],
        ["School or community partner", "Guide learners toward relevant support without maintaining separate lists.", "A single referral destination with standardized scholarship details and mobile access."],
        ["Research and project team", "Evaluate scholarship access, usability, and decision-support quality.", "Structured events and calculation snapshots that support ethical aggregate evaluation after governance approval."],
    ],
    [1800, 3000, 4560],
    font_size=8.9,
)

add_heading("2.2 Why Applicants and Families Should Use It", 2)
add_bullets([
    "One profile can support discovery across multiple programs instead of requiring the learner to interpret every announcement from the beginning.",
    "Filters, target-applicant fields, and match explanations make it easier to distinguish a real fit from a merely attractive title.",
    "Common documents can be prepared before an application and attached to matching requirements when the learner applies.",
    "Program details can show award amount, eligibility, location, deadline, requirements, examination or interview stages, renewal terms, and service obligations before the learner commits.",
    "Application status, schedules, reminders, and provider messages are kept in one account rather than scattered across comments, direct messages, and email threads.",
    "A learner may explore first and complete the profile later, while the submission gate prevents an accidental incomplete application.",
])

add_heading("2.3 Why a Provider Should Use It Instead of Only Social Media", 2)
add_body(
    "Social media is useful for awareness, but it is weak as an application operating system. Posts can reach many "
    "people without confirming eligibility, comments mix questions with private information, attachment collection "
    "is inconsistent, updates are easily missed, and selection records are difficult to audit. A provider's own "
    "website offers control, but building and maintaining identity, applications, document review, notifications, "
    "matching, mobile support, and administration separately can be expensive for a small organization."
)
add_bullets([
    "Matched reach: programs are shown to learners whose education level, location, course or strand, income, school type, and academic information are relevant.",
    "Lower intake cost: required fields and reusable document types reduce incomplete email submissions and repeated follow-up.",
    "Better decisions: a complete applicant profile, private proof, program-specific files, rubric scores, and status history appear in one review flow.",
    "Trust: provider proof and administrator review distinguish a controlled listing from an unverified post.",
    "Repeatability: a provider can edit and resubmit programs using a consistent structure rather than designing a new application process each cycle.",
    "Communication: shared and applicant-specific schedules can be announced in the same place applicants already use to track status.",
    "Accountability: actions, document decisions, status reasons, and decision-support snapshots can be traced for internal review.",
    "Impact evidence: aggregate, privacy-respecting funnel and outcome data can eventually help a provider report reach, completion, selection, and geographic coverage.",
])

add_heading("2.4 Why Administrators and Institutions Should Support It", 2)
add_bullets([
    "Verification and publication review create a trust layer that open social-media groups cannot consistently provide.",
    "Role-based access separates applicant, provider, and administrator responsibilities.",
    "Structured program fields improve consistency without forcing every scholarship to use identical criteria.",
    "Logs and histories support complaints review, quality checks, and accountability.",
    "A common platform can reduce duplicated scholarship lists maintained by individual schools or communities.",
])

add_callout(
    "Provider value principle",
    "The provider should pay only when the platform saves real work, improves qualified reach, strengthens trust, or produces useful operating evidence. Charging merely for visibility would be difficult to justify when free social channels already exist.",
    "info",
)


chapter("3. Why the Process Is Designed This Way", "Part I | Platform Foundation")
add_heading("3.1 Progressive Onboarding", 2)
add_body(
    "The platform does not force a new applicant to complete a long profile before seeing any opportunity. This "
    "reduces friction, allows a learner or parent to judge whether the platform is relevant, and avoids collecting "
    "sensitive data before there is a clear purpose. Completion is enforced at the point where the data is actually "
    "needed: application submission."
)

add_heading("3.2 Separate Authentication from Role Profiles", 2)
add_body(
    "The users table stores only account information needed for sign-in and role enforcement. Student, provider, "
    "and administrator details are stored in separate profile tables. This prevents one oversized user record, makes "
    "role-specific validation clearer, and reduces the risk that applicant-only or provider-only fields are exposed "
    "in the wrong context."
)

add_heading("3.3 Trust Before Publication", 2)
add_body(
    "A provider must verify its email, submit organization proof, and receive administrator approval before it can "
    "publish programs. A submitted program then enters a separate review before publication. The two controls answer "
    "different questions: provider verification asks whether the organization and representative are credible; "
    "program review asks whether a particular scholarship listing is complete, understandable, and acceptable."
)

add_heading("3.4 Reusable Documents with Program-Specific Requirements", 2)
add_body(
    "Applicants often need the same records for several scholarships. A private prepared-document library lets the "
    "learner upload common files in advance. When applying, the system checks the selected program's requirements and "
    "attaches matching prepared files. Provider-specific forms remain visible and can be updated. This approach saves "
    "time without pretending that every provider asks for identical proof."
)

add_heading("3.5 Suitability Is Separate from Document Readiness", 2)
add_body(
    "A learner can be an excellent demographic and academic fit while still missing a document, or can have every "
    "document while not satisfying a target criterion. Combining both into one unexplained percentage would be "
    "misleading. The current DSS therefore scores suitability from eligibility, academics, and financial need, while "
    "displaying document readiness and provider review progress separately. Missing required documents block submission "
    "rather than silently reducing the suitability score."
)

add_heading("3.6 Human Final Decision", 2)
add_body(
    "Scholarships may involve essays, interviews, examinations, financial context, limited slots, and provider values. "
    "The system can organize evidence and explain a match, but the provider must make and own the final decision. "
    "Negative decisions require a reason or note so that a status change is not merely a hidden dropdown action."
)

add_heading("3.7 Shared Program Stages and Individual Schedules", 2)
add_body(
    "A provider defines whether a program uses an examination or interview when creating the scholarship. Screening "
    "and distribution remain part of the standard lifecycle. This prevents staff from manually inventing a different "
    "workflow for each applicant. General events can be announced to all relevant applicants, while individual "
    "schedules remain available when a particular person needs a unique time or place."
)

add_heading("3.8 Web for Privileged Operations, Mobile for Applicants", 2)
add_body(
    "Applicants benefit most from mobile access because they need to discover programs, receive reminders, and check "
    "status while away from a computer. Provider and administrator work includes broad document access, verification, "
    "exports, user controls, and program moderation, so those functions remain web-only in the current scope. This "
    "reduces mobile attack surface and keeps the first mobile release focused."
)

add_heading("3.9 Process Gates", 2)
add_data_table(
    ["Gate", "What it protects", "Current rule"],
    [
        ["Email verification", "Account ownership and reliable notification delivery.", "Required before application submission and before a provider can publish."],
        ["Applicant profile completeness", "Eligibility interpretation and review quality.", "Applicant may browse while incomplete; required to submit."],
        ["Required documents", "Evidence readiness.", "Each listed requirement must have an available prepared file at submission."],
        ["Eligibility blockers", "Applicant time and provider review capacity.", "Hard mismatches prevent submission; open criteria do not create a mismatch."],
        ["Provider verification", "Marketplace trust.", "Provider proof must be approved before program creation or publication operations are unlocked."],
        ["Program review", "Listing quality and legitimacy.", "Submitted program remains pending until an administrator publishes or rejects it."],
        ["Deadline and duplicate checks", "Process integrity.", "Closed programs and duplicate applications cannot be submitted."],
        ["Terms acceptance", "Informed responsibility and version tracking.", "Acceptance is recorded for account, application, document, and provider/program contexts."],
    ],
    [1800, 3240, 4320],
    font_size=8.8,
)

add_callout(
    "Current applicant-verification rule",
    "The current application submission code requires a verified email, complete required profile, no eligibility blocker, valid deadline, required documents, and accepted terms. Administrator verification of the applicant's optional profile proof is visible and useful, but is not presently a universal submission gate. If the project wants it to become mandatory, that should be an explicit policy and code change rather than an assumption.",
    "warning",
)


# ---------------------------------------------------------------------------
# Role workflows
# ---------------------------------------------------------------------------

chapter("4. Applicant and Parent or Guardian Process", "Part II | End-to-End Workflows")
add_heading("4.1 Applicant Journey", 2)
add_numbered([
    ("Register", "Create an applicant account using name, email, username, contact number, password, and terms acceptance. Registration returns to the login page without starting an authenticated session."),
    ("Verify email", "Open the verification message or request another link from the account area. Once verified, the resend prompt is removed. Verification proves control of the notification address."),
    ("Explore", "Browse the dashboard, scholarship finder, program details, maps, and saved opportunities even if the profile is not yet complete."),
    ("Build the profile", "Complete identity, education, household, location, preference, and guardian information across separate profile pages. Conditional fields depend on age and education level."),
    ("Upload profile proof", "Optionally provide private evidence for administrator review. Replacing proof should return the verification state to pending so the new document is reviewed."),
    ("Prepare common documents", "Upload frequently requested files such as enrollment proof, grades, residency, income evidence, or guardian documents before selecting a scholarship."),
    ("Review recommendations", "Use the DSS percentage and explanation as a guide. Review strengths, unmet criteria, distance, deadline, award, and missing files rather than relying on the percentage alone."),
    ("Inspect program details", "Read eligibility, benefits, document requirements, stages, examination or interview information, renewal rules, return-service terms, provider contact, location, and deadline."),
    ("Start the application wizard", "The selected scholarship is carried into the application. The checklist shows that program's requirements and lets the learner upload or replace files before continuing."),
    ("Accept application terms and submit", "The system performs role, email, profile, eligibility, deadline, duplicate, and document checks before creating the application and its initial history."),
    ("Track review and schedules", "Follow status, document decisions, replacement requests, examination or interview schedules, general program events, and notifications. Acknowledge examination or interview schedules when requested; distribution notices are informational and do not require acknowledgment."),
    ("Receive the outcome", "The provider records approval or rejection and later distribution-related updates. Actual examination administration and benefit distribution remain with the provider."),
])

add_heading("4.2 Why the Applicant Is Allowed to Explore First", 2)
add_bullets([
    "It gives immediate value before asking for sensitive household and school information.",
    "It helps a parent decide whether the platform has suitable opportunities for a younger learner.",
    "It reduces form abandonment caused by a long first session.",
    "It supports progressive disclosure: only the fields relevant to the learner's level are shown or required.",
    "It preserves a firm integrity gate at submission, where incomplete information would affect a real provider decision.",
])

add_heading("4.3 Adaptive Profile Requirements", 2)
add_data_table(
    ["Applicant situation", "Required or emphasized", "Conditional or not applicable"],
    [
        ["Preschool or early childhood", "Learner identity, birth date, school or center, location, household context, and guardian details.", "No GWA or college course requirement."],
        ["Elementary", "Current school, grade level, enrollment, location, household context, and guardian details.", "Course and college-specific fields are excluded."],
        ["Junior high school", "School, grade level, education status, household and location information, and guardian details.", "Track, strand, or course appears only when relevant."],
        ["Senior high school", "School, grade, track or strand, grading information where used, household, location, and guardian details for minors.", "College course is not required unless the program explicitly targets a future field."],
        ["TVET or ALS", "Program or training context, institution, level or status, location, household, and any relevant credential information.", "Traditional grade metrics should be optional when the provider accepts competency or pass/fail evidence."],
        ["College or graduate", "Institution, year level, course or program, academic scale/value, household, location, and preferences.", "Guardian information may be optional for an adult self-managed account."],
        ["Parent-managed minor account", "Learner information plus guardian name, relationship, contact information, and account ownership indicator.", "The guardian should not replace the learner's identity with the guardian's identity."],
    ],
    [1950, 3900, 3510],
    font_size=8.7,
)

add_heading("4.4 Parent or Guardian Guidance", 2)
add_bullets([
    "Create or manage the account for the learner only when the platform's policy permits guardian-managed accounts.",
    "Enter the child's name and education information in applicant fields; enter the adult's information only in guardian fields.",
    "Use an email and mobile number that the family checks regularly, because schedules and replacement requests may be time-sensitive.",
    "Explain terms, service obligations, travel, examinations, and interview requirements in language the learner can understand.",
    "Upload only documents required for a legitimate purpose. Avoid adding unrelated health, identity, or family records.",
    "Update account ownership and contact information when the learner becomes capable of managing the account, subject to policy and consent." ,
])

add_heading("4.5 Applicant Submission Rules", 2)
add_bullets([
    "The signed-in account must have the applicant role and must not be suspended.",
    "The email address must be verified.",
    "Required profile fields for the applicant's education level and age must be complete.",
    "The scholarship must be published, open, and within its deadline.",
    "The account must not already have an application for the same program.",
    "No hard eligibility blocker may remain.",
    "Every provider-required document type must have an available file in the applicant's prepared library or application workflow.",
    "The current application terms version must be accepted.",
])

add_heading("4.6 After Approval", 2)
add_body(
    "Approval is not the end of the scholarship lifecycle. The applicant should review the provider's schedule and "
    "conditions, acknowledge an examination or interview schedule when requested, complete any remaining verification, "
    "attend orientation or distribution, sign applicable agreements, and retain copies of relevant notices. A "
    "distribution announcement does not require applicant acknowledgment because the bell and email already deliver "
    "the notice. The provider may later record award, distribution, disbursement, or renewal-related statuses. The "
    "current platform communicates and tracks these events; it does not transfer funds or reconcile payments."
)


chapter("5. Scholarship Provider Process", "Part II | End-to-End Workflows")
add_heading("5.1 Provider Journey", 2)
add_numbered([
    ("Register separately", "Use the provider registration form and supply representative information, organization name, organization type, address, optional website, description, password, and terms acceptance. Registration returns to login without automatically signing in."),
    ("Verify email", "Confirm control of the provider contact address. Email verification is a prerequisite for trustworthy notification and account recovery."),
    ("Complete the organization profile", "Maintain organization description, contact details, address, website, and representative data so applicants and administrators can understand who operates the program."),
    ("Upload proof", "Submit organization registration, authorization letter, representative identification, school or office proof, or another permitted proof type from the provider profile."),
    ("Receive administrator verification", "The administrator approves, rejects, or requests replacement proof. A rejected provider may upload a new file, returning the request to pending."),
    ("Create a draft program", "After verification, define identity, target applicants, academic and financial criteria, locations, documents, award, deadline, contacts, stages, schedules, renewal, and obligations."),
    ("Save or submit", "Keep incomplete work as a draft. Submit only when the program is ready for administrator review and the provider accepts the current scholarship-posting terms."),
    ("Respond to program review", "If rejected or returned, edit the program using administrator notes and resubmit. Material changes to a published program should return to review."),
    ("Publish through approval", "Only an administrator publishes an approved program. The provider should not bypass review with status manipulation."),
    ("Receive applications", "Use the application queue to inspect the learner's profile, proof, required files, DSS information, and review history for programs owned by the provider."),
    ("Review evidence", "Accept a file, reject it, or request a replacement in the file-view workflow. Use the rubric and notes for consistent assessment."),
    ("Decide and schedule", "Approve or reject after eligibility and requirements review. Use program stages and shared or individual schedules for examinations, interviews, and distribution rather than manually inventing a status for each person."),
    ("Close the cycle", "Record award and distribution-related outcomes, communicate conditions, retain only necessary records, and use exports or future aggregate reports for program evaluation."),
])

add_heading("5.2 Provider Verification Proof", 2)
add_data_table(
    ["Proof type", "What it may demonstrate", "Review caution"],
    [
        ["Organization registration", "Legal or recognized existence of the organization.", "Check issuing body, names, date, and consistency with the profile."],
        ["Authorization letter", "Representative is authorized to act for the organization.", "Confirm signatory authority and contact details through an independent channel where feasible."],
        ["Representative valid ID", "Identity of the person operating the account.", "Keep private, limit access, and retain only as long as necessary."],
        ["School or office proof", "Operational address or institutional relationship.", "Do not treat a logo or social-media page alone as sufficient proof."],
        ["Other accepted evidence", "Alternative proof for small community providers.", "Use a documented exception rule and require corroboration rather than lowering trust standards silently."],
    ],
    [2000, 3380, 3980],
    font_size=8.8,
)

add_heading("5.3 Program Creation Structure", 2)
add_bullets([
    "Program identity: title, category, logo, concise summary, full description, and provider attribution.",
    "Target applicants: eligible education levels, courses or strands, school types, year or grade levels, locations, income conditions, and other eligibility rules.",
    "Academic rule: minimum value and declared grading scale, or a non-numeric alternative when grades are not applicable.",
    "Benefit: amount or in-kind value, number of slots, coverage, payment frequency, duration, and exclusions.",
    "Evidence: clear document types using standardized names, with provider-specific forms identified separately.",
    "Process: application mode, screening, optional examination, optional interview, and distribution stages.",
    "Examination details: duration, passing score, venue or delivery method, and what the applicant should prepare when relevant.",
    "Location: address and map pin so applicants can understand distance and whether relocation is necessary.",
    "Obligations: renewal policy, return-service contract, other conditions, and grounds for termination.",
    "Contact and timeline: official email, phone, deadline, expected review window, and schedule publication method.",
    "Review setup: rubric criteria and weights appropriate to the target group, with a maximum of six criteria in the current interface." ,
])

add_heading("5.4 Selection Stage Logic", 2)
add_data_table(
    ["Stage", "When used", "Applicant-facing information"],
    [
        ["Screening", "Every program; checks eligibility and submitted evidence.", "Current status, missing or rejected files, replacement request, and decision note where appropriate."],
        ["Examination", "Only when selected during program creation.", "Purpose, date, time, venue or link source, duration, passing rule, materials, and provider contact."],
        ["Interview", "Only when selected during program creation.", "Date, time, venue or method, expected documents, attendance instructions, and rescheduling policy."],
        ["Approval or rejection", "Provider's final application decision after applicable review stages.", "Clear outcome and an appropriate reason or next action; the DSS never makes this decision."],
        ["Distribution", "Every awarded program; covers orientation, signing, release, or benefit delivery.", "Schedule, address, required identification, agreements, and provider instructions. It is a notice and does not require applicant acknowledgment."],
    ],
    [1600, 3300, 4460],
    font_size=8.8,
)

add_heading("5.5 Review Rubric", 2)
add_body(
    "The current default rubric emphasizes eligibility at 35%, academics at 25%, financial need at 20%, and "
    "document completeness at 20%, totaling 100%. A provider may configure an appropriate rubric within the "
    "current limit, but each criterion should be observable, relevant, non-duplicative, and understandable. The "
    "rubric is a review aid; it should not conceal an automatic decision or reproduce discriminatory criteria."
)
add_bullets([
    "Define what evidence supports each criterion before reviewing applicants.",
    "Do not score the same factor twice under different names.",
    "Use a consistent scale and provide notes for unusual scores.",
    "Separate mandatory eligibility from competitive ranking.",
    "Require a reason for rejection and preserve a status history.",
    "Periodically compare decisions across education level, location, gender, disability, and other lawful fairness dimensions using de-identified data and approved governance." ,
])

add_heading("5.6 Provider Responsibilities", 2)
add_bullets([
    "Fund and administer the scholarship described in the listing.",
    "Publish accurate criteria, dates, benefits, contact details, and obligations.",
    "Use applicant information only for stated scholarship purposes.",
    "Restrict document access to authorized staff and delete records according to policy.",
    "Conduct examinations, interviews, and distribution outside the platform when required, while posting accurate schedules and outcomes in the portal.",
    "Provide reasonable accommodation and accessible communication where applicable.",
    "Avoid hidden fees, misleading claims, pay-to-win selection, or unrelated marketing use of applicant data.",
    "Respond to replacement requests, complaints, and corrections within a published service level.",
])


chapter("6. Administrator Process and Trust Controls", "Part II | End-to-End Workflows")
add_heading("6.1 Administrator Responsibilities", 2)
add_numbered([
    ("Monitor review queues", "Prioritize pending provider proof, applicant proof, and scholarship submissions rather than using the dashboard as a full worklist."),
    ("Verify providers", "Open the provider detail page, compare profile and proof, use independent confirmation when appropriate, and approve, reject, or request replacement with a reason."),
    ("Review programs", "Check provider status, completeness, target criteria, benefits, dates, evidence requirements, stages, obligations, contacts, and potentially harmful or misleading language before publishing."),
    ("Review applicant proof", "Open the dedicated applicant review page, inspect only authorized proof, and verify or reject with a documented reason. A changed proof should require a new review."),
    ("Manage users", "Create or edit accounts, resend or manually confirm email where policy permits, suspend or reactivate accounts, and require a password reset."),
    ("Protect administrator continuity", "Do not suspend the current administrator or remove the final active administrator. The existing guard prevents this operational failure."),
    ("Review logs", "Use paginated activity logs to understand account, review, program, document, and status changes."),
    ("Export cautiously", "Use exports only for a defined administrative purpose, secure the output, and delete it after the approved retention period."),
    ("Handle incidents and complaints", "Preserve evidence, limit further exposure, communicate through approved channels, and follow the incident-response policy."),
])

add_heading("6.2 Provider Review Checklist", 2)
add_bullets([
    "Organization name and type are consistent across profile and proof.",
    "Representative identity and authority are supported.",
    "Address, website, official email, and contact number are plausible and independently confirmable.",
    "The organization has a clear scholarship purpose and no unexplained payment request.",
    "Uploaded proof is readable, current where required, and not visibly altered.",
    "Potential conflicts, duplicate accounts, or high-risk claims are escalated rather than approved from incomplete evidence.",
    "Decision notes explain the outcome without exposing sensitive document content to unauthorized users.",
])

add_heading("6.3 Program Publication Checklist", 2)
add_bullets([
    "Provider is approved and not suspended.",
    "Program title and description clearly identify the opportunity and target learner.",
    "Eligibility fields agree with the narrative and do not accidentally exclude open categories.",
    "Award, slots, duration, payment schedule, and exclusions are understandable.",
    "Deadline and selection schedule are realistic and not already expired.",
    "Document requests are necessary, standardized, and proportionate to the decision stage.",
    "Examination, interview, renewal, service, and distribution terms are disclosed before application.",
    "Official contact and location are provided.",
    "Logo or image is appropriate; the default scholarship image is used only when no valid logo is supplied.",
    "Terms are accepted and no misleading guarantee, hidden charge, or discriminatory condition is present." ,
])

add_heading("6.4 Account and Audit Controls", 2)
add_body(
    "The administrator can create and edit accounts on dedicated pages, suspend or reactivate access, force a "
    "password reset, and manage email verification under policy. Activity logs include the actor, action, IP "
    "address, user agent, and metadata. Application status histories and DSS snapshots provide more specific "
    "records for consequential decisions. These controls support investigation, but log access itself must be "
    "restricted because metadata may contain personal information."
)

add_callout(
    "Separation of duties",
    "For a real deployment, one administrator should not be able to verify a high-risk provider, publish its program, and erase all related evidence without independent review. Add dual approval or supervisory review for sensitive cases as the platform grows.",
    "warning",
)


chapter("7. Notifications, Email, and Mobile Process", "Part II | End-to-End Workflows")
add_heading("7.1 Notification Flow", 2)
add_body(
    "Portal notifications are stored in the database and shown through a shared notification bell with an unread "
    "count. Applicant, provider, and administrator layouts use the same list and mark-as-read behavior. A user may "
    "open one notification, mark it read, or mark all notifications read. Notifications are targeted to the "
    "recipient's account rather than rendered only as page-level messages."
)
add_body(
    "An observer queues email delivery for portal notifications, except email-verification messages, which use the "
    "authentication verification flow. The queue worker must be running for queued email to leave the application. "
    "The scheduler generates deadline reminders at 8:00 AM for configured intervals such as seven, three, one, and "
    "zero days, and uses idempotency controls to avoid duplicate reminders."
)

add_heading("7.2 Recommended Notification Rules", 2)
add_data_table(
    ["Event", "Recipient", "Channel and urgency"],
    [
        ["Email verification or password reset", "Account owner", "Email; time-sensitive security action."],
        ["Provider or applicant proof decision", "Submitted account", "Bell and email; include outcome and next action without attaching private proof."],
        ["Program published or returned", "Provider", "Bell and email; identify program and review note."],
        ["Document replacement requested", "Applicant", "Bell and email; identify requirement, reason, and deadline."],
        ["Application status decision", "Applicant", "Bell and email; provide clear next step and avoid ambiguous status codes."],
        ["Examination, interview, or distribution schedule", "Relevant applicants", "Bell and email; mobile push is a future enhancement. Include date, time, location, preparation, and contact."],
        ["Deadline reminder", "Saved applicants and provider", "Bell and queued email; do not send to applicants who already applied for that program."],
        ["Security event", "Account owner and administrator as appropriate", "Email plus internal alert; never include passwords or full document contents."],
    ],
    [2300, 1900, 5160],
    font_size=8.8,
)

add_heading("7.3 Mobile Scope", 2)
add_bullets([
    "Applicant sign-in uses expiring, hashed mobile API tokens.",
    "Provider and administrator accounts are rejected by the applicant mobile API.",
    "The mobile app uses the same Laravel data source for scholarship discovery, saved programs, profile, prepared documents, applications, schedules, and notifications.",
    "The mobile experience should always disclose when a file upload, map tile, email, or fresh synchronization requires internet access.",
    "Future mobile work should prioritize full profile parity, push notifications, secure token rotation, offline read caching, accessibility, and reliable large-file upload recovery." ,
])


# ---------------------------------------------------------------------------
# Scholarship domain reference
# ---------------------------------------------------------------------------

chapter("8. Complete Scholarship Knowledge Guide", "Part III | Scholarship Domain Reference")
add_callout(
    "How to use this chapter",
    "No single handbook can replace the official rules of every scholarship. This chapter provides the complete information categories, questions, controls, and good practices needed to understand, design, publish, apply for, review, and administer a scholarship responsibly.",
    "info",
)

add_heading("8.1 What a Scholarship Is", 2)
add_body(
    "A scholarship is educational support awarded under defined criteria and conditions. It may cover tuition, fees, "
    "books, transportation, accommodation, devices, meals, uniforms, allowances, training, examinations, or other "
    "education-related needs. Some awards are based mainly on academic merit; others prioritize financial need, "
    "location, talent, identity or community, field of study, service commitment, disability support, crisis, or a "
    "combination of factors."
)

add_heading("8.2 Related Forms of Student Financial Assistance", 2)
add_data_table(
    ["Type", "Typical meaning", "Important distinction"],
    [
        ["Scholarship", "Award based on stated merit, need, target group, service, talent, or mixed criteria.", "Usually not repaid, but may have renewal or service conditions."],
        ["Grant-in-aid", "Support primarily responding to financial need or a defined access problem.", "May use less academic competition but still requires eligibility evidence."],
        ["Subsidy", "Partial support for a defined education cost.", "May be automatic for an eligible group or attached to a public program."],
        ["Allowance or stipend", "Periodic cash or in-kind support for living, transport, supplies, or participation.", "The payment schedule and permitted use should be explicit."],
        ["Student loan", "Funds that must normally be repaid under agreed terms.", "Must never be presented as a scholarship; interest, repayment, and consequences require clear disclosure."],
        ["Fellowship", "Support commonly connected to advanced study, research, professional training, or service.", "May include work, research, publication, or residency obligations."],
        ["Prize or competition award", "One-time reward for achievement or contest performance.", "May not support ongoing education and may have different tax or publicity implications."],
        ["Tuition waiver or discount", "A school reduces all or part of tuition or fees.", "Confirm whether the benefit is renewable and which fees remain payable."],
    ],
    [1800, 3600, 3960],
    font_size=8.7,
)

add_heading("8.3 Ways Scholarships Are Classified", 2)
add_bullets([
    "By basis: merit, financial need, talent, leadership, community service, research, athletics, disability, indigenous or cultural community, geographic priority, crisis, or mixed criteria.",
    "By sponsor: government, school, private foundation, company, professional association, religious or civic organization, alumni group, community organization, or international institution.",
    "By education level: preschool, elementary, junior high, senior high, ALS, TVET, college, graduate, professional, or continuing education.",
    "By benefit: full tuition, partial tuition, fixed cash grant, stipend, in-kind supplies, equipment, transport, housing, mentoring, internship, or bundled support.",
    "By duration: one-time, semester, academic year, multi-year, renewable, or milestone-based.",
    "By geography: national, regional, provincial, city or municipality, barangay, school-specific, or relocation-based.",
    "By application model: open application, school nomination, provider invitation, automatic consideration, competition, examination, interview, or multi-stage review.",
])

add_heading("8.4 Information Every Scholarship Listing Should Contain", 2)
add_data_table(
    ["Information group", "Required questions"],
    [
        ["Identity and legitimacy", "Who funds and administers the program? What is the official name, provider identity, contact, website, address, and proof of authority?"],
        ["Purpose and target", "What access problem or learner group does the program support? Which education levels, places, courses, schools, or circumstances are eligible?"],
        ["Benefit", "What exactly is awarded, how much, how often, for how long, to how many recipients, and which costs are excluded?"],
        ["Eligibility", "Which conditions are mandatory? Which factors are preferences or competitive criteria? What grading scale is used?"],
        ["Requirements", "Which files, forms, essays, references, or certifications are needed, at what stage, in which format, and by what deadline?"],
        ["Selection", "Who reviews, what stages apply, whether there is an examination or interview, how criteria are weighted, and how conflicts are managed?"],
        ["Timeline", "When does the application open and close? When are screening, exam, interview, decision, distribution, and renewal expected?"],
        ["Location and access", "Where is the provider, exam, interview, study, or distribution site? Are travel, relocation, connectivity, or accessibility constraints involved?"],
        ["Conditions", "What academic maintenance, attendance, conduct, reporting, service, internship, publicity, or return obligations apply?"],
        ["Privacy and complaints", "How will data be used, shared, retained, corrected, deleted, and challenged? Where can an applicant report fraud or appeal a process error?"],
    ],
    [2250, 7110],
    font_size=8.8,
)

add_heading("8.5 Eligibility Criteria", 2)
add_body(
    "Eligibility determines who may enter the process; selection determines who receives a limited award. Providers "
    "should distinguish mandatory conditions from preferences. A mandatory criterion should be necessary, lawful, "
    "measurable, and disclosed before application. A preference may influence ranking but should not be presented as "
    "an automatic rejection rule."
)
add_bullets([
    "Education level, grade or year level, enrollment status, school type, course, program, track, or strand.",
    "Academic performance expressed using a declared scale, rank, competency, pass/fail status, or other suitable evidence.",
    "Household income, income bracket, indigency status, dependents, displacement, emergency, or other financial-need indicator.",
    "Residence, school location, provider service area, distance, or willingness and ability to relocate.",
    "Citizenship or residency only where genuinely required by the funding source and lawful program purpose.",
    "Talent, leadership, service, research interest, work experience, disability support, or other program-related factor.",
    "Availability for examination, interview, orientation, internship, service, or distribution obligations.",
])

add_heading("8.6 Education-Level Appropriate Forms", 2)
add_data_table(
    ["Target group", "Appropriate questions", "Usually inappropriate unless justified"],
    [
        ["Preschool and elementary", "Age, grade, school, enrollment, residence, household context, guardian, support need, and basic records.", "College course, college GWA, employment history, or learner-owned contact requirements."],
        ["Junior high school", "Grade level, school, academic or competency evidence, household context, guardian, location, activities, and support need.", "College-major restrictions or adult-only contract language."],
        ["Senior high school", "Grade, track or strand, school, academic scale, location, household context, intended pathway, and guardian where a minor.", "A current college course unless the field is framed as intended study."],
        ["TVET and ALS", "Training or equivalency level, institution, competency or completion evidence, location, employment context where relevant, and support needs.", "Mandatory GWA where competency or pass/fail evidence is the valid measure."],
        ["College", "Course, year level, institution, academic scale, household context, location, activities, intended outcomes, and service availability.", "Guardian consent for an adult unless a specific lawful reason exists."],
        ["Graduate or professional", "Degree or research area, prior qualification, proposal or portfolio, institutional status, professional experience, and research or service commitments.", "Elementary-style parent information or unrelated household details."],
    ],
    [1800, 4050, 3510],
    font_size=8.6,
)

add_heading("8.7 Grading Systems and Alternatives to GWA", 2)
add_body(
    "Programs should never assume that every institution uses the same numeric scale. The provider must state the "
    "accepted grading system and avoid converting values without a documented equivalency method. The current system "
    "supports a minimum academic value paired with a declared scale, and future work should expand structured "
    "equivalency while preserving the original reported value."
)
add_bullets([
    "Percentage average, such as a 0-100 scale.",
    "GPA on a 4.0 scale, where higher values are generally stronger.",
    "Philippine college scales such as 1.0-5.0, where lower values may be stronger depending on the institution.",
    "Letter grades or descriptive ratings.",
    "Class rank or percentile.",
    "Competency-based completed/not-yet-competent evidence.",
    "Pass/fail or satisfactory/unsatisfactory.",
    "Portfolio, audition, performance, research proposal, or examination score where grades are not the best measure.",
    "No academic threshold for need-based, emergency, early-childhood, or school-supplies support when grades are not relevant." ,
])

add_heading("8.8 Common Scholarship Documents", 2)
add_data_table(
    ["Category", "Examples", "Purpose and handling"],
    [
        ["Identity and age", "Birth certificate, school ID, other accepted identity record.", "Confirm applicant identity or age. Collect the least sensitive acceptable evidence and keep it private."],
        ["Enrollment", "Current enrollment certificate, registration form, school certification, learner reference information where necessary.", "Confirm school and level. Do not publish LRN or similar identifiers."],
        ["Academic", "Report card, certificate of grades, transcript, ranking, competency certificate, portfolio.", "Support academic criteria. Preserve the original scale and marking period."],
        ["Residence", "Barangay certificate, billing or official address proof, school location certification.", "Confirm geographic eligibility. Avoid requiring multiple proofs when one reliable document is sufficient."],
        ["Financial need", "Certificate of indigency, income tax return, certificate of employment and compensation, no-income affidavit, social-welfare certification.", "Support need criteria. These records are sensitive and should have short, defined retention."],
        ["Guardian", "Guardian ID, relationship evidence, authorization or consent form.", "Support minor accounts and consent. Limit collection to what the process needs."],
        ["Character or recommendation", "Good moral certificate, teacher or community recommendation.", "Support conduct, engagement, or potential criteria. Explain how recommendations are evaluated."],
        ["Program-specific", "Essay, application form, undertaking, research proposal, portfolio, audition, medical fitness where necessary.", "Request only when directly connected to selection or participation."],
        ["Special circumstance", "Disability, displacement, indigenous-community, solo-parent, disaster, or medical evidence where relevant.", "Highly sensitive; make the purpose explicit, restrict access, and consider less intrusive alternatives."],
        ["Award and distribution", "Bank or payment details, signed agreement, receipt, proof of enrollment after award.", "Collect late in the process from selected applicants, not from every applicant."],
    ],
    [1600, 3440, 4320],
    font_size=8.5,
)

add_callout(
    "Data-minimization rule",
    "Do not request a document merely because another scholarship asks for it. Every file should answer a defined eligibility, selection, award, audit, or legal requirement. Collect high-risk financial, medical, government, and payment data as late as possible and only from people who need to provide it.",
    "warning",
)

add_heading("8.9 Benefits and Funding Terms", 2)
add_bullets([
    "State whether the amount is per month, semester, academic year, milestone, or one-time award.",
    "Separate cash from in-kind support and identify who receives payment: applicant, school, vendor, or another party.",
    "List covered and excluded expenses, release conditions, installment dates, and what happens when school calendars change.",
    "Disclose the number of slots or explain if the number depends on available funds.",
    "Explain whether the award may be combined with other scholarships and whether duplicate public funding is restricted.",
    "Explain refund, suspension, overpayment, withdrawal, and unclaimed-award rules.",
    "Keep provider funding evidence and approval authority separate from applicant-facing marketing claims." ,
])

add_heading("8.10 Selection Stages", 2)
add_numbered([
    ("Administrative screening", "Confirms deadline, completeness, identity, and basic mandatory criteria."),
    ("Eligibility review", "Determines whether the applicant meets disclosed program conditions."),
    ("Document validation", "Checks readability, relevance, authenticity indicators, and replacement needs."),
    ("Competitive assessment", "Uses academic, need, essay, talent, service, leadership, or other disclosed criteria."),
    ("Examination", "Used only when relevant, with clear scope, security, accommodation, passing rule, and retake policy."),
    ("Interview", "Clarifies motivation, context, plans, or obligations using consistent questions and scoring."),
    ("Final deliberation", "Authorized provider reviewers resolve rankings, conflicts, limited slots, and documented exceptions."),
    ("Approval and notification", "Applicants receive a clear outcome, next action, and timeline."),
    ("Distribution and monitoring", "Provider completes agreement, release evidence or recipient confirmation only when policy requires it, and later renewal or completion checks."),
])

add_heading("8.11 Examination Good Practice", 2)
add_bullets([
    "Explain what the examination measures and how it relates to the scholarship purpose.",
    "Publish date, time, venue or delivery method, duration, permitted materials, identity checks, passing or ranking method, and contact for questions.",
    "Use secure handling and avoid exposing answers or individual scores to unauthorized people.",
    "Provide reasonable accommodation and an incident or rescheduling process.",
    "Do not ask the platform to imply it administered an examination that the provider conducted externally.",
    "Record only the status or score needed for the scholarship process and retain it according to policy." ,
])

add_heading("8.12 Interview Good Practice", 2)
add_bullets([
    "Use a consistent question set linked to disclosed criteria.",
    "Train interviewers to avoid irrelevant personal or discriminatory questions.",
    "Use at least two reviewers for high-stakes decisions where feasible.",
    "Record scores and concise evidence-based notes, not unnecessary personal observations.",
    "Disclose location, method, expected duration, documents, dress or equipment only when relevant, and rescheduling rules.",
    "Allow a guardian to support a minor where appropriate without replacing the learner's voice." ,
])

add_heading("8.13 Contracts, Return Service, and Other Obligations", 2)
add_body(
    "A scholarship may require attendance, grade maintenance, reporting, internships, mentoring, community service, "
    "employment, return service, publicity consent, or repayment after a breach. These terms can materially affect a "
    "learner's future and must be visible before application, written in understandable language, and reviewed with a "
    "parent or guardian when the applicant is a minor. The platform should show these obligations; the provider is "
    "responsible for the contract and its lawful enforcement."
)
add_bullets([
    "Identify the exact obligation, duration, location, supervisor, and acceptable proof of completion.",
    "Separate mandatory service from optional volunteering or mentoring.",
    "Explain consequences of withdrawal, academic difficulty, illness, relocation, or force majeure.",
    "State whether repayment is required and how it is calculated; obtain appropriate legal review.",
    "Do not hide an employment bond or marketing consent inside a generic terms checkbox.",
    "Provide a contact and process for questions, accommodation, correction, or dispute." ,
])

add_heading("8.14 Renewal, Suspension, and Termination", 2)
add_bullets([
    "State whether renewal is automatic, competitive, or dependent on available funds.",
    "Publish required academic performance using the correct scale and period.",
    "Explain enrollment load, attendance, conduct, reporting, and service requirements.",
    "Allow correction of records and provide a reasonable notice period for missing renewal evidence.",
    "Define probation, temporary suspension, voluntary withdrawal, transfer, leave of absence, and termination.",
    "Avoid retroactive criteria changes after the applicant accepts the award unless law or a documented emergency requires it." ,
])

add_heading("8.15 Applicant Due-Diligence and Scam Warning Signs", 2)
add_bullets([
    "Verify the provider using an official website, school office, government registry, or independently obtained contact information.",
    "Be cautious when an award is guaranteed without review, when urgency prevents verification, or when the provider asks for secrecy.",
    "Do not pay an unexplained application, reservation, processing, or release fee. A legitimate fee, if ever allowed, must be lawful, disclosed, receipted, and independently verifiable.",
    "Never share a password, one-time code, full banking credential, or unnecessary government identity record.",
    "Check whether the email domain, phone, address, and account name match the provider.",
    "Inspect spelling, dates, benefit details, terms, and contact information; copied logos alone do not prove legitimacy.",
    "Report suspicious listings to the platform administrator and preserve the message, link, date, and evidence." ,
])

add_heading("8.16 Applicant Application Checklist", 2)
add_bullets([
    "Read the complete program details, not only the award amount and match percentage.",
    "Confirm education level, course or strand, school type, year level, location, income, academic scale, and deadline.",
    "Review examination, interview, relocation, service, and renewal obligations with the family.",
    "Use clear scans of current records and confirm that each file opens before submission.",
    "Write essays in the applicant's own words and preserve drafts and submission confirmation.",
    "Use an active email and contact number, check the notification bell, and acknowledge schedules promptly.",
    "Update changed contact, school, grade, address, or proof information before it affects a decision.",
    "Apply to several suitable programs while avoiding duplicate or false submissions." ,
])

add_heading("8.17 Provider Scholarship Design Checklist", 2)
add_bullets([
    "Start with the access problem, target population, funding amount, and outcome the program intends to support.",
    "Choose the fewest mandatory criteria necessary to reach that purpose.",
    "Use age- and education-appropriate forms rather than one college-centered application for everyone.",
    "Distinguish eligibility, competitive scoring, document readiness, and final deliberation.",
    "Request sensitive proof only at the stage when it can affect a decision.",
    "Publish realistic timelines, service levels, examination or interview rules, and distribution dates.",
    "Provide accommodation, correction, complaint, and conflict-of-interest procedures.",
    "Budget not only the award but also review staff, communication, verification, accessibility, data security, and evaluation.",
    "Measure completion, selection, award, renewal, and outcome quality rather than counting only page views." ,
])


chapter("9. Decision Support System", "Part IV | System, Data, and Governance")
add_heading("9.1 Purpose", 2)
add_body(
    "The decision support system helps an applicant understand which scholarships deserve attention and helps the "
    "platform explain why. It is not an automated admissions officer. The current methodology produces a suitability "
    "percentage and explanation from structured applicant and scholarship inputs, then preserves a snapshot for "
    "audit and reproducibility."
)

add_heading("9.2 Current Methodology", 2)
add_data_table(
    ["Component", "Weight", "What it represents"],
    [
        ["Eligibility alignment", "65%", "Education level, course/strand/track, school type, grade or year level, location, income, and other structured fit."],
        ["Academic alignment", "20%", "Applicant academic value interpreted against the provider's declared scale and minimum, when academics apply."],
        ["Financial-need alignment", "15%", "Applicant household context compared with the program's disclosed income or need conditions."],
        ["Document readiness", "Displayed separately", "Whether required files are prepared; it does not inflate or reduce suitability."],
        ["Provider review progress", "Displayed separately", "What the provider has reviewed; it does not alter the original fit calculation."],
    ],
    [2400, 1700, 5260],
    font_size=8.9,
)

add_heading("9.3 Open-Criteria Handling", 2)
add_body(
    "Values such as any, all, no restriction, and nationwide must behave as open criteria. They should not be compared "
    "as literal course or location names. This prevents a STEM learner from being marked as a mismatch when a program "
    "actually accepts all courses. The same normalization principle should apply to education level, school type, "
    "location, grade, course, track, and strand."
)

add_heading("9.4 Hard Blockers and Readiness", 2)
add_bullets([
    "A disclosed mandatory mismatch may block application submission and should be explained in plain language.",
    "A missing required document blocks submission but is shown as a readiness task rather than a lower personal suitability score.",
    "A missing optional preference should not be presented as an absolute ineligibility.",
    "Uncertain or unstructured criteria should be shown for manual review rather than converted into a false precise score.",
    "The applicant should always be able to open the program and read the provider's original rule." ,
])

add_heading("9.5 Transparency Output", 2)
add_bullets([
    "Overall suitability percentage.",
    "Matched strengths and the criteria that support them.",
    "Items needing attention, such as an uncertain grade scale or location condition.",
    "Missing required documents shown separately.",
    "A next action, such as completing profile information, preparing a file, or reviewing program terms.",
    "A methodology version so later changes do not rewrite the meaning of an old application silently." ,
])

add_heading("9.6 DSS Snapshot and Auditability", 2)
add_body(
    "Each calculation snapshot can preserve applicant inputs, scholarship inputs, document inputs, methodology "
    "version, input hash, and resulting output. This supports reproducibility: a reviewer can determine which "
    "information and method produced the recommendation at that time. The snapshot must be protected as personal "
    "data and should not become a public profile."
)

add_heading("9.7 Risks and Fairness Controls", 2)
add_data_table(
    ["Risk", "Example", "Control"],
    [
        ["False precision", "A percentage appears more certain than incomplete profile data allows.", "Show missing inputs, confidence or readiness, and plain-language reasons."],
        ["Scale mismatch", "A 1.5 grade is treated as worse than 85 without understanding the scale.", "Store the original scale and value; use provider-declared equivalency or manual review."],
        ["Proxy discrimination", "Location or school type indirectly excludes a protected or disadvantaged group.", "Require purpose justification, fairness review, and aggregate outcome monitoring."],
        ["Historical bias", "Future machine learning copies prior provider decisions.", "Do not deploy predictive ranking until data quality, lawful basis, bias testing, and human oversight are established."],
        ["Automation bias", "Provider accepts the system score without reading evidence.", "Display the score as support, require human decision and notes, and audit score-decision disagreement."],
        ["Applicant gaming", "Profile values are changed only to increase a score.", "Use proof where necessary, preserve histories, and verify consequential fields proportionately."],
    ],
    [1800, 3300, 4260],
    font_size=8.7,
)

add_callout(
    "Final authority",
    "The provider always makes the final scholarship decision. The DSS should inform, explain, prioritize, and document; it should never silently award or reject a learner.",
    "success",
)


chapter("10. Technical Architecture", "Part IV | System, Data, and Governance")
add_heading("10.1 Architecture Summary", 2)
add_data_table(
    ["Layer", "Technology and responsibility"],
    [
        ["Web front end", "Vue 3 components styled with Tailwind CSS 4 and compiled by Vite; responsive applicant navigation and role-specific provider/admin side panels."],
        ["Backend", "Laravel 12 on PHP 8.2; routing, validation, authentication, authorization, persistence, notifications, queues, scheduler, exports, and mobile API."],
        ["Database", "SQLite in the current local setup; production should use managed or properly administered MySQL with backup and restore procedures."],
        ["Mobile", "Flutter/Dart applicant application using Laravel mobile API endpoints and expiring hashed tokens."],
        ["Files", "Private scholarship evidence stored outside the public web root under Laravel storage, served only through authorized controller actions."],
        ["Email", "SMTP-compatible mail with queued portal notifications; verification and password-reset mail use authentication flows."],
        ["Scheduling", "Laravel scheduler for deadline reminders; operating system cron or equivalent must call it in production."],
        ["Maps", "Leaflet-based maps and geocoding/address search using internet-delivered tiles or services; address and selected coordinates are stored for distance visualization."],
        ["Health and operations", "Laravel health endpoint, logs, queue worker, failed-job handling, cache/config optimization, and deployment checks."],
    ],
    [2100, 7260],
    font_size=8.9,
)

add_heading("10.2 Request and Data Flow", 2)
add_numbered([
    ("User request", "Browser or mobile app sends an authenticated request to Laravel."),
    ("Middleware", "Laravel checks session or token, active-account status, role, CSRF for web requests, and route-specific authorization."),
    ("Validation and policy", "Controller or request validation checks required fields, file constraints, ownership, status, deadline, and process rules."),
    ("Domain operation", "Models and services read or update profiles, scholarships, documents, applications, DSS snapshots, statuses, events, and notifications."),
    ("Persistence", "Structured data is written to the database; private files are written to protected storage."),
    ("Response", "Vue receives a page or JSON response; Flutter receives API JSON."),
    ("Asynchronous work", "Queued notifications and email are processed by a worker; scheduled reminders are triggered by the scheduler."),
    ("Audit", "Relevant actions write activity logs, application histories, funnel events, or DSS snapshots."),
])

add_heading("10.3 Internet Dependency", 2)
add_data_table(
    ["Function", "Can work without internet on one local machine?", "Internet-dependent part"],
    [
        ["Core Laravel and SQLite", "Yes, after dependencies and assets are installed.", "External users cannot connect unless the server is network-accessible."],
        ["Compiled Vue/Tailwind assets", "Yes when built and served locally.", "CDN-only fonts, icons, or images will fail unless stored locally."],
        ["Leaflet map interface", "The code can load locally.", "Map tiles, address search, and remote geocoding normally need internet unless self-hosted or cached."],
        ["Email verification and notification mail", "Not to real external inboxes.", "SMTP or mail API and DNS/network connectivity are required."],
        ["Mobile on emulator or phone", "Only if the device can reach the Laravel server on the local network.", "Remote use needs hosting, HTTPS, and internet or an accessible private network."],
        ["Database", "SQLite works offline; MySQL can also work on a local network.", "Managed cloud MySQL requires network access."],
    ],
    [2000, 3160, 4200],
    font_size=8.7,
)

add_heading("10.4 Route Capability Summary", 2)
add_bullets([
    "Public and authentication routes cover landing content, terms, login, applicant registration, provider registration, email verification, and password reset.",
    "Applicant routes cover dashboard, profile sections, documents, scholarships, saved programs, program detail and map, applications, uploads, schedules, notifications, and account settings.",
    "Provider routes cover dashboard, profile and proof, programs, program creation and editing, applications, applicant and file review, schedules/events, exports, notifications, and account settings.",
    "Administrator routes cover dashboard, provider/applicant/program review, detail pages, user creation and editing, account controls, logs, exports, notifications, and account settings.",
    "Mobile API routes expose applicant functions only and reject provider or administrator use." ,
])


chapter("11. Database and Data Lifecycle", "Part IV | System, Data, and Governance")
add_heading("11.1 Account and Profile Separation", 2)
add_body(
    "The users table is the authentication core: email, username, password hash, role, verification, suspension, and "
    "account-control information. Role-specific details are stored in student_profiles, provider_profiles, and "
    "admin_profiles. This structure makes it clear which data is needed for login and which data belongs to the "
    "applicant, provider organization, or administrator role."
)

add_heading("11.2 Domain Table Dictionary", 2)
domain_rows = [
    ["users", "All roles", "Authentication identity, role, email verification, password controls, suspension, and core contact fields."],
    ["student_profiles", "Applicant", "Identity extensions, birth date, gender, education, school, grading, household, preferences, address/map, and guardian information."],
    ["provider_profiles", "Provider", "Organization identity, type, website, address, description, representative information, location, and verification state."],
    ["admin_profiles", "Administrator", "Administrator-specific profile details separated from authentication."],
    ["scholarships", "Provider-owned", "Program description, target criteria, grading rule, award, documents, stages, exam settings, location, deadline, status, contacts, contracts, and terms acceptance."],
    ["scholarship_events", "Provider/program", "Shared screening, exam, interview, distribution, or other events visible to relevant applicants."],
    ["scholarship_applications", "Applicant + program", "Application record, current status, decision details, DSS summary, and submission information."],
    ["application_documents", "Application", "Files attached to a specific requirement, review status, provider notes, and replacement decisions."],
    ["student_documents", "Applicant", "Reusable prepared common documents uploaded before or during applications."],
    ["applicant_verification_documents", "Applicant/admin", "Private proof used for applicant profile verification."],
    ["provider_verification_documents", "Provider/admin", "Private organization and representative proof used for provider verification."],
    ["application_status_histories", "Application", "Chronological status changes, actor, reason, notes, and timestamps."],
        ["application_schedules", "Application", "Applicant-specific exam, interview, distribution, or other schedule details; acknowledgment applies to exam/interview when requested, not distribution notices."],
    ["scholarship_bookmarks", "Applicant", "Saved scholarship relationships used for applicant lists and reminders."],
    ["portal_notifications", "All roles", "In-web notification title, message, link, type, read state, and delivery context."],
    ["activity_logs", "Administration/audit", "Actor, action, target, IP, user agent, and contextual metadata."],
    ["mobile_api_tokens", "Applicant mobile", "Hashed API token, expiration, device/use metadata, and revocation context."],
    ["scholarship_funnel_events", "Product/research", "Structured events such as profile completion, application start, document activity, submission, and status change."],
    ["dss_calculation_snapshots", "DSS/audit", "Inputs, methodology version, hash, explanation, and result at calculation time."],
]
add_data_table(["Table", "Owner/context", "Primary data"], domain_rows, [2200, 1760, 5400], font_size=8.2)

add_heading("11.3 Laravel Framework Tables", 2)
add_data_table(
    ["Table", "Purpose"],
    [
        ["migrations", "Records which database schema migrations have run."],
        ["sessions", "Stores web sessions when the configured session driver uses the database."],
        ["password_reset_tokens", "Supports time-limited password reset requests."],
        ["cache and cache_locks", "Database-backed cache values and locks when configured."],
        ["jobs", "Queued work waiting for a worker."],
        ["failed_jobs", "Queued work that failed and needs inspection or retry."],
        ["job_batches", "Metadata for grouped queued jobs."],
    ],
    [2500, 6860],
    font_size=9.0,
)

add_heading("11.4 File Storage", 2)
add_bullets([
    "Public program images or logos may be exposed through approved public storage paths and should fall back to the configured default scholarship image when no valid image exists.",
    "Applicant prepared documents, application files, applicant verification proof, and provider verification proof belong in private Laravel storage rather than the public web root.",
    "Authorized controller actions should check role, ownership, and relationship before streaming a file, and should return no-store and content-type protection headers for proof documents.",
    "Database rows store metadata and paths; the file bytes remain in storage. A database backup without the storage backup is incomplete." ,
])

add_heading("11.5 Data Lifecycle", 2)
add_numbered([
    ("Collect", "Ask only for data connected to an account, eligibility, application, verification, award, security, or approved research purpose."),
    ("Validate", "Check format, ownership, file type and size, relationship, and status before saving."),
    ("Use", "Make data available only to the applicant, relevant provider, authorized administrator, or processor that needs it."),
    ("Update", "Allow correction and preserve consequential histories where required; changed proof should trigger new review."),
    ("Retain", "Keep active records for the approved operational and legal period, with shorter periods for high-risk unsuccessful-applicant documents."),
    ("Archive or aggregate", "Separate de-identified research or impact data from active operational identity where possible."),
    ("Delete", "Remove database rows, storage files, exports, backups according to policy and document exceptions or legal holds."),
])

add_heading("11.6 Recommended Data Classification", 2)
add_data_table(
    ["Class", "Examples", "Minimum handling"],
    [
        ["Public", "Published scholarship title, provider public profile, deadline, award, eligibility summary, logo.", "Review before publication; permit correction and expiration."],
        ["Internal", "Operational counts, non-sensitive configuration, generic review notes.", "Authenticated access and appropriate role restrictions."],
        ["Confidential", "Applicant profile, contact, education, household, application, schedules, rubric scores.", "Need-to-know access, encryption in transit, logging, retention policy."],
        ["Highly sensitive", "Identity proof, income records, health/disability evidence, guardian proof, bank details, government identifiers.", "Strong access controls, late collection, short retention, encrypted storage, export restriction, incident priority."],
    ],
    [1700, 3560, 4100],
    font_size=8.7,
)


chapter("12. Security, Privacy, and Governance", "Part IV | System, Data, and Governance")
add_heading("12.1 Current Safeguards", 2)
add_bullets([
    "Role-based middleware for applicant, provider, and administrator workspaces, plus active-account checks.",
    "Password hashing, session regeneration, password reset, email verification, and account suspension controls.",
    "CSRF protection for web forms and throttling for authentication, upload, and mobile API activity.",
    "Hashed, expiring mobile API tokens and rejection of privileged roles from the applicant mobile API.",
    "Private document storage with role, ownership, and provider-application relationship checks.",
    "No-store and content-sniffing protection when sensitive proof is viewed through authorized routes.",
    "Versioned terms acceptance with timestamps for account, application, document, provider proof, and program contexts.",
    "Activity logs, application status histories, and DSS calculation snapshots for traceability.",
    "Protection against suspending or demoting the last active administrator." ,
])

add_heading("12.2 Privacy Principles", 2)
add_body(
    "The Philippine Data Privacy Act and its implementing framework emphasize legitimate and declared purposes, fair "
    "and lawful processing, accurate and relevant information, data that is adequate but not excessive, retention only "
    "for as long as necessary, and reasonable protection. The platform should translate these principles into visible "
    "notices, field-level purpose, access rules, retention schedules, correction, deletion, complaints, and incident "
    "handling rather than relying on a single generic terms checkbox."
)

add_heading("12.3 Minors and Guardian Governance", 2)
add_bullets([
    "Define when a guardian may create or manage an account and how the learner's assent or understanding is supported.",
    "Keep learner identity separate from guardian identity and record the relationship and account owner.",
    "Use age-appropriate explanations for matching, documents, schedules, obligations, and decisions.",
    "Avoid public profiles, unnecessary location precision, and unnecessary government identifiers for minors.",
    "Set a process for transferring control, correcting guardian information, family disputes, and the learner reaching the applicable age of independent account control.",
    "Require additional review before using a minor's photo, story, or award information for publicity." ,
])

add_heading("12.4 Controls Needed Before Public Hosting", 2)
add_data_table(
    ["Control", "Why it is needed", "Recommended implementation"],
    [
        ["Malware and file-content scanning", "Extensions and MIME declarations alone do not prove a file is safe.", "Quarantine uploads, inspect magic bytes, scan with a maintained engine, and release only clean files."],
        ["Multi-factor authentication", "Provider and administrator accounts can access consequential records.", "Require TOTP or passkey-based MFA for privileged roles and recovery codes under policy."],
        ["Retention and deletion", "Sensitive unsuccessful-applicant files should not remain forever.", "Approve a table-by-table schedule, automate deletion, document holds, and test backup expiration."],
        ["Encryption and key management", "Database and file compromise could expose education and financial records.", "Use HTTPS, encrypted disks/object storage, managed secrets, rotation, and field encryption where risk justifies it."],
        ["Backup restoration", "A backup is not useful until restoration is proven.", "Automate encrypted backups, separate locations, restoration drills, recovery objectives, and evidence."],
        ["Monitoring and incident response", "Queues, mail, login abuse, and file access can fail silently.", "Central logs, alerts, on-call ownership, breach workflow, evidence preservation, and communication templates."],
        ["Accessibility", "Learners may use assistive technology or low-end mobile devices.", "Audit toward WCAG 2.2 AA, keyboard use, focus, labels, contrast, reflow, error messages, and document accessibility."],
        ["Vendor and service governance", "Email, hosting, maps, and storage processors handle or observe data.", "Document processors, contracts, data locations, breach duties, availability, and exit plans."],
    ],
    [2160, 3100, 4100],
    font_size=8.6,
)

add_heading("12.5 Role Access Principles", 2)
add_bullets([
    "Applicants access their own profile, files, applications, schedules, saved programs, and notifications.",
    "Providers access their own organization, programs, and applicants who submitted to those programs; they do not browse unrelated applicant records.",
    "Administrators access verification and moderation data only for an approved purpose and should not use records for provider selection on behalf of a provider.",
    "Mobile endpoints expose applicant functions only.",
    "Exports require the same or stronger restrictions as the screen from which they were generated.",
    "Logs record access and decisions, but administrators should not place full sensitive documents or unnecessary personal text into log metadata." ,
])

add_heading("12.6 Terms and Consent", 2)
add_body(
    "The system supports context-specific modal terms rather than sending users to one distant page. Account terms "
    "cover truthful registration and account security; application terms cover accuracy and provider decision authority; "
    "document terms cover permission and secure use; provider proof terms cover authenticity; and scholarship-posting "
    "terms cover program accuracy, funding responsibility, privacy, and applicant treatment. Acceptance must identify "
    "the terms version and time. Material changes should trigger renewed acceptance."
)


chapter("13. Hosting, Operations, and Quality Assurance", "Part IV | System, Data, and Governance")
add_heading("13.1 Production Readiness Checklist", 2)
add_bullets([
    "Use a supported PHP runtime, Composer dependencies installed without development packages, and a production web server whose public root points only to Laravel's public directory.",
    "Set APP_ENV to production, APP_DEBUG to false, generate a unique application key, use HTTPS, and store secrets outside version control.",
    "Use a production MySQL database with least-privilege credentials, migrations, backups, monitoring, and tested restoration.",
    "Build front-end assets with a clean dependency install and production Vite build; do not rely on the development server in hosting.",
    "Create the approved public storage link only for public assets; keep applicant and provider proof private.",
    "Run a supervised queue worker for email and notifications, monitor failed jobs, and define retry behavior.",
    "Schedule Laravel's scheduler every minute so daily deadline reminders and future maintenance tasks run.",
    "Configure SMTP or a mail service using a dedicated account, verified sender, SPF, DKIM, DMARC, bounce handling, and rate monitoring.",
    "Configure logs, uptime and health checks, error reporting, disk-space alerts, database alerts, and security monitoring.",
    "Run migrations, automated tests, role-based smoke tests, accessibility checks, and backup restoration before launch." ,
])

add_heading("13.2 Typical Deployment Commands", 2)
add_callout(
    "Production sequence",
    "composer install --no-dev --optimize-autoloader; npm ci; npm run build; php artisan migrate --force; php artisan storage:link where approved; php artisan config:cache; php artisan route:cache; php artisan view:cache. Run a supervised queue worker and configure the scheduler separately. Adjust commands to the host and deployment policy.",
    "info",
)

add_heading("13.3 Backup and Recovery", 2)
add_bullets([
    "Back up both database and private/public storage; one without the other cannot restore a complete application record.",
    "Encrypt backups, restrict access, keep at least one independent copy, and document retention.",
    "Define recovery point objective, recovery time objective, and the person authorized to restore production.",
    "Test restoration into an isolated environment and verify accounts, applications, file paths, notifications, and migrations.",
    "Protect backup logs and avoid sending sensitive content in alert messages.",
    "Test deletion propagation according to policy, including expired backups where legally and technically possible." ,
])

add_heading("13.4 Automated and Manual Quality", 2)
add_body(
    "At the audit point, the project test suite passed 73 tests with 591 assertions, and the production front-end build "
    "completed. This is a strong engineering signal but not proof that every scholarship policy, browser, assistive "
    "technology, mobile device, provider workflow, and real-world document is correct. Automated tests should be "
    "combined with structured user acceptance testing."
)
add_data_table(
    ["Test area", "Examples"],
    [
        ["Authentication and RBAC", "Register, verify, login, reset, suspension, last-admin guard, provider/admin mobile rejection, cross-role route denial."],
        ["Applicant workflow", "Browse incomplete profile, conditional fields by education level, document preparation, match explanation, application gates, replacement upload, exam/interview acknowledgment, and distribution notice visibility."],
        ["Provider workflow", "Proof replacement, verification lock, draft program, review return, publication, applicant ownership, file decision, rubric, stage events, final reason."],
        ["Administrator workflow", "Provider/applicant/program detail pages, publish/reject, account controls, logs pagination, exports, notification bell."],
        ["Security", "CSRF, throttling, unauthorized file access, MIME/content checks, token expiry, session fixation, export access, XSS, SQL injection, upload abuse."],
        ["Accessibility and usability", "Keyboard, focus, labels, validation recovery, contrast, zoom, responsive layouts, screen readers, plain language, low-bandwidth behavior."],
        ["Operations", "Mail queue, failed jobs, scheduler idempotency, health endpoint, storage permissions, backup restore, cache clear, production debug disabled."],
    ],
    [2500, 6860],
    font_size=8.8,
)


# ---------------------------------------------------------------------------
# Sustainability, limitations, roadmap, and research
# ---------------------------------------------------------------------------

chapter("14. Sustainability and Provider Adoption", "Part V | Strategy and Future Development")
add_heading("14.1 Sustainability Definition", 2)
add_body(
    "Sustainability means the platform can continue operating, protecting data, supporting users, and improving "
    "scholarship outcomes after the capstone demonstration. It includes financial sustainability, provider participation, "
    "technical maintainability, governance, trustworthy data, staff capacity, and measurable learner value. A platform "
    "with many listings but weak verification or abandoned applications is not sustainable."
)

add_heading("14.2 Provider Return on Participation", 2)
add_data_table(
    ["Provider gain", "Operational effect", "Evidence to measure"],
    [
        ["Relevant reach", "Program is discovered by applicants who fit target criteria rather than only existing social followers.", "Qualified views, geographic reach, save-to-application conversion."],
        ["Fewer incomplete submissions", "Structured fields and pre-uploaded documents reduce clarification messages.", "Completeness at submission, replacement rate, staff follow-up time."],
        ["Faster review", "Profiles, proof, files, DSS context, rubric, and history are available in one workflow.", "Median time from submission to decision, applications reviewed per staff hour."],
        ["Consistent process", "Configured stages and criteria reduce ad hoc treatment.", "Rubric completion, reason coverage, status consistency, exception rate."],
        ["Applicant communication", "One notification and schedule system reduces missed updates.", "Delivery/read rate, schedule acknowledgment, no-show rate."],
        ["Trust and reputation", "Verification and reviewed listings distinguish the provider from scams or incomplete posts.", "Provider verification completion, complaints, applicant trust survey."],
        ["Impact reporting", "Structured outcomes can support board, donor, or CSR reporting after privacy controls.", "Awards, renewal, target-group reach, outcome follow-up, data-quality rate."],
    ],
    [1900, 3900, 3560],
    font_size=8.6,
)

add_heading("14.3 Ethical Revenue Options", 2)
add_bullets([
    "Freemium provider model: free verified basic listing and application intake; paid advanced workflow, team seats, integrations, or enhanced support.",
    "Provider subscription: predictable fee for repeated scholarship cycles, document workflows, scheduling, exports, and service-level support.",
    "Managed-service fee: onboarding, program configuration, applicant communication, or document administration performed by trained staff under contract.",
    "Institution or consortium license: schools, LGUs, foundations, or associations fund access for a group of providers and learners.",
    "API or integration fee: mature providers pay for approved connections to their own CRM, student information, identity, or reporting systems.",
    "Impact and compliance service: privacy-respecting aggregate program evaluation, donor reporting, and audit support after governance and analytics maturity.",
    "Sponsorship: clearly labeled support for platform operations without allowing sponsors to alter match scores or buy applicant decisions." ,
])

add_heading("14.4 Fee and Convenience-Fee Position", 2)
add_body(
    "The safest sustainability position is to keep discovery, matching, document preparation, and scholarship "
    "application free for applicants and families. If a fee is introduced, it should normally be a fixed provider-side "
    "service or application-cycle administration fee tied to optional work such as assisted onboarding, managed intake, "
    "bulk communication, integrations, or support. Calling it a provider service fee is clearer than calling it a "
    "convenience fee because the provider is paying for defined administrative value rather than simple website access."
)
add_bullets([
    "Never deduct a platform fee from the learner's scholarship award.",
    "Never charge an applicant to improve matching, ranking, review priority, verification, or probability of approval.",
    "Prefer a fixed or cost-based provider fee over a percentage of the award, which can discourage larger benefits and create conflicts of interest.",
    "If a payment processor adds a transaction charge, show the full amount before confirmation, issue a receipt, and identify which party receives the fee.",
    "Provide a no-fee or institution-funded path for applicants and consider waivers for small community providers during pilot adoption.",
    "Obtain legal, accounting, tax, payment-provider, and consumer-protection review before collecting real money." ,
])

add_heading("14.5 Revenue Practices to Avoid", 2)
add_bullets([
    "Charging applicants to improve match scores, appear more eligible, or move ahead in provider queues.",
    "Selling applicant contact or profile data to marketers without a separate lawful purpose and valid consent.",
    "Paid provider placement that looks like an objective recommendation; sponsored content must be labeled and must not affect eligibility calculations.",
    "Taking a percentage of a learner's award without clear legal, ethical, and provider authorization.",
    "Locking essential security, privacy, correction, or accessibility functions behind a paid tier.",
    "Promising providers guaranteed applicants or recipients regardless of the actual target population." ,
])

add_heading("14.6 Provider Adoption Plan", 2)
add_numbered([
    ("Start with small credible providers", "Pilot with organizations that have real programs but limited technical capacity; this makes the workflow value visible."),
    ("Offer assisted onboarding", "Help define criteria, document names, schedules, and terms so the first program is high quality."),
    ("Set verification service levels", "Publish what proof is required and how long review normally takes."),
    ("Prove time savings", "Measure incomplete submissions, replacement rate, review time, and communication effort before asking for payment."),
    ("Build trust artifacts", "Use verification dates, program review dates, clear contacts, transparent terms, and complaint handling."),
    ("Create renewal value", "Make repeated cycles, archived templates, applicant communication, and impact summaries easier than rebuilding the process elsewhere."),
    ("Use social media as a channel", "Providers can still announce programs socially, but direct applicants to the structured platform for official details and application."),
])

add_heading("14.7 Sustainability Metrics", 2)
add_bullets([
    "Applicant activation, profile completion, document readiness, suitable-program views, application start and completion, and schedule acknowledgment.",
    "Provider verification completion, time to first program, program review turnaround, program renewal, and provider retention.",
    "Application completeness, document replacement rate, review time, decision time, no-show rate, award rate, and renewal rate.",
    "Support volume, queue failures, mail delivery, uptime, restore-test success, security incidents, and accessibility defects.",
    "Fairness indicators across appropriate groups using de-identified aggregate data and minimum group-size protections.",
    "Cost per active applicant, cost per completed application, cost per provider, and provider staff time saved." ,
])


chapter("15. Current Limitations and Detailed Gap Analysis", "Part V | Strategy and Future Development")
add_heading("15.1 Known Boundaries", 2)
add_bullets([
    "The seeded providers, programs, and accounts are fictional demonstration data and must not be presented as active public scholarship offers.",
    "SQLite is suitable for local development and demonstrations but not the preferred database for a multi-user production rollout.",
    "The platform records and communicates examination and interview details; it does not administer secure online examinations.",
    "The platform records approval, distribution schedules, and later statuses; it does not currently transfer funds or reconcile payments.",
    "The provider is responsible for program funding, selection, contracts, and benefit distribution.",
    "Applicant profile proof verification is not currently a universal prerequisite for application submission; email, profile completeness, eligibility, documents, terms, and deadline are the active submission gates.",
    "Email delivery depends on valid SMTP configuration, internet connectivity, and a continuously running queue worker.",
    "Deadline reminders depend on the scheduler running and system time being correct.",
    "Leaflet map tiles and address search normally depend on external network services; a map failure must not prevent access to the written address.",
    "Mobile parity is applicant-focused and may lag behind complex web profile, proof, or document-review capabilities.",
    "Uploaded files do not yet have a complete production-grade malware quarantine and content-disarm pipeline.",
    "Analytics pages are intentionally deferred, although funnel events and DSS snapshots create a foundation for later approved reporting.",
    "Legacy schema or migrations, including older assessment concepts superseded by scholarship-stage fields, should be reviewed and cleaned before a long-lived production database is established.",
    "A comprehensive retention/deletion engine, data export for applicants, account deletion workflow, MFA, push notifications, and formal appeals are not yet complete." ,
])

add_callout(
    "Demonstration statement",
    "During a defense or pilot demonstration, say clearly which records are fictional, which operations happen outside the platform, and which roadmap items are planned. Honest boundaries strengthen the project because they show governance and engineering judgment.",
    "warning",
)

add_heading("15.2 Current Capability Coverage", 2)
add_body(
    "The following matrix separates working behavior from partial coverage, missing capabilities, and operational "
    "dependencies. Implemented means represented in the current source code and tested behavior; it does not mean "
    "that the related policy, hosting service, staffing, or external provider process is automatically production-ready."
)
capability_gap_rows = [
    ["Public platform information", "Implemented", "Landing, login, registration, provider registration, terms modals, and public navigation exist.", "Continue plain-language, accessibility, and real-provider content validation."],
    ["Role authentication and RBAC", "Implemented", "Applicant, provider, and administrator routes are separated; mobile authentication is applicant-only.", "Add formal permission review and automated authorization coverage for every new route."],
    ["Email verification and recovery", "Implemented", "Signed verification, resend, password reset, throttling, and admin controls exist.", "Production SMTP, bounce handling, sender-domain controls, queue monitoring, and delivery metrics remain operational work."],
    ["Applicant exploration", "Implemented", "An applicant may browse before finishing the profile; submission applies the integrity gates.", "Validate with younger applicants and parents that the boundary is understood."],
    ["Adaptive applicant profile", "Implemented with gaps", "Education level, school, grading scale, household, location, preferences, support needs, and guardian context are stored.", "Add program-specific questions, richer non-college pathways, and approved data-minimization rules."],
    ["Guardian-managed minor account", "Partial", "Account manager and guardian details are required for a minor profile.", "Add guardian identity/authority, consent and learner assent, account transfer, dispute, communication, and age-transition rules."],
    ["Applicant profile verification", "Implemented", "Private proof can be uploaded, reviewed by admin, shown as verified, and reset to pending when replaced.", "Define which programs require verification, acceptable proof, review service level, expiry, and appeal/correction policy."],
    ["Provider onboarding and verification", "Implemented", "Separate registration, email verification, organization profile, proof upload, admin decision, and program lock exist.", "Add verification expiry, periodic re-check, registry checks, beneficial ownership where needed, and complaint escalation."],
    ["Program authoring", "Implemented", "Structured identity, target, academics, location, requirements, benefit, stages, events, rubric, renewal, contacts, and terms exist.", "Add cycle versioning, richer benefits, conditional questions/requirements, and more complete post-award terms."],
    ["Program moderation", "Implemented", "Provider submits; admin views a detail page and publishes, rejects, or requests correction. Material changes return to review.", "Add dual approval for high-risk programs, change comparison, review assignment, and review service levels."],
    ["Search and discovery", "Partial", "Published programs, recommendations, saved programs, key program cards, and details exist.", "Add stronger keyword/filter combinations, saved searches, closing-soon calendar, school/service-area filters, and sorting transparency."],
    ["Map and distance", "Partial", "Applicant/program addresses, map pins, Leaflet display, geocoding, and distance visualization exist.", "Add provider service areas, travel mode/time, geocode confidence, offline written fallback, and resilient/self-hosted map options."],
    ["Decision support", "Implemented with limits", "Eligibility and suitability are separated, open criteria do not penalize, grading alternatives are handled, explanations and versioned snapshots exist.", "Validate weights with stakeholders, monitor fairness and error rates, add change explanations, and keep final decisions human."],
    ["Prepared document library", "Implemented", "Common files can be uploaded before applying and reused by requirement type.", "Add issue date, expiry, owner, issuer, version, validity, reminder, aliases, and per-program reuse consent."],
    ["Application wizard and gates", "Implemented", "Selected scholarship, checklist, uploads, terms, email/profile/deadline/duplicate/eligibility/document gates, and history exist.", "Add draft application recovery, autosave, withdrawal, correction window, and clearer exceptional-case handling."],
    ["Application document review", "Implemented", "Provider can view files and accept, reject, or request replacement with notes.", "Add file-version history, reviewer assignment, document-expiry logic, malware quarantine, and review due dates."],
    ["Provider applicant review", "Implemented", "Provider sees the applicant profile, authorized proof, requirements, DSS context, rubric, notes, and decision controls.", "Add committee workflow, conflict declarations, dual scoring, calibration, ranking, and waitlist management."],
    ["Selection stages", "Implemented", "Screening plus optional exam/interview and distribution are configured when the program is authored; providers approve or reject within the plan.", "Add provider-defined stage service levels, exception approval, and clearer stage-specific evidence rules."],
    ["Program and individual schedules", "Implemented", "Shared events can be previewed before applying and propagated to relevant applicants; individual schedules support attendance tracking.", "Add rescheduling requests, calendar export, accommodation request, cancellation escalation, timezone clarity, and no-show policy."],
    ["Distribution notice", "Implemented with limits", "Approved applicants receive distribution details through the portal/email without redundant acknowledgment; provider can record release.", "Add recipient confirmation only if policy truly requires it, benefit-item detail, installment records, proof of release, and reconciliation."],
    ["Notifications", "Implemented with dependency", "Shared bell, unread counts, mark one/all read, deduplication, portal records, queued email, and deadline reminders exist.", "Add delivery/bounce visibility, preferences, digest rules, SMS/push options, escalation, and template governance."],
    ["Administrator oversight", "Implemented", "Provider, program, applicant proof, user account, exports, notification, and paginated activity-log functions exist.", "Add case assignment, supervisory review, complaint cases, retention jobs, and stronger privileged-account controls."],
    ["Applicant mobile app", "Partial parity", "Applicant registration/login, profile, scholarships, saved items, documents, applications, schedules, and notifications are represented.", "Complete proof/document-review parity, accessibility, deep links, push notifications, offline reading, and device testing."],
    ["Analytics and reporting", "Foundation only", "Funnel events, outcome fields, exports, and DSS snapshots exist; analytics pages are intentionally hidden/deferred.", "Approve metric definitions, consent, minimum group sizes, data quality, and governance before exposing dashboards."],
    ["Provider teams", "Not implemented", "A provider account acts as one organizational user.", "Add organization membership, invitations, owner/reviewer/finance/auditor roles, revocation, and action attribution."],
    ["Appeals, corrections, and complaints", "Not implemented", "Notes and replacement requests exist, but there is no complete case workflow.", "Add issue type, deadline, evidence, assigned reviewer, outcome, escalation, and distinction between factual correction and provider discretion."],
    ["Payment and financial reconciliation", "Not implemented", "Awarded amount and distribution status may be recorded; no money moves through the platform.", "Add only after legal and financial review: provider funding account, installment ledger, payment reference, reconciliation, refund, and least-privilege access."],
    ["Renewal management", "Partial", "A renewal policy and recognized renewed status exist.", "Add renewal cycle, period, checklist, academic/enrollment monitoring, provider decision, reminder, probation, and termination history."],
    ["Data-subject rights", "Partial", "Users can edit profiles and admins can export operational records, but a complete self-service rights workflow is absent.", "Add access/export request, correction, deletion/account closure, objection where applicable, retention holds, identity verification, and audit."],
    ["Uploaded-file security", "Partial", "Authorization, private paths, MIME/size validation, and controlled view/download routes exist.", "Add malware scanning, quarantine, safe preview/content disarm, checksum, encryption strategy, access telemetry, and tested deletion."],
    ["Privileged-account security", "Partial", "Role middleware, throttling, suspension, password-reset controls, and last-admin protection exist.", "Add MFA, recovery codes, reauthentication for sensitive actions, session/device controls, alerting, and periodic access review."],
    ["Language and accessibility", "Partial", "Responsive web/mobile interfaces and simplified language are present in many areas.", "Complete WCAG 2.2 AA testing, Filipino/plain-language options, map/upload alternatives, error summaries, and assistive-technology UAT."],
    ["External integrations", "Not implemented", "Email and external map services are the main current integrations.", "Add APIs only after contracts and governance for school records, provider CRM, identity, payments, calendar, or public scholarship feeds."],
    ["Production operations", "Partial", "Environment examples, health route, tests, queues, scheduler, and deployment guidance exist.", "Provision production database/storage, HTTPS, monitoring, backup restore, incident response, capacity tests, release process, and named operators."],
]
add_data_table(
    ["Capability", "Status", "Current coverage", "Missing or required next step"],
    capability_gap_rows,
    [1850, 1450, 3010, 3050],
    font_size=8.0,
)

add_heading("15.3 Scholarship Information and Lifecycle Not Yet Fully Represented", 2)
add_body(
    "The platform already stores a useful scholarship listing, but a mature scholarship-management service needs more "
    "than title, eligibility, amount, deadline, requirements, and a decision. The following catalog identifies material "
    "scholarship information or lifecycle behavior that is absent or only partially structured in the current system."
)
scholarship_gap_rows = [
    ["Program cycle and version", "No complete immutable cycle/version model.", "Cycle name, academic year, version, copied-from program, effective dates, change reason, and archived snapshot.", "High"],
    ["Opening date and exact closing time", "A deadline date exists.", "Opening timestamp, closing timestamp, timezone, rolling/first-come policy, grace period, and late-submission rule.", "High"],
    ["Funding authorization", "Provider verification and award amount exist.", "Funding source, approved budget, authorization evidence, funded/unfunded status, and responsible finance contact kept private where necessary.", "High"],
    ["Benefit components", "One award amount and narrative description are available.", "Separate tuition, allowance, books, device, transport, housing, meals, insurance, mentoring, internship, and other in-kind benefits.", "High"],
    ["Benefit cadence and duration", "Renewal text can describe it.", "One-time/monthly/term/annual frequency, number of installments, maximum duration, covered school periods, and release conditions.", "High"],
    ["Payment recipient and method", "Not structured.", "Applicant, school, vendor, or guardian recipient; payment channel; masked reference; and required confirmation without storing unnecessary banking data.", "Medium"],
    ["Slots, reserves, and quotas", "Total slots exist.", "Reserved categories, waitlist capacity, geographic/school allocation, unused-slot transfer, and transparent lawful quota rules.", "Medium"],
    ["Other-aid compatibility", "Not structured.", "Whether another scholarship is allowed, prohibited combinations, disclosure timing, stacking cap, and conflict resolution.", "High"],
    ["Age and citizenship/residency rules", "Some context can be placed in narrative criteria.", "Structured minimum/maximum age, citizenship/residency category and duration, legal justification, and open/any option.", "Medium"],
    ["Special target or priority groups", "Support needs and narrative rules are available.", "Structured but optional lawful priorities such as disability support, indigenous community, solo-parent household, displacement, disaster, or out-of-school status, with privacy controls.", "Medium"],
    ["Nomination or referral", "Application mode is basic.", "School nomination, provider invitation, referral source, nominee consent, eligibility confirmation, and conversion to an application.", "Medium"],
    ["Custom application questions", "Profile and standard program fields are used.", "Provider questions with type, required/optional state, target education level, validation, privacy purpose, and immutable answer snapshot.", "High"],
    ["Essays, portfolio, audition, or proposal", "Can be requested as a generic document.", "Dedicated prompt, word/file limits, blind-review option, scoring criteria, authorship declaration, and permitted assistance policy.", "Medium"],
    ["Reference or recommendation workflow", "A recommendation file may be uploaded manually.", "Referee invitation, secure link, deadline, confidentiality choice, completion state, reminder, and applicant visibility policy.", "Medium"],
    ["Conditional document requirements", "Requirements are listed by program.", "Rules by age, education level, school type, income route, stage, or applicant response so irrelevant proof is not requested.", "High"],
    ["Document validity", "File date/status are available, but expiry is incomplete.", "Issuer, issue date, expiry date, academic period, owner, certification, accepted alternatives, translation rule, and stale-document reminder.", "High"],
    ["Application draft and amendment", "Submission is supported; prepared files can be replaced.", "Autosaved draft, applicant withdrawal, correction window, material amendment, provider-requested update, and immutable submitted versions.", "High"],
    ["Reviewer assignment", "Provider account reviews its applications.", "Named reviewers, workload, due date, recusal, reassignment, read-only observer, and action ownership.", "Medium"],
    ["Committee and conflict controls", "Rubric and history exist.", "Conflict declaration, dual review, score reconciliation, committee minutes, quorum, override authority, and exception reason.", "High"],
    ["Ranking, tie-break, and waitlist", "DSS and rubric scores exist but do not form a formal ranked award list.", "Provider-controlled ranking policy, tie-break rule, slot allocation, waitlist position/privacy, offer expiry, and movement history.", "High"],
    ["Appeal and factual correction", "Replacement and decision notes exist.", "Appealable issue types, deadline, evidence, independent reviewer, final outcome, and clear boundary around provider discretion.", "High"],
    ["Exam governance", "Duration, passing score, event, schedule, and attendance exist.", "Scope, allowed materials, identity check, accommodation, incident, reschedule, retake, score release, external-provider responsibility, and retention.", "High"],
    ["Interview governance", "Interview stage and schedule exist.", "Question set, scoring anchors, panel members, guardian role, accommodation, recording policy, reschedule, and evidence-based notes.", "Medium"],
    ["Award offer and acceptance", "Student response and contract acceptance fields exist.", "Offer letter, acceptance deadline, decline reason, alternate recipient, agreement version, e-signature policy, and non-digital alternative.", "High"],
    ["Installment and disbursement ledger", "Distribution scheduling and awarded amount exist.", "Installment amount/date, funding source, payee, masked reference, approval, release evidence, failed/reversed payment, and reconciliation.", "High"],
    ["In-kind fulfillment", "Not structured.", "Item inventory, size/specification, vendor, pickup/delivery, recipient confirmation, replacement, and unclaimed-item handling.", "Medium"],
    ["Renewal cycle", "Renewal policy text and status exist.", "Renewal period, opening/closing dates, required evidence, maintained conditions, decision, new amount, and next review.", "High"],
    ["Scholar monitoring", "Outcome notes and statuses are basic.", "Enrollment, attendance, academic progress using original scale, milestone reports, mentoring, internship, and support intervention with proportionate collection.", "Medium"],
    ["Leave, transfer, and deferment", "Can be described in renewal terms.", "Request type, dates, reason category, evidence, approval, adjusted benefit, and return plan.", "Medium"],
    ["Suspension, termination, and recovery", "Narrative terms and negative outcomes exist.", "Policy reason, notice, cure period, final decision, overpayment/recovery amount, waiver, appeal, and audit trail after legal review.", "Medium"],
    ["Return-service compliance", "Contract terms are displayed and accepted.", "Assigned service, supervisor, hours/milestones, evidence, completion, deferral, breach, dispute, and closure.", "Medium"],
    ["Complaints and scam reporting", "Administrator review and contact information exist.", "Report button, case type, evidence, urgency, assigned handler, provider response, resolution, notification, and trend monitoring.", "High"],
    ["Program closure and archive", "Status and deadline exist.", "Closed, cancelled, suspended, exhausted, completed, and archived states; applicant communication; public archive rule; and cycle outcome summary.", "High"],
    ["Accessible and multilingual listing", "Current content is mainly English and provider-authored.", "Plain-language summary, Filipino translation, reading level, accessible document alternative, contact accommodation, and translation-version ownership.", "Medium"],
]
add_data_table(
    ["Missing or partial area", "Current position", "Needed addition", "Priority"],
    scholarship_gap_rows,
    [2100, 2250, 3650, 1360],
    font_size=8.0,
)

add_heading("15.4 Process, Policy, and Operational Gaps", 2)
process_gap_rows = [
    ["Named system owner", "Not established by code", "Assign accountable institution, product owner, technical owner, privacy officer, and escalation contacts.", "Before pilot"],
    ["Privacy inventory and retention schedule", "Partial guidance", "Map purpose, lawful basis, fields, recipients, storage, retention, deletion, backup handling, and data-subject rights.", "Before pilot"],
    ["Minor and guardian policy", "Profile logic only", "Approve age bands, guardian authority, consent/assent, communications, account transfer, and dispute procedure.", "Before pilot"],
    ["Provider verification SOP", "Workflow implemented", "Approve proof standards, independent checks, expiry/re-check, reviewer authority, escalation, and badge wording.", "Before pilot"],
    ["Program review SOP", "Workflow implemented", "Approve content standards, hidden-fee prohibition, harmful criteria review, material-change rule, dual approval, and service level.", "Before pilot"],
    ["Applicant correction/complaint process", "Not complete", "Create cases for factual correction, file review, accessibility, privacy, scam report, and process complaint.", "Before pilot"],
    ["File-security operations", "Partial technical checks", "Quarantine, antivirus/content validation, safe preview, staff procedure, alerting, and incident evidence.", "Before pilot"],
    ["Privileged MFA and access review", "Not complete", "Require MFA, recovery, reauthentication, session controls, quarterly access review, and emergency administrator continuity.", "Before pilot"],
    ["Backup and disaster recovery", "Guidance only", "Set recovery objectives, encrypted independent backup, file/database consistency, restoration runbook, and recurring drills.", "Before pilot"],
    ["Mail, queue, scheduler, and map monitoring", "Runtime dependent", "Supervise workers, monitor failures and latency, alert named staff, retry safely, and provide written fallback.", "Before pilot"],
    ["Accessibility and low-bandwidth acceptance", "Partial", "Run WCAG 2.2 AA, keyboard, screen-reader, zoom, mobile, slow-network, and assisted-use testing with real users.", "Before pilot"],
    ["DSS validation and governance", "Versioned logic exists", "Approve criteria definitions and weights, test false blocks/open values/scales, monitor fairness, publish explanation, and review changes.", "Pilot"],
    ["Provider service model", "Strategy only", "Define free core, optional provider services, fee basis, waivers, invoicing, support level, and measurable provider value.", "Pilot"],
    ["User support and training", "Not formalized", "Create applicant/guardian help, provider onboarding, admin SOP, service hours, escalation, and issue knowledge base.", "Pilot"],
    ["Data quality ownership", "Events/snapshots exist", "Define metrics, valid values, missing-data rules, deduplication, correction responsibility, and recurring quality review.", "Pilot"],
    ["Release and change management", "Development process", "Use approved releases, migration rehearsal, rollback, change log, policy sign-off, and post-release checks.", "Before hosting"],
    ["Legal and financial review", "Disclaimers only", "Review terms, provider contracts, minors, fees, taxes, disbursement, consumer disclosures, and data-sharing before real transactions.", "Before payments"],
]
add_data_table(
    ["Gap", "Current position", "Required action", "When"],
    process_gap_rows,
    [2000, 1750, 4250, 1360],
    font_size=8.0,
)

add_heading("15.5 Features Intentionally Deferred or Not Recommended Yet", 2)
add_bullets([
    "Do not add a platform-hosted online examination engine yet. Providers currently administer examinations externally while the platform records the plan, schedule, attendance, and outcome.",
    "Do not add automatic grade scanning or OCR now. It was explicitly deferred; any future extraction must preserve the source, show the value to the applicant, require confirmation, and avoid automatic final decisions.",
    "Do not let the DSS approve, reject, or rank applicants secretly. It should explain fit and support screening while an authorized provider makes the decision.",
    "Do not charge applicants for discovery, matching, required verification, application submission, or better placement. Any sustainability fee should normally purchase an optional provider service.",
    "Do not expose applicant documents through public links, provider-wide browsing, or analytics exports without a defined purpose and authorization.",
    "Do not launch predictive analytics, payment processing, or broad integrations before privacy, data quality, security, appeal, monitoring, and ownership controls are operating.",
    "Do not add social-feed or gamification features merely to increase activity; prioritize accurate opportunities, completion, trust, accessibility, and award outcomes." ,
])


chapter("16. Prioritized Future Features and Improvements", "Part V | Strategy and Future Development")
add_heading("16.1 Priority 0: Before a Controlled Pilot", 2)
add_data_table(
    ["Improvement", "Reason", "Completion signal"],
    [
        ["Final privacy, retention, deletion, complaints, and guardian policies", "Current controls need approved operating rules.", "Policies published, accepted, mapped to fields/tables, and tested."],
        ["Malware scanning and upload quarantine", "Scholarship proof is a high-risk file surface.", "Unsafe files are blocked, clean files released, failures logged, and staff trained."],
        ["MFA for provider and administrator accounts", "Privileged account compromise has broad impact.", "Enrollment, recovery, enforcement, and audit tests pass."],
        ["Backup and restoration drill", "Hosting readiness requires recoverability, not only backup creation.", "Database and storage restored within approved objectives."],
        ["Accessibility and low-bandwidth audit", "Applicants use varied devices and abilities.", "Critical WCAG 2.2 AA issues resolved and slow-network flows tested."],
        ["Provider verification and program-review SOP", "Trust decisions must be repeatable.", "Review checklist, escalation, service level, and evidence standards approved."],
        ["Production monitoring", "Mail, queues, scheduler, storage, and errors can otherwise fail silently.", "Alerts, owners, runbooks, and incident test are operational."],
    ],
    [3300, 3100, 2960],
    font_size=8.4,
)

add_heading("16.2 Priority 1: Pilot Learning and Usability", 2)
add_bullets([
    "Dedicated guardian-managed account workflow with consent, account ownership, transfer, and age-appropriate explanations.",
    "Scholarship templates by target group so elementary, junior high, senior high, TVET, ALS, and college programs request appropriate fields.",
    "Document expiry, issue date, replacement reminders, and requirement aliases so prepared files remain accurate across programs.",
    "Provider trust profile with verification date, review scope, official contacts, complaint channel, and clearly bounded badge meaning.",
    "Improved search, saved filters, deadline calendar, location radius, travel mode, and written-address fallback.",
    "Appeal or correction workflow for factual errors, document review, and process complaints without promising an appeal of provider discretion where none exists.",
    "Mobile push notifications with preference controls and a reliable deep link into the relevant application or schedule.",
    "Multilingual and plain-language content, beginning with Filipino and English terminology validated by users.",
    "Accessibility improvements across keyboard, screen reader, focus, contrast, reflow, error summary, file upload, and maps." ,
])

add_heading("16.3 Priority 2: Provider Operations and Scale", 2)
add_bullets([
    "Team accounts and scoped provider permissions for program owners, reviewers, finance staff, and read-only auditors.",
    "Conflict-of-interest declarations, reviewer assignment, dual review, and calibration for competitive scholarships.",
    "Program cloning and cycle versioning so criteria changes do not rewrite old applications.",
    "Structured provider service levels and automatic escalation for overdue proof, document, and application decisions.",
    "Secure integrations with school, identity, CRM, or payment systems only after data-sharing agreements and purpose review.",
    "E-signature or agreement acknowledgment for award terms, with legal review and a non-digital alternative.",
    "Disbursement tracking and reconciliation without exposing full banking information to unnecessary roles.",
    "Provider impact dashboards only after metrics, governance, minimum group size, and data-quality rules are mature.",
    "Institutional exports and APIs with scopes, rate limits, audit, revocation, and retention controls." ,
])

add_heading("16.4 Priority 3: Advanced Research and Ecosystem Features", 2)
add_bullets([
    "Outcome follow-up for persistence, completion, employment, service, or other program-specific goals using voluntary and proportionate data.",
    "Fairness and access evaluation across geography, education level, disability, income, and other approved dimensions.",
    "Recommendation personalization only after enough high-quality data, with transparent controls and no hidden eligibility decisions.",
    "Fraud and duplicate detection with human review, careful false-positive handling, and appeal or correction.",
    "Self-hosted or resilient map services, school and community reference data, and offline-first mobile reading.",
    "OCR or grade-document extraction may be explored later, but it is explicitly deferred for now. Any future use must preserve the original file, show extracted values to the applicant, require confirmation, and never make an unreviewed decision.",
    "Public scholarship data exchange standards so providers can publish once and distribute accurate structured listings across approved channels." ,
])

add_heading("16.5 Recommended Roadmap Order", 2)
add_numbered([
    ("Trust and safety", "Policies, provider review, document security, privileged authentication, and incident response."),
    ("Operational reliability", "Backups, queues, scheduler, email, monitoring, and hosting."),
    ("Applicant usability", "Age-appropriate profiles, documents, search, accessibility, language, and mobile parity."),
    ("Provider efficiency", "Team workflow, review calibration, templates, service levels, and program cycle management."),
    ("Measurement", "Data quality and definitions before analytics pages or predictive features."),
    ("Automation", "Only automate high-volume tasks after the policy, data, exception, and appeal processes are stable."),
])


chapter("17. Research, Evaluation, and Data Analysis", "Part V | Strategy and Future Development")
add_heading("17.1 Evaluation Questions", 2)
add_bullets([
    "Can applicants and parents find a relevant scholarship without assistance?",
    "Do users understand why a scholarship is or is not a strong fit?",
    "Does progressive onboarding improve exploration without lowering application completeness?",
    "Does the prepared-document library reduce repeated upload work and provider replacement requests?",
    "Do providers review applications faster and more consistently than through email or social media?",
    "Do administrator verification and program review improve applicant trust?",
    "Are match explanations accurate across open criteria, different grading scales, education levels, and locations?",
    "Are there systematic differences in discovery, completion, review, or award outcomes that require design or policy changes?",
    "Which platform functions create enough provider value to support sustainable funding?" ,
])

add_heading("17.2 Applicant Measures", 2)
add_bullets([
    "Time to first suitable program, profile completion, document readiness, saved programs, application start and completion.",
    "Task success, error rate, time on task, satisfaction, trust, perceived clarity, and willingness to recommend.",
    "Understanding of match guidance, provider final authority, requirements, schedules, and obligations.",
    "Drop-off by profile step and application step, with careful interpretation rather than blaming the user.",
    "Access by device, connectivity, education level, age group, location, and guardian-managed status." ,
])

add_heading("17.3 Provider Measures", 2)
add_bullets([
    "Verification completion, time to first program, draft-to-publication time, and review return rate.",
    "Application completeness, document replacement rate, median review time, and staff time per applicant.",
    "Use of rubric, reasons, schedules, and notifications; no-show and communication-failure rates.",
    "Qualified applicant reach, geographic reach, award fill rate, renewal, and provider retention.",
    "Provider satisfaction, perceived trust, operating cost, and willingness to pay for specific time-saving features." ,
])

add_heading("17.4 DSS Evaluation", 2)
add_bullets([
    "Criterion-level accuracy against manually checked scholarship rules.",
    "False blocker rate and false strong-match rate.",
    "Consistency across grading scales and open-criteria wording.",
    "Applicant comprehension of explanation and next action.",
    "Provider agreement and disagreement with suitability, with reasons rather than treating disagreement as model error automatically.",
    "Outcome and fairness comparison across approved groups, with minimum group-size and privacy controls.",
    "Reproducibility of the same snapshot and methodology version." ,
])

add_heading("17.5 Study Design", 2)
add_numbered([
    ("Define participants", "Include learners across education levels, parents or guardians, small providers, reviewers, and administrators."),
    ("Use realistic tasks", "Register, explore, complete relevant profile sections, prepare a file, understand a match, apply, review evidence, schedule, and decide."),
    ("Observe behavior", "Record task success and confusion with consent; do not rely only on self-reported satisfaction."),
    ("Interview after tasks", "Ask what information was missing, which terms were unclear, and what created or reduced trust."),
    ("Analyze by role and context", "Separate mobile/desktop, education level, guardian status, connectivity, and provider size."),
    ("Prioritize evidence", "Fix safety, blocking errors, and misunderstanding before cosmetic preferences."),
    ("Repeat", "Re-test changed workflows and preserve version information."),
])

add_callout(
    "Analytics timing",
    "The system already records useful funnel events and DSS snapshots, but full analytics pages should remain deferred until event definitions, privacy rules, data quality, minimum group sizes, and decision ownership are approved. Measurement infrastructure should mature before polished dashboards.",
    "info",
)


chapter("18. Conclusions and Recommended Position")
add_body(
    "The platform should be presented as a scholarship access and operations system, not merely a list of links. Its "
    "value comes from connecting discovery, adaptive applicant information, reusable documents, transparent matching, "
    "verified providers, reviewed programs, structured applications, human decisions, schedules, notifications, and "
    "audit evidence. The process is deliberately staged because different controls answer different trust questions."
)
add_body(
    "Applicants and families should use it because it reduces fragmented searching, repeated preparation, hidden "
    "requirements, and missed updates. Providers should use it because it offers matched reach, cleaner intake, "
    "consistent review, communication, evidence, and future impact reporting that a social-media post alone cannot "
    "provide. Administrators and institutions should support it because provider verification, program moderation, "
    "role separation, and traceability protect the scholarship ecosystem."
)
add_body(
    "The next phase should prioritize trust, privacy, file security, privileged-account protection, recovery, "
    "accessibility, and real-user validation. Advanced analytics, integrations, payment, OCR, and predictive features "
    "should follow only after the underlying process, governance, and data quality are mature."
)
add_callout(
    "Recommended defense statement",
    "The system does not replace scholarship providers or guarantee awards. It reduces information and process barriers, helps applicants identify and prepare for suitable opportunities, gives providers a structured and auditable workflow, and keeps final scholarship decisions with authorized human reviewers.",
    "success",
)


# ---------------------------------------------------------------------------
# Appendices
# ---------------------------------------------------------------------------

chapter("Appendix A. Master Scholarship Listing Template")
add_body(
    "Use this checklist when creating, reviewing, importing, or auditing a scholarship listing. Fields marked as "
    "conditional should appear only when they apply to the target group or program process."
)
master_fields = [
    ["Program identity", "Official title; short title; category; program cycle; provider; logo; public status; version; source or authority."],
    ["Purpose", "Problem addressed; intended learner outcome; target population; geographic or institutional context."],
    ["Provider identity", "Organization name and type; verified contact; address; website; representative; verification date; complaint channel."],
    ["Education target", "Preschool, elementary, JHS, SHS, ALS, TVET, college, graduate, professional, or mixed; eligible schools and enrollment status."],
    ["Academic target", "Grade/year level; course; track/strand; school type; grading scale; minimum value; alternatives such as competency or portfolio."],
    ["Financial target", "Income limit or bracket; evidence; household factors; priority circumstances; whether financial need is mandatory or ranked."],
    ["Location", "Eligible region/province/city/barangay/school; nationwide/open rule; provider address; study, exam, interview, and distribution locations; map pin."],
    ["Other criteria", "Age; citizenship/residency where lawful; talent; leadership; service; disability support; affiliation; availability; relocation."],
    ["Benefit", "Cash amount; currency; frequency; duration; covered costs; in-kind items; slot count; direct-to-school payment; exclusions; funding limits."],
    ["Application", "Opening and closing date/time/timezone; application mode; nomination or open entry; late policy; duplicate rule; contact."],
    ["Documents", "Standard type; description; mandatory/optional; stage requested; accepted format/size; issue/expiry rule; replacement process."],
    ["Selection", "Screening; examination; interview; rubric; reviewer roles; conflict management; ranking; tie handling; final authority."],
    ["Examination", "Conditional: purpose, scope, duration, passing/ranking method, date, venue or method, identity check, accommodation, retake."],
    ["Interview", "Conditional: format, criteria, date, venue or method, expected duration, guardian role, accommodation, rescheduling."],
    ["Timeline", "Screening window; document replacement deadline; exam/interview; decision; orientation; distribution; renewal."],
    ["Award conditions", "Acceptance deadline; agreement; enrollment confirmation; attendance; grades; reporting; conduct; publicity; combination with other aid."],
    ["Return service", "Conditional: service type, duration, location, supervisor, proof, deferral, breach consequence, dispute process."],
    ["Renewal", "Eligibility; documents; review date; available funds; probation; leave/transfer; termination; maximum duration."],
    ["Privacy", "Purpose; recipients; processors; retention; correction; deletion; complaints; minor/guardian handling; consent version."],
    ["Communications", "Official email/phone; notification channels; response time; language; accessibility; emergency update method."],
    ["Review and publication", "Provider approval; admin reviewer; submitted/reviewed/published dates; notes; material-change rule; expiry/archive date."],
]
add_data_table(["Section", "Complete information"], master_fields, [2100, 7260], font_size=8.4)


chapter("Appendix B. Applicant Profile and Prepared-Document Reference")
add_heading("B.1 Applicant Profile Information", 2)
profile_rows = [
    ["Account", "Email, username, password hash, contact number, role, email verification, suspension, password controls, terms acceptance."],
    ["Identity", "First, middle, last, suffix, birth date, gender where appropriate, and identity/proof status."],
    ["Education", "Education level, school, school type, enrollment status, grade/year, track/strand/course/training, learner reference only when needed."],
    ["Academic", "Original grading scale and value, ranking/competency/portfolio alternative, marking period, and supporting record."],
    ["Household", "Income bracket or provider-required value, household size, dependents, and relevant support needs."],
    ["Location", "Address, city/municipality, province, region, postal code, map pin, and willingness to relocate where relevant."],
    ["Preferences", "Education goal, preferred fields, program type, location preference, award need, and accessibility/support needs."],
    ["Guardian", "Name, relationship, contact, email, account owner/manager, consent or authorization, and control-transfer context."],
    ["Verification", "Private proof file, type, status, review reason, reviewer, timestamps, and replacement history."],
]
add_data_table(["Profile group", "Data and purpose"], profile_rows, [2050, 7310], font_size=8.7)

add_heading("B.2 Prepared Document Rules", 2)
add_bullets([
    "Show common document types as ready-to-upload cards rather than forcing the applicant to type a document name.",
    "Keep provider-specific forms within the selected program's application checklist.",
    "Show file date, status, programs that can use it, expiry where applicable, and replacement action.",
    "Allow the applicant to replace a file before submission and preserve review history after submission.",
    "Do not automatically reuse a document whose type, owner, date, or validity does not satisfy the new requirement.",
    "Never make a private document publicly discoverable or directly linkable without authorization." ,
])


chapter("Appendix C. Provider and Program Review SOP")
add_heading("C.1 Provider Verification Decision", 2)
add_numbered([
    ("Open the provider detail page", "Confirm the signed-in reviewer is authorized and the account is not being reviewed through an unsafe export."),
    ("Compare identity", "Match organization name, type, address, website, representative, and contact across profile and proof."),
    ("Independently confirm", "Use an official source or separately obtained contact for high-risk or unfamiliar providers."),
    ("Inspect proof", "Check readability, issuer, dates, consistency, alteration indicators, and whether the proof actually supports authority."),
    ("Assess risk", "Escalate payment requests, guaranteed awards, unverifiable contacts, copied identity, duplicate accounts, or unusual data requests."),
    ("Decide", "Approve, reject, or request replacement. Record a concise reason and next action."),
    ("Notify", "Send the outcome without attaching the proof or disclosing unnecessary sensitive details."),
    ("Retain", "Keep the record only for the approved retention period and restrict subsequent access."),
])

add_heading("C.2 Program Review Decision", 2)
add_numbered([
    ("Confirm provider status", "Only an approved, active provider may proceed to publication."),
    ("Check purpose and target", "Ensure the intended learner and program objective are understandable."),
    ("Check criteria consistency", "Compare structured fields, description, documents, stages, and DSS implications."),
    ("Check benefit and funding claims", "Confirm amount, frequency, duration, slots, exclusions, and provider responsibility."),
    ("Check applicant burden", "Remove unnecessary fields or proof, especially for minors and early-stage applicants."),
    ("Check process", "Review deadline, exam/interview details, distribution, renewal, service, complaint, and contact."),
    ("Check trust and safety", "Reject misleading guarantees, hidden fees, excessive data, unsafe contact instructions, or discriminatory conditions."),
    ("Publish or return", "Record approval or a specific correction note. Material later changes should return to review."),
])

add_heading("C.3 Applicant Proof Review", 2)
add_bullets([
    "Review only proof relevant to profile verification, not every private document by default.",
    "Confirm the applicant and proof correspond without copying full identifiers into notes.",
    "Approve, reject, or request replacement with an understandable reason.",
    "If proof changes, remove the verified state or mark it pending until the replacement is reviewed.",
    "Providers may view applicant proof only when the applicant applied to a program owned by that provider and the proof is relevant to review." ,
])


chapter("Appendix D. Application and Review Status Dictionary")
status_rows = [
    ["submitted", "Application passed submission gates and entered provider screening.", "Review eligibility and documents."],
    ["under_review / screening", "Provider is checking profile, eligibility, and evidence.", "Wait or respond to a replacement request."],
    ["qualified", "Basic eligibility and required evidence are acceptable.", "Proceed to the next configured stage."],
    ["shortlisted", "Applicant remains under competitive consideration.", "Review schedule and provider instructions."],
    ["exam_qualified", "Applicant may take the program's configured examination.", "Wait for or review exam details."],
    ["exam_scheduled", "Examination date/time/location or method has been set.", "Acknowledge and attend as instructed."],
    ["exam_taken", "Provider recorded completion; result is pending.", "Wait for result."],
    ["exam_passed", "Examination condition was satisfied.", "Proceed to interview or final review."],
    ["exam_failed", "Examination condition was not satisfied under provider rules.", "Review provider outcome and any permitted correction or retake policy."],
    ["interview", "Applicant is in the interview stage or has an interview schedule.", "Review and acknowledge details; attend."],
    ["approved", "Provider approved the application subject to disclosed award steps.", "Review acceptance, agreement, and distribution information."],
    ["awarded", "Applicant is recorded as an award recipient.", "Complete remaining award or distribution requirements."],
    ["distribution_scheduled", "Benefit release or orientation has been scheduled.", "Review the date, address, identification, and provider instructions; no portal acknowledgment is required."],
    ["disbursed", "Provider recorded release of the benefit.", "Retain confirmation and follow renewal or monitoring terms."],
    ["renewed", "Provider approved support for another period.", "Review updated terms and schedule."],
    ["rejected / not_awarded", "Provider made a negative final decision or slots were not awarded.", "Read the reason or note and use the record for future applications."],
]
add_data_table(["Status", "Meaning", "Expected next action"], status_rows, [2300, 3900, 3160], font_size=8.3)


chapter("Appendix E. Role Access Matrix")
access_rows = [
    ["Browse public scholarships", "Yes", "Yes", "Yes"],
    ["Manage applicant profile/documents", "Own only", "No", "Verification view only"],
    ["Create and edit programs", "No", "Own only after verification", "Review; no provider authorship by default"],
    ["Publish programs", "No", "Submit for review", "Yes after review"],
    ["View applicant application", "Own only", "Applicants to owned programs", "Oversight/review under policy"],
    ["View applicant proof", "Own only", "Only applicant to owned program", "Verification role"],
    ["Review application files", "View own status", "Owned program", "Oversight where authorized"],
    ["Set application decision", "No", "Owned program; final provider authority", "No routine provider decision"],
    ["Verify provider", "No", "Submit proof", "Yes"],
    ["Verify applicant proof", "Submit proof", "No independent browsing", "Yes"],
    ["Manage all users", "No", "No", "Yes with last-admin safeguards"],
    ["View logs", "No", "Own relevant history only", "Yes under policy"],
    ["Use applicant mobile API", "Yes", "No", "No"],
]
add_data_table(["Capability", "Applicant", "Provider", "Administrator"], access_rows, [3200, 1800, 2200, 2160], font_size=8.4)


chapter("Appendix F. Deployment and Operations Checklist")
add_heading("F.1 Before Deployment", 2)
add_bullets([
    "Review source control status and deploy an approved commit or release artifact.",
    "Replace all demo passwords, remove or label demo data, and verify no real secret is committed.",
    "Create production environment variables and validate dotenv syntax, quoting values that contain spaces.",
    "Confirm PHP extensions, Composer, Node build output, database driver, file permissions, and storage paths.",
    "Configure domain, DNS, HTTPS certificate, secure cookies, trusted proxies, and canonical application URL.",
    "Configure MySQL, least-privilege user, migration backup, connection limits, and timezone.",
    "Configure SMTP sender, verification link domain, queue worker, scheduler, and monitoring.",
    "Run test suite, production asset build, database migration rehearsal, and vulnerability/dependency review." ,
])

add_heading("F.2 After Deployment", 2)
add_bullets([
    "Open the health endpoint and landing, login, registration, terms, and role dashboards over HTTPS.",
    "Register and verify a test applicant and provider using controlled addresses.",
    "Test provider proof, admin verification, draft program, admin review, publication, applicant application, file review, schedule, approval, and distribution notice.",
    "Confirm bell counts, individual and all-read actions, queued email, failed-job monitoring, and deadline scheduler idempotency.",
    "Verify unauthorized direct file URLs and cross-provider application access fail.",
    "Create and restore a backup in isolation; compare record and file counts.",
    "Check browser, mobile device, keyboard, screen reader, zoom, and low-bandwidth behavior.",
    "Record release version, migration batch, owner, rollback plan, known issues, and approval." ,
])

add_heading("F.3 Recurring Operations", 2)
add_bullets([
    "Daily: health, errors, queue, scheduler, failed jobs, disk, database, mail, and security alerts.",
    "Weekly: pending review age, bounced mail, provider complaints, storage growth, dependency notices, and sample access-log review.",
    "Monthly: restore test or rotating restoration exercise, account review, permission review, retention jobs, accessibility issues, and provider service metrics.",
    "Per cycle: archive expired programs, preserve immutable historical applications, update criteria versions, and review scholarship information against official provider rules.",
    "Annually or after major change: privacy impact review, penetration test, disaster-recovery exercise, policy review, and DSS fairness/accuracy assessment." ,
])


chapter("Appendix G. End-to-End Demonstration and Testing Guide")
add_callout(
    "Demo account warning",
    "The default seeded accounts and scholarships are fictional. Change all passwords and remove or clearly isolate demo records before public hosting. Never claim that a fictional provider or program is accepting real applications.",
    "warning",
)
add_heading("G.1 Seeded Demonstration Accounts", 2)
add_data_table(
    ["Role", "Email", "Username/password note"],
    [
        ["Administrator", "admin@scholarship.test", "Demo username may be admin; the seeded default password is password123 unless the seeder was changed."],
        ["Applicant", "student@scholarship.test", "Demo username may be student; the seeded default password is password123 unless the seeder was changed."],
        ["Provider 1", "tulayaral@scholarship.test", "Fictional community provider; seeded password is password123 unless changed."],
        ["Provider 2", "bukasfoundation@scholarship.test", "Fictional community provider; seeded password is password123 unless changed."],
    ],
    [1900, 3000, 4460],
    font_size=8.8,
)
add_body(
    "Older project notes may mention role-name passwords such as admin, student, tulayaral, or bukasfoundation. The "
    "current seeder should be treated as the source of truth. Confirm the actual local credentials immediately before "
    "a demonstration rather than publishing uncertain passwords.",
    italic=True,
)

add_heading("G.2 Full Role Test", 2)
add_numbered([
    ("Public", "Open the landing page, scholarship carousel, terms modals, applicant registration, provider registration, login, and responsive navigation."),
    ("Applicant registration", "Create an applicant, verify email or use an approved demo verification path, explore while profile is incomplete, and confirm the resend control disappears after verification."),
    ("Applicant profile", "Test a minor/basic-education profile and an adult college profile; verify conditional fields, common selects, map pin/address, proof upload, and replacement state."),
    ("Applicant documents", "Upload common files from the visible document cards, replace a file, and confirm number-summary cards are absent."),
    ("Scholarship discovery", "Open recommendations, verify consistent card alignment and logo placement, save a program, view details, map/distance, stages, and obligations."),
    ("Application", "Start from a selected scholarship, upload missing requirements in the checklist, accept terms, submit, and verify duplicate/incomplete/deadline/eligibility gates."),
    ("Provider verification", "Register a provider, verify email, upload organization proof, confirm program creation is locked until admin approval, then approve from admin."),
    ("Program review", "Create a draft with a target education level, grading alternative, documents, logo, location, stages, exam/interview details, distribution, and terms; submit; publish or return from admin."),
    ("Provider application review", "Open the applicant detail page, inspect profile/proof, view each file in the modal or page, accept/reject/request replacement, use rubric, and record a reasoned decision."),
    ("Schedules", "Post a general examination/interview/distribution event and an individual schedule; verify applicants can preview planned dates, acknowledge exam/interview when requested, and cannot redundantly acknowledge distribution."),
    ("Notifications", "Confirm unread count in applicant, provider, and admin layouts; open one, mark read, mark all read, and verify queued email where configured."),
    ("Account controls", "Test suspension, reactivation, forced password reset, email verification control, and last-admin protection."),
    ("Mobile", "Sign in as applicant, verify provider/admin rejection, review scholarships, saved programs, profile, documents, application, schedules, and notifications."),
    ("Security", "Attempt cross-role routes, cross-provider application/file access, expired token use, invalid upload, CSRF failure, throttled login, and direct private file access."),
    ("Operations", "Run tests and build, trigger queue/scheduler, inspect failed jobs and logs, and restore a backup in a non-production environment."),
])


chapter("Appendix H. Risk Register")
risk_rows = [
    ["Fraudulent provider or unfunded program", "High", "Provider proof, independent review, program approval, complaints, rapid suspension, evidence retention."],
    ["Sensitive document exposure", "High", "Private storage, authorization, no-store, encryption, malware scan, access log, retention/deletion, incident response."],
    ["Privileged account compromise", "High", "MFA, strong recovery, least privilege, session controls, monitoring, administrator continuity safeguards."],
    ["Incorrect eligibility advice", "High", "Structured criteria, open-value normalization, explanations, snapshots, provider rule link, human final decision, regression tests."],
    ["Minor or guardian misuse", "High", "Guardian workflow, age-appropriate notice, account ownership, consent/assent, minimal data, transfer and dispute policy."],
    ["File malware", "High", "Quarantine, content validation, antivirus, safe preview, restricted download, staff training."],
    ["Mail or scheduler failure", "Medium", "Supervised workers, health checks, alerts, failed-job handling, idempotent reminders, visible in-web notifications."],
    ["Database/storage loss", "High", "Encrypted backups of both, independent copy, restore drill, recovery objectives, deployment rollback."],
    ["Provider abandonment", "Medium", "Assisted onboarding, service levels, time-saving metrics, repeatable cycles, provider support and feedback."],
    ["Applicant overload", "Medium", "Progressive disclosure, concise pages, age-appropriate forms, plain language, consistent design, actionable next step."],
    ["Biased selection or DSS", "High", "Criteria justification, fairness testing, score transparency, human review, conflict controls, aggregate monitoring."],
    ["Third-party map/mail outage", "Medium", "Written-address fallback, retry, alternate provider, status communication, vendor governance."],
    ["Unsustainable operating cost", "Medium", "Pilot metrics, ethical provider pricing, institutional funding, automation only after process maturity."],
    ["Scope confusion", "Medium", "Clear disclosures that exams, final decisions, contracts, and payments are provider-managed."],
]
add_data_table(["Risk", "Severity", "Primary controls"], risk_rows, [2800, 1200, 5360], font_size=8.4)


chapter("Appendix I. Glossary")
glossary_rows = [
    ["Applicant", "Learner seeking scholarship support; may be assisted or represented by a parent or guardian under policy."],
    ["Provider", "Organization that funds or administers a scholarship and makes the final selection decision."],
    ["Administrator", "Platform role responsible for verification, moderation, user controls, logs, and governance."],
    ["Eligibility", "Mandatory conditions that determine whether an applicant may proceed."],
    ["Selection criterion", "Factor used to compare eligible applicants when awards or slots are limited."],
    ["DSS", "Decision Support System; an explainable aid that organizes fit information without replacing human authority."],
    ["Suitability score", "Current weighted representation of eligibility, academic, and financial-need alignment."],
    ["Document readiness", "Whether the files required by a scholarship are prepared; shown separately from suitability."],
    ["Screening", "Initial provider review of eligibility, profile, and evidence."],
    ["Prepared document", "Reusable private applicant file uploaded before or during applications."],
    ["Application document", "A file attached to a specific program requirement and reviewed in that application."],
    ["Verification", "Platform review of submitted proof under policy; not a guarantee of every future action or claim."],
    ["Publication review", "Administrator check of a provider-submitted program before it becomes publicly available."],
    ["Program event", "Schedule or announcement shared for a scholarship stage such as exam, interview, or distribution."],
    ["Application schedule", "Applicant-specific schedule associated with one application."],
    ["Status history", "Chronological record of application status changes and reasons."],
    ["Audit log", "Record of an action, actor, time, target, and relevant technical context."],
    ["Snapshot", "Stored DSS inputs, version, hash, and output representing the calculation at a specific time."],
    ["RBAC", "Role-Based Access Control; permissions based on applicant, provider, or administrator responsibility."],
    ["Queue worker", "Background process that handles jobs such as queued notification email."],
    ["Scheduler", "Laravel process trigger used for timed tasks such as deadline reminders."],
    ["Data minimization", "Collecting only information necessary for a specific legitimate purpose."],
    ["Retention", "Approved period for keeping a record before archive or deletion."],
    ["WCAG", "Web Content Accessibility Guidelines used to evaluate accessible digital experiences."],
]
add_data_table(["Term", "Meaning"], glossary_rows, [2300, 7060], font_size=8.7)


chapter("Appendix J. Official References and Further Reading")
add_body(
    "Official criteria and laws may change. Always verify the current program cycle and applicable legal guidance. "
    "The sources below support the policy context used in this handbook; they do not certify fictional demo records."
)
add_source(
    "1",
    "National Privacy Commission - Republic Act No. 10173, Data Privacy Act of 2012",
    "https://privacy.gov.ph/data-privacy-act/",
    "Primary reference for Philippine personal-data principles and obligations.",
)
add_source(
    "2",
    "National Privacy Commission - Data Privacy Act and Implementing Rules and Regulations",
    "https://privacy.gov.ph/the-data-privacy-act-and-its-irr/",
    "Use for detailed privacy governance and current Commission guidance.",
)
add_source(
    "3",
    "UniFAST - History and Mandate",
    "https://unifast.gov.ph/uni-hist.html",
    "Explains the harmonization of government-funded student financial assistance modalities in the Philippines.",
)
add_source(
    "4",
    "UniFAST - Official Program and Guidelines Portal",
    "https://www.unifast.gov.ph/guidelines.html",
    "Illustrates why current program-specific guidance must be checked for each cycle.",
)
add_source(
    "5",
    "Commission on Higher Education Regional Office V - Grants Portal",
    "https://grants.r5.ched.gov.ph/",
    "An official example of a profile, browse, application, and update-oriented scholarship portal.",
)
add_source(
    "6",
    "DOST-SEI - 2026 Undergraduate Scholarship Primer",
    "https://science-scholarships.ph/pdf/2026_DOST-SEI_UG_Primer.pdf",
    "Example of an official provider primer whose cycle-specific criteria should remain authoritative.",
)
add_source(
    "7",
    "Department of Education - Learner Information System Support",
    "https://support.lis.deped.gov.ph/support/",
    "Includes privacy reminders relevant to learner identifiers and education records.",
)
add_source(
    "8",
    "World Wide Web Consortium - Web Content Accessibility Guidelines Overview",
    "https://www.w3.org/WAI/standards-guidelines/wcag/",
    "Primary accessibility reference for perceivable, operable, understandable, and robust experiences.",
)
add_source(
    "9",
    "Lawphil - Republic Act No. 11765, Financial Products and Services Consumer Protection Act",
    "https://lawphil.net/statutes/repacts/ra2022/ra_11765_2022.html",
    "Reference for fair treatment, transparency, disclosure, and responsible pricing when regulated financial services or payments are involved.",
)
add_source(
    "10",
    "Bangko Sentral ng Pilipinas - Guidelines on Pricing of Electronic Payments",
    "https://www.bsp.gov.ph/Regulations/Issuances/2024/M-2024-015.pdf",
    "Supports transparent, fair, competitive, and cost-justified electronic-payment pricing; obtain professional review for the platform's actual payment model.",
)

add_heading("Project Sources Reviewed", 2)
add_bullets([
    "Laravel routes, controllers, middleware, models, policies, notifications, observers, migrations, seeders, console scheduling, and tests.",
    "Vue components and role layouts for public, applicant, provider, and administrator experiences.",
    "Flutter applicant mobile application and Laravel mobile API behavior.",
    "README, environment examples, hosting guidance, local SQLite schema, and current automated test/build results." ,
])

add_callout(
    "Maintenance note",
    "Update this handbook whenever program statuses, required profile fields, DSS weights or methodology version, verification policy, notification channels, database schema, terms version, hosting architecture, or mobile scope changes.",
    "info",
)


# Final document settings and save.
for section in doc.sections:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

settings = doc.settings._element
update_fields = settings.find(qn("w:updateFields"))
if update_fields is None:
    update_fields = OxmlElement("w:updateFields")
    settings.append(update_fields)
update_fields.set(qn("w:val"), "true")

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(r"D:\XAMPP\htdocs\scholarship_documents\Scholarship_Providers_Complete_Guide_Motivations_Operations_and_Platform_Value.docx")

NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
GOLD = "B88900"
INK = "1F2937"
MUTED = "5B677A"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"
BORDER = "CBD5E1"
POSITIVE = "166534"
CAUTION = "7A5A00"
RISK = "9B1C1C"


def rgb(hex_value):
    return RGBColor.from_string(hex_value)


def set_run_font(run, size=11, color=INK, bold=False, italic=False, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    run.bold = bold
    run.italic = italic


def set_spacing(paragraph, before=0, after=6, line=300):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line / 240
    ppr = paragraph._p.get_or_add_pPr()
    spacing = ppr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        ppr.append(spacing)
    spacing.set(qn("w:before"), str(before * 20))
    spacing.set(qn("w:after"), str(after * 20))
    spacing.set(qn("w:line"), str(line))
    spacing.set(qn("w:lineRule"), "auto")


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_tokens = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.keep_together = True

    if "Provider Callout" not in styles:
        callout_style = styles.add_style("Provider Callout", WD_STYLE_TYPE.PARAGRAPH)
    else:
        callout_style = styles["Provider Callout"]
    callout_style.font.name = "Calibri"
    callout_style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    callout_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    callout_style.font.size = Pt(10.5)
    callout_style.font.color.rgb = rgb(NAVY)
    callout_style.paragraph_format.left_indent = Inches(0.18)
    callout_style.paragraph_format.right_indent = Inches(0.08)
    callout_style.paragraph_format.space_before = Pt(8)
    callout_style.paragraph_format.space_after = Pt(10)
    callout_style.paragraph_format.line_spacing = 1.2


def configure_page(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_spacing(p, before=0, after=0, line=240)
    run = p.add_run("SCHOLARSHIP PROVIDER HANDBOOK  |  REFERENCE GUIDE")
    set_run_font(run, size=8.5, color=MUTED, bold=True)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=0, after=0, line=240)
    run = p.add_run("Scholarship Finder and Eligibility Platform  |  Page ")
    set_run_font(run, size=8.5, color=MUTED)
    add_page_field(p)


def add_page_field(paragraph):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    set_run_font(run, size=8.5, color=MUTED)
    run._r.extend([begin, instr, separate, text, end])


def add_paragraph(doc, text, bold_lead=None, italic=False, align=None, after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    set_spacing(p, before=0, after=after, line=300)
    p.paragraph_format.widow_control = True
    if bold_lead and text.startswith(bold_lead):
        lead = p.add_run(bold_lead)
        set_run_font(lead, bold=True)
        rest = p.add_run(text[len(bold_lead):])
        set_run_font(rest, italic=italic)
    else:
        run = p.add_run(text)
        set_run_font(run, italic=italic)
    return p


def add_heading(doc, text, level=2):
    p = doc.add_paragraph(text, style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    return p


def add_part(doc, text, subtitle):
    doc.add_page_break()
    p = doc.add_paragraph()
    set_spacing(p, before=26, after=4, line=240)
    run = p.add_run(text.upper())
    set_run_font(run, size=10, color=GOLD, bold=True)
    p = doc.add_paragraph(subtitle, style="Heading 1")
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(14)


def add_callout(doc, label, text, color=NAVY, fill=CALLOUT):
    p = doc.add_paragraph(style="Provider Callout")
    ppr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    ppr.append(shd)
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "22")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), GOLD)
    borders.append(left)
    ppr.append(borders)
    lead = p.add_run(f"{label}: ")
    set_run_font(lead, size=10.5, color=color, bold=True)
    body = p.add_run(text)
    set_run_font(body, size=10.5, color=color)
    return p


def numbering_id(doc, kind):
    numbering = doc.part.numbering_part.element
    num_ids = [int(el.get(qn("w:numId"))) for el in numbering.findall(qn("w:num"))]
    num_id = max(num_ids, default=0) + 1
    abstract_id = 8 if kind == "bullet" else 7
    abstract = next(
        element
        for element in numbering.findall(qn("w:abstractNum"))
        if int(element.get(qn("w:abstractNumId"))) == abstract_id
    )
    lvl = abstract.find(qn("w:lvl"))
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "540")
    tabs.append(tab)
    ppr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    ppr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    ppr.append(spacing)
    existing_ppr = lvl.find(qn("w:pPr"))
    if existing_ppr is not None:
        lvl.remove(existing_ppr)
    lvl.append(ppr)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    if kind == "decimal":
        override = OxmlElement("w:lvlOverride")
        override.set(qn("w:ilvl"), "0")
        start_override = OxmlElement("w:startOverride")
        start_override.set(qn("w:val"), "1")
        override.append(start_override)
        num.append(override)
    numbering.append(num)
    return num_id


def add_list_item(doc, num_id, label, detail=""):
    p = doc.add_paragraph()
    ppr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    ppr.append(num_pr)
    set_spacing(p, before=0, after=4, line=300)
    p.paragraph_format.widow_control = True
    lead = p.add_run(label)
    set_run_font(lead, bold=bool(detail))
    if detail:
        body = p.add_run(f" {detail}")
        set_run_font(body)
    return p


def set_cell_shading(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tcpr = cell._tc.get_or_add_tcPr()
    tc_mar = tcpr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tcpr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    total = sum(widths)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tblpr = table._tbl.tblPr
    tblw = tblpr.find(qn("w:tblW"))
    if tblw is None:
        tblw = OxmlElement("w:tblW")
        tblpr.append(tblw)
    tblw.set(qn("w:w"), str(total))
    tblw.set(qn("w:type"), "dxa")
    tblind = tblpr.find(qn("w:tblInd"))
    if tblind is None:
        tblind = OxmlElement("w:tblInd")
        tblpr.append(tblind)
    tblind.set(qn("w:w"), "120")
    tblind.set(qn("w:type"), "dxa")
    layout = tblpr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblpr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = Inches(widths[index] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            tcpr = cell._tc.get_or_add_tcPr()
            tcw = tcpr.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                tcpr.append(tcw)
            tcw.set(qn("w:w"), str(widths[index]))
            tcw.set(qn("w:type"), "dxa")


def set_row_cant_split(row):
    trpr = row._tr.get_or_add_trPr()
    if trpr.find(qn("w:cantSplit")) is None:
        trpr.append(OxmlElement("w:cantSplit"))


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        set_spacing(p, before=0, after=0, line=260)
        run = p.add_run(header)
        set_run_font(run, size=9.5, color=NAVY, bold=True)
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    set_row_cant_split(table.rows[0])

    for row_values in rows:
        new_row = table.add_row()
        set_row_cant_split(new_row)
        cells = new_row.cells
        for index, value in enumerate(row_values):
            p = cells[index].paragraphs[0]
            set_spacing(p, before=0, after=0, line=260)
            run = p.add_run(value)
            set_run_font(run, size=9.5, color=INK, bold=index == 0)
    set_table_geometry(table, widths)
    p = doc.add_paragraph()
    set_spacing(p, before=0, after=4, line=240)
    return table


def add_cover(doc):
    spacer = doc.add_paragraph()
    set_spacing(spacer, before=0, after=82, line=240)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=0, after=14, line=240)
    run = p.add_run("PROVIDER REFERENCE HANDBOOK")
    set_run_font(run, size=10, color=GOLD, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=0, after=10, line=260)
    run = p.add_run("Scholarship Providers")
    set_run_font(run, size=30, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=0, after=26, line=280)
    run = p.add_run("Why Organizations Fund Learners, How Programs Work, and What Providers Gain")
    set_run_font(run, size=15, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.left_indent = Inches(0.55)
    p.paragraph_format.right_indent = Inches(0.55)
    set_spacing(p, before=0, after=54, line=300)
    run = p.add_run(
        "A comprehensive guide for capstone defense, provider onboarding, platform design, "
        "program governance, stakeholder interviews, and sustainability planning."
    )
    set_run_font(run, size=11, color=MUTED, italic=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=0, after=5, line=240)
    run = p.add_run("Scholarship Finder and Eligibility Platform for Filipino Learners")
    set_run_font(run, size=11, color=NAVY, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=0, after=0, line=240)
    run = p.add_run("Provider-focused reference | August 2026")
    set_run_font(run, size=9.5, color=MUTED)
    doc.add_page_break()


def build_document():
    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    properties = doc.core_properties
    properties.title = "Scholarship Providers: Motivations, Operations, and Platform Value"
    properties.subject = "Comprehensive provider reference for a scholarship finder and eligibility platform"
    properties.author = "Scholarship Finder and Eligibility Platform"
    properties.keywords = "scholarship providers, education access, program design, DSS, provider value, sustainability"
    properties.comments = "Prepared as a capstone and stakeholder reference."

    bullet_id = numbering_id(doc, "bullet")

    add_cover(doc)

    add_heading(doc, "Purpose of this guide", 1)
    add_paragraph(
        doc,
        "This guide explains scholarship providers as organizations with missions, strategies, obligations, constraints, and expected outcomes. It does not treat a provider as merely a source of money. A provider may be a government agency, local government unit, school, company, foundation, nonprofit organization, religious or civic group, professional association, employer, international donor, alumni group, family, or individual. Each type enters scholarship work for different reasons and operates under different rules."
    )
    add_paragraph(
        doc,
        "The guide also explains why a provider would use a shared scholarship platform even when it already has social media pages, an office, or its own website. The answer is not simply visibility. The deeper value is coordinated discovery, structured eligibility checking, reusable applicant records, controlled review, consistent communication, auditable decisions, and evidence of impact."
    )
    add_callout(
        doc,
        "Central idea",
        "A scholarship is the benefit received by a learner; a provider is the organization that converts money, mission, policy, people, and accountability into a fair and workable program. Understanding the provider is therefore essential to understanding the entire scholarship system."
    )

    add_heading(doc, "Contents", 2)
    for label, detail in [
        ("Part I - Provider foundations.", "Who providers are, why they participate, and what value they seek."),
        ("Part II - Program design and operations.", "How providers turn intent and funding into a transparent, manageable scholarship program."),
        ("Part III - Platform value and sustainability.", "Why a shared platform matters, how the DSS should support decisions, and how the system can be funded fairly."),
        ("Part IV - Practical reference.", "Provider scenarios, onboarding questions, success measures, common misconceptions, and future improvements."),
    ]:
        add_list_item(doc, bullet_id, label, detail)
    add_paragraph(
        doc,
        "Scope note: The platform uses the words scholarship and program broadly because providers may offer cash, tuition support, school supplies, devices, transport, accommodation, mentoring, training, internships, research support, examination fees, or combinations of these. Legal, tax, accounting, and safeguarding requirements still depend on the provider and the exact arrangement."
    )

    add_part(doc, "Part I", "Provider Foundations")

    add_heading(doc, "What a scholarship provider is", 2)
    add_paragraph(
        doc,
        "A scholarship provider is the accountable entity that establishes or administers assistance for learners. It decides the purpose of the program, secures or controls resources, defines who may qualify, sets the benefit package, approves selection rules, appoints reviewers, communicates decisions, releases benefits, monitors compliance, and evaluates results. A provider can perform all these functions itself or share them with a funder, school, implementing partner, or platform operator."
    )
    add_paragraph(
        doc,
        "The word provider should not be used carelessly. The entity whose logo appears on a program may not be the entity holding the funds; the funder may not review applications; the school may verify enrollment but not choose awardees; and the platform may organize the workflow without owning the scholarship. Clear roles prevent applicants from being misled and prevent organizations from assuming obligations they did not accept."
    )
    add_table(
        doc,
        ["Role", "Main responsibility", "What the role does not automatically mean"],
        [
            ("Provider", "Owns or administers the program and is accountable for its rules and awards.", "Not necessarily the original source of every peso or benefit."),
            ("Funder or donor", "Supplies money, goods, or other resources.", "Does not automatically control applicant selection or personal data."),
            ("Sponsor", "Supports a program or the shared platform for mission, impact, or partnership value.", "Should not receive unfair ranking or private applicant data in return."),
            ("Implementer", "Runs outreach, intake, review, delivery, or monitoring on behalf of another entity.", "Does not own the funds or final policy unless authorized."),
            ("School or verifier", "Confirms enrollment, grades, identity, or institutional records where lawful.", "Does not automatically endorse or select an applicant."),
            ("Platform operator", "Provides discovery, matching, workflow, communication, and records.", "Does not replace the provider's final human decision."),
        ],
        [1550, 3900, 3910],
    )

    add_heading(doc, "Providers are not limited to large scholarship foundations", 2)
    add_paragraph(
        doc,
        "A provider can operate one local program with five awards or a national portfolio with thousands of beneficiaries. Small organizations are especially important because they often understand needs that national programs cannot see: transportation barriers in one municipality, equipment costs for a specific trade, school supplies for younger learners, fees for licensure examinations, or support for students from a particular community. A good platform should therefore accommodate both simple and complex programs without forcing every provider into the same process."
    )

    provider_types = [
        ("National government agencies", "These providers implement law, public policy, and national development priorities. They may target poverty reduction, teacher supply, agriculture, science and technology, public service, regional equity, or other strategic sectors. Their selection, budgeting, reporting, and audit requirements are usually formal because public funds are involved."),
        ("Local government units", "LGUs may support residents based on local ordinances and community priorities. They often value residency, household need, school attendance, and contribution to local development. Their advantage is proximity to applicants; their challenge is maintaining consistent rules across political terms and preventing patronage."),
        ("Schools and higher education institutions", "Schools may offer entrance scholarships, tuition discounts, grants-in-aid, athletic or performing-arts awards, research support, and retention assistance. Their goals can include access, enrollment, completion, campus diversity, academic excellence, and support for institutional priority programs."),
        ("Private companies and corporate foundations", "Companies may fund education through corporate social responsibility, sustainability programs, community relations, workforce planning, industry development, employee engagement, or long-term talent investment. A corporate scholarship can be philanthropic, strategic, or both, but the rules must not create hidden labor obligations or misleading promises of employment."),
        ("Independent foundations and charitable trusts", "These organizations are created to pursue a continuing social purpose. They may focus on poverty, a geographic area, a field of study, a religious mission, disability inclusion, indigenous communities, disaster recovery, or another cause. Their main concern is converting limited resources into credible, measurable impact over time."),
        ("Nonprofits, civic groups, and religious organizations", "These providers often combine financial help with mentoring, formation, community service, psychosocial support, or local referrals. They may know applicants personally, which improves context but also requires conflict-of-interest rules and consistent documentation."),
        ("Professional associations and industry groups", "Associations may develop future professionals, address shortages in a field, support licensure, strengthen standards, encourage research, or expand representation. Benefits can include exam fees, review support, equipment, internships, and mentorship rather than tuition alone."),
        ("Employers and workforce partnerships", "Employers may support employees, dependents, trainees, or future workers. Their programs may build hard-to-find skills, improve retention, strengthen communities around operating sites, or support reskilling. Any return-service or employment condition must be explicit, reasonable, and lawful."),
        ("Alumni groups and school communities", "Alumni often give because they benefited from education and want to preserve access for later learners. Their programs can build a culture of giving, deepen alumni connection, and provide mentoring networks, but volunteer-led selection still needs structure and continuity."),
        ("International organizations and development partners", "These providers may support education as part of humanitarian response, development cooperation, peacebuilding, inclusion, technical capacity, or regional exchange. They usually require strong monitoring, safeguarding, and evidence that funds reached the intended population."),
        ("Individuals and families", "A person or family may create a scholarship to honor someone, repay a personal opportunity, support a hometown, preserve a legacy, or respond to a direct need. Even a small personal program benefits from written criteria, privacy safeguards, transparent decisions, and a plan for continuity."),
    ]
    for title, text in provider_types:
        add_heading(doc, title, 3)
        add_paragraph(doc, text)

    add_heading(doc, "Why providers give scholarships", 2)
    add_paragraph(
        doc,
        "There is no single provider motive. Most programs combine several motives, and the balance can change over time. A responsible provider should name its motives clearly because purpose determines target applicants, benefits, evidence, selection criteria, renewal rules, and evaluation. Vague motives produce vague programs; vague programs are difficult to explain, defend, and improve."
    )

    motivations = [
        ("Expanding access and reducing inequality", "Education has direct and indirect costs. Even when tuition is low or subsidized, learners may struggle with transport, devices, connectivity, supplies, uniforms, food, accommodation, examinations, or lost household income. Providers fund scholarships to remove one or more of these barriers. Need-based support is not a reward for being poor; it is a mechanism for correcting unequal access to opportunity."),
        ("Fulfilling a public or institutional mission", "Government agencies, schools, foundations, and nonprofits often have an explicit mandate related to education, youth, poverty, science, community development, faith, or public service. A scholarship is one instrument for carrying out that mandate. The provider gains mission fulfillment and evidence that its resources were used for the intended public purpose."),
        ("Building human capital", "Education increases the knowledge and skills available to communities, industries, and public institutions. A provider may support learners in fields where qualified people are scarce, such as science, teaching, health, agriculture, engineering, technology, or skilled trades. The DOST-SEI scholarship model is a clear Philippine example: scholarship funding is used to expand the supply of scientific and technical human resources for research, innovation, and national development."),
        ("Creating a long-term talent pipeline", "A company or professional association may support learners before recruitment because waiting until graduation is too late to address a persistent skills gap. Scholarships can help more learners enter and complete relevant programs, become familiar with an industry, receive mentoring, and develop competencies. This is an indirect and long-term return, not a guarantee that every scholar becomes an employee."),
        ("Strengthening a community or region", "Local providers may want more teachers, nurses, technicians, entrepreneurs, researchers, or professionals to remain connected to their communities. They may prioritize residents, local schools, or applicants willing to serve locally. The intended return is stronger local capability, family stability, and economic participation rather than immediate revenue."),
        ("Supporting merit, talent, leadership, or innovation", "Some providers reward academic performance, artistic or athletic talent, leadership, research potential, entrepreneurship, or technical skill. These programs encourage excellence and help high-potential learners continue developing. Merit criteria should still account for unequal access to preparation and should not assume that grades alone represent potential."),
        ("Improving representation and inclusion", "Providers may intentionally support learners who are underrepresented because of disability, gender, geography, ethnicity, income, displacement, or other structural barriers. The goal is not to lower standards. It is to remove unequal barriers, diversify participation, and ensure that capable people are not excluded by circumstances unrelated to their potential."),
        ("Improving school access, retention, and completion", "An educational institution may use scholarships to attract qualified students, prevent stop-out, encourage continuation in priority programs, or improve completion. This can support both the learner and the institution. Ethical practice requires the school to explain the real cost after aid, renewal conditions, and what happens if the learner changes programs or cannot maintain a threshold."),
        ("Meeting sustainability and social-impact commitments", "Companies increasingly measure social performance alongside financial performance. Education support can contribute to community investment, workforce development, and broader sustainability objectives. Philippine SEC sustainability reporting guidance emphasizes measuring economic, environmental, and social impacts and using non-financial information to improve competitiveness and long-term success. A scholarship can therefore be both social investment and accountable organizational practice."),
        ("Building trust and a social license to operate", "Organizations depend on relationships with communities, employees, schools, regulators, customers, and local leaders. A well-governed scholarship can show long-term commitment to a community. However, trust is earned through fair access, reliable delivery, and respectful treatment; publicity cannot compensate for unclear rules, delayed awards, or misuse of applicant information."),
        ("Engaging employees, members, donors, or alumni", "People are more likely to remain involved when they can see a credible pathway from contributions to learner outcomes. Staff may volunteer as mentors, professionals may guide students, alumni may donate, and members may help with career exposure. The provider gains stronger internal participation and a shared sense of purpose."),
        ("Creating a legacy or honoring a person", "Family and memorial scholarships preserve values, community ties, or the memory of a person. The benefit to the provider is continuity of purpose and a meaningful legacy. Long-term credibility depends on documented governance, successor arrangements, and a budget that can support commitments beyond a single announcement."),
        ("Responding to disruption or urgent need", "Disasters, conflict, displacement, health emergencies, school closures, or sudden household hardship can interrupt education. Providers may create temporary programs to protect continuity. Emergency programs need simplified evidence, rapid decisions, and safeguards against excluding applicants who lost documents or connectivity."),
        ("Pooling resources and multiplying impact", "A provider may collaborate because no single organization can fund, find, assess, support, and monitor every learner alone. Shared programs combine money, outreach, school verification, mentoring, technology, and local knowledge. The return is leverage: each partner achieves more than it could independently."),
        ("Receiving lawful recognition or incentives", "Some arrangements may provide public recognition, partnership value, or tax treatment. For example, the Philippine Adopt-a-School Act encourages qualified private-sector support to public schools and provides defined incentives for covered arrangements. This should be treated as a secondary enabler, not the universal reason for scholarship giving. Not every scholarship qualifies, and providers should obtain legal and tax advice before making claims."),
    ]
    for title, text in motivations:
        add_heading(doc, title, 3)
        add_paragraph(doc, text)

    add_callout(
        doc,
        "Why this matters",
        "A provider does not need to earn direct profit from a scholar for the program to create value. The return may be public impact, mission achievement, stronger communities, a future skills base, institutional trust, employee engagement, credible sustainability outcomes, or a lasting legacy."
    )

    add_heading(doc, "What providers receive in return", 2)
    add_paragraph(
        doc,
        "The phrase return on scholarship should not be limited to money. Providers normally create a portfolio of social, strategic, institutional, and learning returns. A strong program makes these returns visible without turning the learner into a product or promising benefits the scholarship cannot deliver."
    )
    returns = [
        ("Social return.", "More learners remain in education, complete a level, acquire useful skills, improve employability, or contribute to their families and communities."),
        ("Mission return.", "The organization can demonstrate that it fulfilled a public, charitable, educational, religious, professional, or community purpose."),
        ("Talent return.", "The relevant field gains a larger and more diverse pool of trained people, mentors, researchers, professionals, or future employees."),
        ("Institutional return.", "Schools can improve access and persistence; associations can strengthen professions; LGUs can build local capability; foundations can deepen program learning."),
        ("Relationship return.", "The provider builds credible connections with schools, communities, local leaders, alumni, employees, professional networks, and partner organizations."),
        ("Reputation return.", "A transparent and reliable program can strengthen trust. This is earned value, not purchased praise, and it can be lost quickly through unfair selection or poor delivery."),
        ("Evidence return.", "Structured records allow the provider to explain who was reached, how decisions were made, what benefits were delivered, and what outcomes followed."),
        ("Organizational learning.", "Application patterns reveal unmet needs, geographic gaps, document barriers, program demand, and reasons otherwise qualified learners are excluded."),
        ("Financial or tax return in limited cases.", "A provider may lawfully recover administration costs, receive grant funding, or qualify for incentives under a specific arrangement. These possibilities require separate legal and accounting review and should never be assumed."),
    ]
    for label, detail in returns:
        add_list_item(doc, bullet_id, label, detail)

    add_heading(doc, "What providers should not expect", 3)
    for label, detail in [
        ("Ownership of the learner.", "Support does not give the provider unrestricted control over a student's personal choices, identity, records, or future."),
        ("Guaranteed employment or loyalty.", "A scholarship may create goodwill and a talent connection, but employment remains a separate decision unless a lawful and clearly disclosed agreement applies."),
        ("Unlimited access to personal data.", "Providers should receive only the information necessary for the published program and only through authorized roles."),
        ("Pay-to-win visibility.", "Paying the platform should not allow a provider to hide competitors, manipulate match scores, bypass verification, or receive preferential approval."),
        ("Publicity without accountability.", "A logo, announcement, or photo is not proof of impact. The provider must deliver promised benefits and report results honestly."),
    ]:
        add_list_item(doc, bullet_id, label, detail)

    add_table(
        doc,
        ["Provider type", "Typical primary motive", "Common organizational value"],
        [
            ("Government or LGU", "Equity, public mandate, regional or sector development", "Public outcomes, policy delivery, accountable use of funds"),
            ("School or university", "Access, enrollment, retention, academic development", "Student success, mission, program vitality"),
            ("Company", "Community impact, workforce, sustainability, trust", "Talent ecosystem, relationships, measurable social value"),
            ("Foundation or nonprofit", "Mission, equity, targeted social change", "Demonstrated impact, donor confidence, program learning"),
            ("Association or employer", "Professional pipeline and skills", "Stronger field, qualified talent, member engagement"),
            ("Alumni, family, or individual", "Giving back, legacy, local connection", "Continuity of values and visible learner outcomes"),
        ],
        [1900, 3830, 3630],
    )

    add_part(doc, "Part II", "Program Design and Operations")

    add_heading(doc, "A scholarship program begins with a problem, not a form", 2)
    add_paragraph(
        doc,
        "Before asking applicants for documents, the provider should define the problem it is trying to solve. Examples include learners leaving school because of transport costs, insufficient STEM graduates in a region, low participation of students with disabilities, teachers unable to pursue graduate study, or talented students lacking equipment. The application form should collect only the information needed to identify and address that problem."
    )
    add_paragraph(
        doc,
        "A useful program logic is: resources are invested; activities are delivered; qualified learners receive benefits; barriers are reduced; education or training continues; and longer-term outcomes become possible. The provider should distinguish what it controls, such as timely release and mentoring, from what it only influences, such as graduation, employment, or community development."
    )

    add_heading(doc, "Core design decisions", 2)
    decisions = [
        ("Define the purpose.", "State the barrier, opportunity, or capability the program addresses and the population it intends to reach."),
        ("Identify the target applicants.", "Choose education levels, locations, schools, fields, income groups, talents, circumstances, or other criteria that directly support the purpose."),
        ("Set the number and duration of awards.", "Align slots with confirmed funding, staff capacity, monitoring requirements, and possible renewals."),
        ("Design the benefit package.", "Match the benefit to the actual barrier. Tuition alone does not solve transport, device, food, or accommodation problems."),
        ("Define evidence and requirements.", "Ask only for documents that prove an eligibility rule or support a fair decision. Avoid duplicate or provider-irrelevant files."),
        ("Choose the selection process.", "Use document review, scoring, exam, interview, portfolio, audition, or community validation only when each stage has a clear purpose."),
        ("Set renewal and completion rules.", "Explain academic expectations, reporting, changes in circumstances, leaves, transfers, and the consequences of non-compliance."),
        ("Plan communication and delivery.", "Name contact channels, decision dates, schedule announcements, disbursement methods, and escalation procedures."),
        ("Define outcomes and indicators.", "Decide what success means before applications open, not after the provider needs a report."),
        ("Approve governance and safeguards.", "Assign decision authority, conflict rules, privacy roles, retention periods, complaint handling, and protection for minors."),
    ]
    decision_number_id = numbering_id(doc, "decimal")
    for label, detail in decisions:
        add_list_item(doc, decision_number_id, label, detail)

    add_heading(doc, "Benefits are not limited to an award amount", 2)
    add_paragraph(
        doc,
        "A scholarship platform should not assume that every program is a cash award. Some of the most useful programs combine financial and non-financial support. The provider should describe each benefit separately, including amount or quantity, frequency, duration, recipient, delivery method, conditions, and whether it is guaranteed or subject to actual cost."
    )
    benefit_groups = [
        ("Direct education costs", "Tuition, school fees, laboratory fees, books, uniforms, learning materials, review fees, certification fees, and research expenses."),
        ("Living and participation support", "Transport, meals, accommodation, monthly allowance, connectivity, child-care support, and accessibility-related costs."),
        ("Equipment and technology", "Laptop, tablet, calculator, tools, safety equipment, internet device, data allowance, or course-specific equipment."),
        ("Academic and professional development", "Tutoring, mentoring, coaching, internships, workshops, licensure review, language support, conference participation, and research supervision."),
        ("Well-being and continuity support", "Counseling, emergency assistance, health support, psychosocial services, disability accommodations, and referrals."),
        ("Opportunity and network access", "Introductions to professionals, alumni networks, work exposure, volunteer opportunities, leadership development, and community projects."),
    ]
    for label, detail in benefit_groups:
        add_list_item(doc, bullet_id, f"{label}.", detail)
    add_callout(
        doc,
        "Design rule",
        "The benefit should follow the barrier. If a learner's main barrier is distance, a transport or housing benefit may be more effective than a small one-time cash grant. If the goal is workforce capability, mentoring, equipment, internships, and certification support may matter as much as tuition."
    )

    add_heading(doc, "Eligibility and targeting", 2)
    add_paragraph(
        doc,
        "Eligibility rules define who may enter the selection process. Selection rules define who receives an award when eligible applicants exceed available slots. Mixing these concepts causes confusion. A learner can be eligible but not selected because capacity is limited; a high score cannot make an ineligible learner eligible unless the provider changes the published rule."
    )
    selection_number_id = numbering_id(doc, "decimal")
    for label, detail in [
        ("Need-based targeting.", "Uses household income, vulnerability, dependents, displacement, or documented financial barriers. Income should be interpreted with household size, location, and actual education costs where feasible."),
        ("Merit-based targeting.", "Uses grades, achievements, assessment results, portfolios, talent, or leadership. Providers should state the grading scale and allow equivalent evidence across school systems."),
        ("Field-based targeting.", "Limits support to courses, strands, trades, research topics, or professional areas connected to the provider's purpose."),
        ("Geographic targeting.", "Prioritizes residents, schools, service areas, rural locations, or communities. The rule should explain whether location is strict eligibility, a preference, or simply the delivery site."),
        ("Identity or circumstance targeting.", "Supports groups facing documented barriers, such as persons with disabilities, indigenous learners, solo-parent households, displaced families, or dependents of a defined group. Sensitive data require strong necessity and privacy review."),
        ("Open or broad eligibility.", "Keeps criteria minimal when the provider's purpose is wide access. Open criteria should not reduce a match score merely because an applicant has a specific course or background."),
    ]:
        add_list_item(doc, bullet_id, label, detail)

    add_heading(doc, "Applicant data: what providers need and why", 2)
    add_paragraph(
        doc,
        "Providers need enough information to verify eligibility, compare applicants fairly, contact selected learners, deliver benefits, and report results. They do not need every possible fact at the start. Collection should be progressive: basic matching data for discovery, necessary evidence for application, additional details for finalists, and payment or contract data only after selection."
    )
    add_table(
        doc,
        ["Data category", "Legitimate program purpose", "Minimum-design rule"],
        [
            ("Identity and contact", "Create a reliable record and communicate decisions.", "Collect accurate contact details; avoid unrelated identifiers."),
            ("Age and guardian", "Protect minors and identify an authorized adult contact.", "Collect guardian data only when age or account context requires it."),
            ("Education level and school", "Match the learner to the correct program form and confirm current status.", "Use level-specific fields; do not ask elementary learners for college-only data."),
            ("Grade or academic result", "Apply a published academic rule or support merit review.", "Store the grading scale and accept equivalent records, not one assumed scale."),
            ("Course, strand, or training", "Match field-specific opportunities.", "Use standardized choices plus a controlled option for valid unlisted programs."),
            ("Household and income", "Assess financial need or grant-in-aid eligibility.", "Request only when need affects eligibility or priority; explain acceptable proof."),
            ("Location", "Apply residency or coverage rules and estimate practical access.", "Use address for usability; protect precise coordinates and avoid unnecessary tracking."),
            ("Documents", "Verify claims and prepare application requirements.", "Reuse common documents and request provider-specific files only in the relevant application."),
            ("Preferences and support needs", "Improve discovery and understand barriers.", "Treat preferences as guidance unless explicitly part of selection."),
            ("Sensitive circumstances", "Provide targeted support or accommodations.", "Collect only with clear necessity, consent or another lawful basis, and restricted access."),
        ],
        [2020, 3760, 3580],
    )
    add_paragraph(
        doc,
        "In the Philippines, education records can be sensitive personal information. The Data Privacy Act requires transparency, legitimate purpose, and proportionality. In practical terms, the provider and platform should tell applicants what is collected, why, who may view it, how long it is retained, and how a privacy concern can be raised. Access should follow role and application ownership rather than convenience."
    )

    add_heading(doc, "Selection and review", 2)
    add_paragraph(
        doc,
        "A defensible selection process separates evidence from judgment. Reviewers first verify whether required facts and documents are present and credible. They then apply published criteria consistently. A rubric is useful when reviewers must compare qualitative evidence, but it should not create false precision. Scores should support discussion and documentation, not replace responsible human review."
    )
    for label, detail in [
        ("Administrative review.", "Checks completeness, duplicates, deadlines, consent, and basic program eligibility."),
        ("Document review.", "Checks whether each file supports the claim it is meant to prove, with accepted, rejected, or replacement-requested outcomes."),
        ("Academic or eligibility review.", "Applies grade, course, level, location, income, and other published rules."),
        ("Exam, interview, portfolio, or audition.", "Adds evidence only when the program purpose requires it. Detailed links can remain private until the learner reaches that stage."),
        ("Final decision.", "Balances evidence, rubric results, available slots, conflicts, and program priorities. The provider remains responsible for approval or rejection."),
        ("Feedback and record.", "Records the reason for adverse decisions in language that is useful, respectful, and consistent with the published process."),
    ]:
        add_list_item(doc, selection_number_id, label, detail)
    add_paragraph(
        doc,
        "Reviewers should not see the approval or rejection control before they can see the evidence required to decide. A clean workspace should present applicant identity and verification status, program fit, eligibility findings, required documents, profile evidence where authorized, scoring or notes, conflict declarations, and then the final action."
    )

    add_heading(doc, "Renewal, contracts, and service obligations", 2)
    add_paragraph(
        doc,
        "Some scholarships end after one distribution; others continue for several terms or years. Renewal rules may include enrollment, academic standing, progress reports, attendance, approved program changes, or continued need. Providers should distinguish automatic renewal, conditional renewal, and reapplication. They should also explain what happens during illness, leave, school transfer, curriculum change, or temporary academic difficulty."
    )
    add_paragraph(
        doc,
        "Contracts can document benefits, responsibilities, privacy, release schedules, conduct, reporting, and dispute procedures. A service obligation may be appropriate when the program intentionally builds public or sector capacity, as seen in some DOST-supported programs. It must be disclosed before acceptance, proportionate to the support, understandable to the learner, and reviewed for legality. The platform should present possible obligations before application and the final agreement only after selection."
    )
    add_callout(
        doc,
        "Fairness limit",
        "A scholarship agreement should not hide repayment, employment, exclusivity, publicity, or service conditions inside general terms. Material obligations must be prominent, specific, and accepted separately at the correct stage."
    )

    add_heading(doc, "Benefit delivery and ongoing support", 2)
    add_paragraph(
        doc,
        "Selection is not the end of provider work. The provider must confirm the award, announce relevant exam or interview results, communicate distribution details, deliver all cash and non-cash benefits, keep transaction or acknowledgment records where appropriate, monitor renewal requirements, and respond to problems. A program that selects awardees but delays or changes benefits without explanation damages applicants and provider credibility."
    )
    for label, detail in [
        ("Award confirmation.", "State the exact benefits, duration, conditions, contacts, and next action."),
        ("Distribution planning.", "Record date, mode, address or online instructions, items to bring, accessibility needs, and who is included."),
        ("Disbursement control.", "Separate authorization, release, acknowledgment, reconciliation, and exception handling."),
        ("Communication.", "Use platform notifications and email for important decisions while keeping an official record inside the system."),
        ("Support and escalation.", "Provide a route for schedule conflicts, missing benefits, document corrections, accessibility requests, and complaints."),
        ("Renewal monitoring.", "Request only the evidence needed for the next period and notify learners before deadlines."),
        ("Completion and alumni connection.", "Close records accurately, measure outcomes, and invite voluntary mentoring or alumni engagement without making it coercive."),
    ]:
        add_list_item(doc, bullet_id, label, detail)

    add_heading(doc, "Governance, ethics, and safeguarding", 2)
    add_paragraph(
        doc,
        "Scholarship decisions affect education, dignity, and trust. Governance is therefore not an administrative extra. It is the structure that makes a program fair when staff change, demand exceeds supply, applicants know reviewers, data are sensitive, or public attention is high."
    )
    governance_items = [
        ("Clear authority.", "Identify who drafts criteria, verifies providers, reviews applications, approves awards, releases funds, handles appeals, and audits changes."),
        ("Conflict management.", "Require reviewers to disclose family, employment, school, financial, or personal relationships and recuse when impartiality is doubtful."),
        ("Privacy and security.", "Restrict records by role, protect document routes, log access and changes, use secure storage, and define deletion or retention schedules."),
        ("Safeguarding younger applicants.", "Use age-appropriate language, guardian support where needed, controlled communication, and limits on unnecessary direct contact or publicity."),
        ("Accessibility and inclusion.", "Make forms usable on small screens, allow alternative evidence, provide accommodations, avoid color-only meaning, and support applicants with limited connectivity."),
        ("Transparent criteria.", "Publish eligibility, selection stages, benefits, capacity, deadlines, obligations, and reasons a decision may change."),
        ("Human oversight.", "Keep the provider's final decision accountable to named people even when the system provides scores or recommendations."),
        ("Complaint and correction routes.", "Allow applicants to report program issues, correct inaccurate data, and request review of process errors without promising automatic reversal."),
        ("Auditability.", "Record versions of criteria, documents, scores, status changes, notices, schedules, and decisions so the process can be reconstructed."),
    ]
    for label, detail in governance_items:
        add_list_item(doc, bullet_id, label, detail)

    add_heading(doc, "Common operational risks and suitable controls", 2)
    add_table(
        doc,
        ["Risk", "What can go wrong", "Practical control"],
        [
            ("Unfunded promise", "More awards are published than the provider can deliver.", "Verify budget and slots before publication; lock awards already committed."),
            ("Biased review", "Personal relationships or subjective preferences affect decisions.", "Conflict disclosure, rubrics, multiple reviewers, reasons, and audit logs."),
            ("Fraud or altered proof", "Documents or claims are false or reused improperly.", "Secure viewing, verification status, replacement history, and selective source checks."),
            ("Excessive data", "The provider collects information unrelated to the program.", "Purpose review, stage-based forms, role restrictions, and retention rules."),
            ("Process overload", "Reviewers cannot manage large applicant volumes.", "Compact queues, filters, readiness indicators, bulk actions with safeguards, and assignment."),
            ("Missed communication", "Applicants do not see schedules or decisions.", "In-platform records plus email, clear dates, and resend or escalation procedures."),
            ("Criteria drift", "Rules change after learners apply.", "Version and snapshot criteria; return material program changes to admin review."),
            ("Stage inconsistency", "A learner is advanced before documents or attendance are complete.", "Workflow gates based on configured stages and accepted evidence."),
            ("Benefit delay", "Approved learners wait without an explanation.", "Distribution planning, responsible owner, release status, exception notes, and reporting."),
            ("Small-provider discontinuity", "A volunteer or founder leaves and records are lost.", "Shared account roles, written procedures, successor access, and exportable records."),
        ],
        [1700, 3830, 3830],
    )

    add_part(doc, "Part III", "Platform Value and Sustainability")

    add_heading(doc, "Why providers would use a shared platform", 2)
    add_paragraph(
        doc,
        "A provider's website and social media remain useful for branding, announcements, and direct relationships. They are not designed to solve the entire scholarship market. Applicants must already know the provider, interpret unstructured posts, compare different rules, repeat information, track deadlines manually, and trust that an account is legitimate. Providers then receive inconsistent emails, messages, spreadsheets, and file links that are difficult to review and audit."
    )
    add_paragraph(
        doc,
        "The shared platform should complement, not replace, provider channels. It gives providers a common operating layer while preserving their identity, eligibility rules, benefit package, final decision, and external exam or interview process. The strongest reason to use it is coordinated administration across the full applicant journey, not simply another place to post an announcement."
    )
    add_table(
        doc,
        ["Provider problem", "Typical fragmented approach", "Shared-platform value"],
        [
            ("Hard-to-reach eligible learners", "Separate posts and referrals", "Searchable catalog, profile-based discovery, and location-aware matching"),
            ("Repeated applicant questions", "Private messages and email", "Structured details, benefits, criteria, requirements, schedules, and notices"),
            ("Inconsistent applications", "Different forms and attachments", "Level-aware profile, reusable documents, and provider-specific application steps"),
            ("Large review workload", "Spreadsheets and folders", "Filtered queues, assignments, readiness indicators, document review, and bulk-safe actions"),
            ("Difficulty explaining decisions", "Scattered notes", "Eligibility findings, rubric records, status history, and decision reasons"),
            ("Schedule confusion", "Individual messages", "Program-level exam, interview, and distribution announcements applied to eligible participants"),
            ("Weak accountability", "No complete activity history", "Role-based access, notifications, email records, logs, and versioned decisions"),
            ("Poor outcome evidence", "Counts prepared manually", "Consistent process and outcome data suitable for program learning"),
        ],
        [2050, 3440, 3870],
    )

    add_heading(doc, "Provider journey in the platform", 2)
    platform_steps = [
        ("Register and verify the organization.", "The provider creates an account, confirms email, completes its organization profile, uploads proof, and waits for admin verification before publishing."),
        ("Create a program as a draft.", "The provider adds the title, description, target learners, benefits, criteria, requirements, application period, contact details, process stages, schedules, possible obligations, capacity, and logo."),
        ("Use DSS guidance while designing criteria.", "The form explains how fields affect applicant matching and warns against unnecessarily narrow or inconsistent rules."),
        ("Submit the program for admin review.", "The administrator checks provider status, program completeness, clarity, dates, requirements, benefits, and practical terms before publication."),
        ("Publish and receive applications.", "Eligible learners discover the program, view details and schedules, prepare required files, accept application terms, and submit."),
        ("Review applicants and evidence.", "Authorized provider reviewers see the applicant profile, verification status, fit findings, required documents, optional rubric, notes, and history."),
        ("Approve or reject the current decision stage.", "The provider makes a human decision after eligibility and documents are checked. If the program includes an exam or interview, approved learners proceed to that configured stage."),
        ("Manage shared schedules.", "The provider posts exam, interview, or distribution details once for the program and applies them to applicants who have reached the relevant stage."),
        ("Record outcomes and delivery.", "The provider records attendance or results when needed, confirms final approval, and completes benefit distribution for the intended group."),
        ("Monitor, resolve reports, and close the cycle.", "The provider handles program-related reports, communicates changes, preserves records, evaluates outcomes, and improves the next cycle."),
    ]
    journey_number_id = numbering_id(doc, "decimal")
    for label, detail in platform_steps:
        add_list_item(doc, journey_number_id, label, detail)

    add_heading(doc, "What the Decision Support System should do", 2)
    add_paragraph(
        doc,
        "The DSS is an advisory layer that compares an applicant's saved profile with the provider's published criteria. It can calculate fit, identify matched criteria, show gaps, distinguish strict eligibility from preferences, and help applicants decide where to invest time. It can also help providers design criteria consistently and focus manual review on real uncertainties."
    )
    add_paragraph(
        doc,
        "The DSS should not choose awardees. It cannot fully evaluate authenticity, context, resilience, interview performance, portfolio quality, or exceptional circumstances. The provider makes the final decision and remains accountable for it. A useful DSS is transparent enough to explain why a result occurred, tolerant of equivalent course names and grading scales, and careful not to penalize applicants when a program accepts any course, level, income, or location."
    )
    for label, detail in [
        ("Separate eligibility from suitability.", "A strict failed rule can block an application; a preference can lower fit without declaring the learner ineligible."),
        ("Treat open criteria as matched.", "Any course, any income, or nationwide coverage should not create a mismatch."),
        ("Normalize equivalent values.", "Examples include BSIT and BS Information Technology, percentage and GWA conventions, and grade-level labels."),
        ("Show evidence gaps.", "Tell applicants which profile field or document is missing rather than showing only a percentage."),
        ("Preserve decision snapshots.", "The system should record the criteria and profile used at submission so later edits do not rewrite history."),
        ("Require human review for uncertainty.", "Conflicting scales, unusual school systems, or sensitive circumstances should be flagged rather than guessed."),
    ]:
        add_list_item(doc, bullet_id, label, detail)

    add_heading(doc, "Why provider verification matters", 2)
    add_paragraph(
        doc,
        "Applicants may submit identity records, grades, income proof, and other sensitive files. The platform cannot responsibly allow any newly created account to publish a program and collect those records. Provider verification establishes that an organization or responsible person exists, supplied suitable proof, has authorized contacts, and passed an administrative review. It does not guarantee that every future action is perfect, so continued logs, reporting, suspension controls, and program review remain necessary."
    )
    add_paragraph(
        doc,
        "Verification also benefits legitimate providers. It separates them from impersonators, improves applicant confidence, reduces repeated questions about legitimacy, and gives platform administrators a clear contact for complaints or changes. Provider staff accounts should inherit the organization boundary and receive only the permissions needed for their role."
    )

    add_heading(doc, "A suitable sustainability model", 2)
    add_paragraph(
        doc,
        "The strongest sustainability model for this platform is sponsor-supported core access with optional paid services, institutional partnerships, and grant or public-interest funding. Core discovery, provider verification, program posting, matching, application submission, and fair review should remain accessible without a mandatory provider subscription whenever feasible. This supports inclusion and keeps the platform useful to small foundations, LGUs, schools, community groups, and new programs."
    )

    add_heading(doc, "Why mandatory provider subscriptions are a weak default", 3)
    for label, detail in [
        ("Providers already fund the public benefit.", "A mandatory recurring fee reduces the resources available for learners and can feel like charging an organization for giving."),
        ("Small providers are often the most price-sensitive.", "A subscription may exclude local programs that add geographic reach and applicant choice."),
        ("The platform needs broad supply.", "A useful finder depends on many legitimate opportunities; high posting barriers make the catalog incomplete."),
        ("Subscriptions can encourage adverse selection.", "Only well-funded or commercially motivated providers may remain, while community programs return to fragmented channels."),
        ("Paying must not influence trust.", "If payment appears to buy ranking, approval, or applicant access, match credibility and admin neutrality are weakened."),
        ("Value may be seasonal.", "Many providers operate one cycle per year, making a continuous subscription difficult to justify."),
    ]:
        add_list_item(doc, bullet_id, label, detail)

    add_heading(doc, "Why institutions should sponsor the platform", 3)
    add_paragraph(
        doc,
        "A platform sponsor funds shared infrastructure rather than one award. This can be attractive to universities, companies, foundations, banks, telecommunications firms, technology companies, professional associations, development partners, and local institutions because one investment improves access, coordination, privacy controls, and evidence across many providers and learners. Sponsorship can therefore produce sector-wide value that a single scholarship cannot."
    )
    sponsor_reasons = [
        ("Scalable education impact.", "The sponsor helps many learners find appropriate opportunities, including programs the sponsor does not directly fund."),
        ("Reduced duplication.", "Shared identity, profile, document, communication, and workflow tools lower repeated administrative effort across institutions."),
        ("Better ecosystem evidence.", "Aggregated and privacy-safe data can reveal unmet demand, common barriers, underserved locations, and document problems."),
        ("Mission and sustainability alignment.", "Education access, inclusion, digital public infrastructure, workforce development, and community resilience can support institutional social goals."),
        ("Credible partnership visibility.", "The sponsor can receive transparent recognition for enabling the system without controlling applicant rankings or decisions."),
        ("Innovation and capability building.", "The platform demonstrates practical use of role-based access, structured workflows, maps, notifications, and decision support for a public-interest problem."),
        ("Long-term relationships.", "Sponsorship connects institutions with schools, providers, LGUs, associations, and communities in a neutral operating environment."),
        ("Research and program learning.", "With proper governance, anonymized trends can support program improvement, policy discussion, and future investment decisions."),
    ]
    for label, detail in sponsor_reasons:
        add_list_item(doc, bullet_id, label, detail)

    add_heading(doc, "Revenue and funding sources that preserve fairness", 3)
    for label, detail in [
        ("Institutional sponsorship.", "Annual support for infrastructure, outreach, accessibility, security, or geographic expansion with clear recognition rules."),
        ("Government or university partnership.", "Service agreements or public-interest funding for deployment, onboarding, verification, and program coordination."),
        ("Foundation and development grants.", "Funding for access, inclusion, research, safeguarding, rural reach, or capacity building."),
        ("Optional provider services.", "Paid onboarding, data migration, branded campaign support, staff training, priority technical support, custom exports, or integrations. Core ranking and approval remain unaffected."),
        ("Managed administration.", "A provider may pay for the platform team to help configure a program or coordinate a cycle, with the service and cost disclosed."),
        ("Ethical transaction support.", "A convenience or payment-processing charge may be passed to the purchasing organization only when transparent, proportionate, and not deducted unexpectedly from learner benefits."),
        ("Research and training projects.", "Funded studies, workshops, and provider-capacity programs can support maintenance while improving practice."),
        ("In-kind support.", "Hosting, email, security review, mapping, accessibility testing, or communications support can reduce operating costs."),
    ]:
        add_list_item(doc, bullet_id, label, detail)
    add_callout(
        doc,
        "Sustainability guardrail",
        "No revenue source should allow a sponsor or paying provider to see unauthorized personal data, bypass verification, alter DSS logic in its favor, suppress competing programs, or move ahead of qualified applicants. Financial sustainability must protect platform neutrality."
    )

    add_heading(doc, "Measuring provider and platform success", 2)
    add_paragraph(
        doc,
        "Counting applications is not enough. A high application count can indicate strong reach, unclear targeting, or an unnecessarily difficult process. Providers should use a balanced measurement framework that covers resources, process quality, outputs, outcomes, equity, and risk. Measures should be interpreted with context and never encourage staff to accept unqualified applicants merely to improve a number."
    )
    add_table(
        doc,
        ["Measurement level", "Useful examples", "Question answered"],
        [
            ("Inputs", "Budget, staff time, benefits, slots, partner contributions", "What resources were committed?"),
            ("Reach", "Views, eligible profiles, applications by level and location", "Who discovered and entered the process?"),
            ("Process", "Completion rate, review time, replacement rate, response time", "Was the process understandable and manageable?"),
            ("Selection", "Eligible-to-award rate, reasons for rejection, unused slots", "Did criteria and capacity work as intended?"),
            ("Delivery", "Benefits released, schedule completion, delays, exceptions", "Did the provider deliver what was promised?"),
            ("Education outcome", "Enrollment continuity, renewal, progression, completion", "Did support help the learner continue or succeed?"),
            ("Capability outcome", "Certification, licensure, research, employment, service", "Did the program build the intended capability?"),
            ("Equity", "Reach and outcomes across location, income, level, disability, or other relevant groups", "Were intended barriers reduced fairly?"),
            ("Trust and safety", "Complaints, privacy incidents, appeals, conflicts, corrections", "Was the program governed responsibly?"),
            ("Learning", "Criteria changes, benefit changes, repeated gaps, partner actions", "What should improve in the next cycle?"),
        ],
        [1760, 4420, 3180],
    )

    add_part(doc, "Part IV", "Practical Provider Reference")

    add_heading(doc, "Questions every provider should answer before publishing", 2)
    questions = [
        ("Purpose", "What problem are we solving, for whom, and why are we the right organization to address it?"),
        ("Authority", "Who owns the program, controls the funds, approves decisions, and can change the rules?"),
        ("Funding", "Are all announced slots and benefits secured, including possible renewals and administration costs?"),
        ("Target", "Which learners should qualify, and how does every criterion connect to the program purpose?"),
        ("Benefits", "What exactly will each awardee receive, when, for how long, and through which delivery method?"),
        ("Evidence", "What is the minimum proof needed for each rule, and can common documents be reused?"),
        ("Selection", "How will eligible applicants be compared if demand exceeds slots?"),
        ("Capacity", "How many applications can staff review fairly within the published timeline?"),
        ("Process stages", "Are an exam, interview, portfolio, or audition genuinely necessary, and what does each add?"),
        ("Obligations", "What renewal, service, conduct, reporting, or repayment terms could apply?"),
        ("Privacy", "Who may view each data type, how long will it be kept, and how can an applicant raise a concern?"),
        ("Safeguarding", "How will the program protect minors and applicants in vulnerable circumstances?"),
        ("Communication", "Who answers questions, sends decisions, announces schedules, and handles delays?"),
        ("Complaints", "How will reports, conflicts, corrections, and process errors be handled?"),
        ("Measurement", "Which results will show that the program solved the stated problem rather than merely distributing funds?"),
        ("Continuity", "What happens if a reviewer leaves, funding changes, or the next cycle must be delayed?"),
    ]
    questions_number_id = numbering_id(doc, "decimal")
    for label, detail in questions:
        add_list_item(doc, questions_number_id, f"{label}.", detail)

    add_heading(doc, "Illustrative provider scenarios", 2)

    scenarios = [
        ("Community foundation supporting basic education", "A local foundation observes that elementary and junior-high learners miss school because of transport and supplies. It targets residents of several barangays, asks for enrollment and a simple need profile, and offers transport support, school supplies, and guardian communication. It measures attendance continuity and renewal rather than requiring a college-style GWA form. The platform helps it reach eligible families, use age-appropriate profiles, verify documents, and schedule distribution."),
        ("Small company building a technical talent pathway", "A regional technology company needs more entry-level network and software talent. It supports senior-high STEM and TVL learners plus college IT students with devices, certification fees, mentoring, and internships. The scholarship does not promise employment. The company measures completion, certification, internship participation, and voluntary applications for future jobs. The platform standardizes course matching, preserves applicant consent, and keeps final hiring separate."),
        ("LGU multi-level education assistance", "An LGU offers assistance from senior high school through college and TVET. Different levels require different forms and evidence. Residency is strict eligibility, financial need affects priority, and academic standing supports renewal. The platform prevents elementary or TVET applicants from being asked college-only questions, provides map-based coverage context, and gives administrators an auditable queue instead of disconnected spreadsheets."),
        ("Alumni-funded school retention grant", "An alumni group wants to prevent current students from stopping because of temporary hardship. The school confirms enrollment; a small committee reviews a limited need statement; and the alumni group funds one-term grants. The program uses conflict declarations because reviewers may know applicants. The platform provides role boundaries, a consistent record, and continuity when volunteer officers change."),
        ("Professional association licensure support", "An association supports graduating students and recent graduates with review fees, examination fees, mentoring, and professional orientation. Academic records establish readiness, but selection also considers financial need and commitment to the profession. The provider's return is a stronger and more inclusive professional pipeline, not ownership of the licensee's future career."),
        ("Emergency continuity fund", "After a disaster, a nonprofit provides devices, connectivity, transport, and temporary living support. Normal documents may be missing, so it accepts alternative verification and conducts rapid manual review. The platform can identify affected locations and communicate decisions, but the provider relaxes evidence rules to avoid excluding the very learners the program intends to help."),
    ]
    for title, text in scenarios:
        add_heading(doc, title, 3)
        add_paragraph(doc, text)

    add_heading(doc, "Common misconceptions", 2)
    misconceptions = [
        ("A scholarship provider is only a donor.", "Providers also design policy, manage risk, evaluate evidence, deliver benefits, communicate, monitor, and remain accountable."),
        ("Providers give only because they are charitable.", "Programs can combine compassion with public mandate, workforce strategy, community development, institutional mission, sustainability, and long-term capability building."),
        ("A company scholarship is automatically recruitment.", "Some programs build a talent ecosystem, but employment is separate unless openly and lawfully included."),
        ("More requirements make a program more secure.", "Unnecessary requirements increase exclusion and data risk. Security comes from relevant evidence, controlled access, and verification."),
        ("The highest grade always identifies the best applicant.", "Program purpose may also require need, field fit, location, talent, service potential, or contextual review."),
        ("A DSS score is the final decision.", "The score is guidance. The provider makes and explains the final decision."),
        ("Posting on social media is enough.", "Posting creates awareness but does not provide structured matching, document reuse, review controls, history, or outcome evidence."),
        ("A provider must pay a subscription to prove commitment.", "Verification, reliable delivery, and accountable decisions prove commitment. Funding the award already represents a material contribution."),
        ("Tax benefits apply to every scholarship.", "Tax treatment depends on the legal structure and applicable program. Providers must not assume eligibility from a general education-support law."),
        ("Once selected, the provider's work is finished.", "Award confirmation, distribution, support, renewal, issue handling, measurement, and closure are essential parts of the program."),
    ]
    for label, detail in misconceptions:
        add_list_item(doc, bullet_id, label, detail)

    add_heading(doc, "Recommended provider onboarding package", 2)
    for label, detail in [
        ("Organization guide.", "Explain verification, roles, permissions, proof standards, privacy, and support contacts."),
        ("Program design worksheet.", "Translate purpose into target learners, benefits, criteria, documents, stages, capacity, obligations, and indicators."),
        ("Level-specific form guide.", "Show which fields apply to preschool, elementary, junior high, senior high, college, TVET, ALS, and other learners."),
        ("Requirement minimization checklist.", "Require a purpose beside each requested document and remove duplicate or late-stage files from initial application."),
        ("Reviewer handbook.", "Cover evidence review, rubrics, conflicts, decisions, feedback, applicant dignity, and secure document handling."),
        ("Schedule and distribution guide.", "Explain shared program events, participant eligibility, attendance or result tracking, and bulk-safe actions."),
        ("Outcome plan.", "Define success measures, reporting frequency, privacy-safe aggregation, and next-cycle learning."),
        ("Continuity plan.", "Name backup staff, maintain organization-owned access, document funding authority, and preserve records."),
    ]:
        add_list_item(doc, bullet_id, label, detail)

    add_heading(doc, "Future improvements that would add provider value", 2)
    add_paragraph(
        doc,
        "The current platform already covers provider verification, program creation, admin review, applicant matching, document review, scheduling, notifications, email, decision history, reports, role permissions, and benefit packages. Future work should deepen quality and scale rather than simply add more pages."
    )
    future_items = [
        ("Configurable renewal cycles.", "Term-by-term continuation, updated evidence, changed circumstances, deferral, transfer, and renewal decisions."),
        ("Outcome follow-up.", "Privacy-aware surveys and milestones for enrollment continuity, completion, certification, employment, or service."),
        ("Provider portfolio view.", "Cross-program capacity, overlapping criteria, delivery commitments, timeline conflicts, and reusable templates."),
        ("Duplicate and fraud signals.", "Careful detection of repeated files, conflicting claims, or unusual activity with human review and due process."),
        ("School verification partnerships.", "Authorized confirmation of enrollment or records without distributing documents more widely than necessary."),
        ("Accessible alternative evidence.", "Structured affidavits, school confirmations, or assisted intake for applicants who cannot obtain standard documents."),
        ("Data-retention controls.", "Provider-specific schedules, legal holds, deletion workflows, applicant requests, and auditable disposal."),
        ("Interoperability and APIs.", "Secure integration with provider systems and the mobile app without exposing the database directly."),
        ("Provider learning reports.", "Operational and outcome summaries focused on program improvement, not vanity metrics or automated ranking."),
        ("Independent fairness review.", "Periodic testing of DSS mappings, application burden, exclusion patterns, and reviewer consistency."),
        ("Document text extraction with verification.", "OCR can assist data entry, but the uploaded record remains the evidence and an authorized human confirms extracted values."),
        ("Multilingual and assisted access.", "Filipino and local-language guidance, plain-language terms, and supported application for younger or low-connectivity users."),
    ]
    for label, detail in future_items:
        add_list_item(doc, bullet_id, label, detail)

    add_heading(doc, "Conclusion", 2)
    add_paragraph(
        doc,
        "Providers give scholarships because education can advance missions that are larger than a single payment: equal opportunity, public service, local development, professional capacity, innovation, workforce readiness, institutional success, community trust, and intergenerational change. The provider's return is therefore a combination of impact, capability, evidence, relationships, reputation, and learning. Financial benefit may exist in limited lawful forms, but it should not depend on charging applicants, owning their future, or purchasing influence over the platform."
    )
    add_paragraph(
        doc,
        "A shared scholarship finder and eligibility platform is valuable when it reduces fragmentation while preserving provider responsibility. It helps the right learners find opportunities, lets providers state criteria and benefits clearly, supports reusable profiles and documents, organizes review and communication, keeps a decision trail, and produces evidence that can improve later programs. Its DSS guides rather than decides. Its sustainability model should invite providers to participate, protect applicant trust, and fund the shared infrastructure through sponsors, partnerships, grants, optional services, and other neutral sources."
    )
    add_callout(
        doc,
        "Final position",
        "A successful scholarship ecosystem does not ask only how much money was offered. It asks whether the right barrier was identified, the right learners could access the opportunity, decisions were fair, benefits were delivered, personal data were protected, and the program created the outcome the provider intended."
    )

    doc.add_page_break()
    add_heading(doc, "Glossary", 1)
    glossary = [
        ("Applicant", "A person seeking consideration for a scholarship or education-support program."),
        ("Awardee or scholar", "An applicant selected to receive a defined benefit, subject to the program's accepted terms."),
        ("Benefit package", "The full set of cash and non-cash support promised by a program."),
        ("Decision Support System", "An advisory system that compares profile information with published criteria and explains fit or gaps."),
        ("Eligibility", "The minimum rules an applicant must satisfy to enter or remain in the selection process."),
        ("Funder", "The entity supplying money, goods, or another resource."),
        ("Grant-in-aid", "Assistance commonly directed to eligible learners with financial need, distinct from pure merit or talent scholarship models."),
        ("Implementer", "The entity performing program operations on behalf of a provider or funder."),
        ("Provider", "The accountable entity that owns or administers the scholarship program and its decisions."),
        ("Selection", "The comparative decision among eligible applicants when awards or capacity are limited."),
        ("Sponsor", "An institution supporting the program or platform in exchange for mission, impact, partnership, or recognition value under defined safeguards."),
        ("Student financial assistance", "The broader family of scholarships, grants-in-aid, subsidies, loans, and related education-support mechanisms."),
    ]
    for term, definition in glossary:
        add_paragraph(doc, f"{term}. {definition}", bold_lead=f"{term}.")

    add_heading(doc, "Evidence base and further reading", 1)
    add_paragraph(
        doc,
        "The sources below provide policy and institutional context for the major principles used in this guide. They support the explanation of education access, coordinated financial assistance, human-capital development, private-sector participation, sustainability reporting, and personal-data protection."
    )
    sources = [
        ("Republic Act No. 10687 - UniFAST Act.", "Defines a unified student financial assistance system and links access, efficiency, regional equity, labor-market relevance, quality, sustainability, and national development. https://lawphil.net/statutes/repacts/ra2015/ra_10687_2015.html"),
        ("Republic Act No. 8525 - Adopt-a-School Act.", "Shows Philippine policy support for qualified private-sector participation in education and defines incentives for covered arrangements. It should not be interpreted as applying automatically to every scholarship. https://lawphil.net/statutes/repacts/ra1998/ra_8525_1998.html"),
        ("Republic Act No. 10173 - Data Privacy Act of 2012.", "Establishes privacy protections and the principles of transparency, legitimate purpose, and proportionality relevant to applicant and education records. https://privacy.gov.ph/data-privacy-act/"),
        ("National Privacy Commission - Implementing Rules and Regulations of the Data Privacy Act.", "Provides operational requirements for organizations processing personal data. https://privacy.gov.ph/implementing-rules-regulations-data-privacy-act-2012/"),
        ("United Nations Sustainable Development Goal 4.", "Frames inclusive and equitable quality education, lifelong learning, technical and vocational skills, vulnerable groups, and scholarship expansion as development priorities. https://sdgs.un.org/goals/goal4"),
        ("DOST-SEI Accelerated Science and Technology Human Resource Development Program.", "Illustrates scholarships as a strategy for producing high-level scientific and technical human resources and strengthening research and innovation. https://asthrdp.science-scholarships.ph/"),
        ("DOST-SEI Engineering Research and Development for Technology Program.", "Illustrates scholarship investment in a critical mass of engineers and researchers for inclusive development and industrialization. https://erdt.science-scholarships.ph/"),
        ("Philippine SEC Sustainability Reporting Guidelines.", "Explains social, economic, environmental, and governance-related performance as relevant to competitiveness and long-term success. https://www.sec.gov.ph/wp-content/uploads/2019/10/2019MCNo04.pdf"),
    ]
    for title, detail in sources:
        add_list_item(doc, bullet_id, title, detail)

    add_paragraph(
        doc,
        "Disclaimer: This document is an educational, system-design, and stakeholder reference. It is not legal, tax, accounting, or regulatory advice. Providers should confirm the rules that apply to their organization, funding source, applicants, benefits, contracts, and data-processing activities.",
        italic=True,
        after=0,
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
